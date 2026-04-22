from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
STP_DIR = BASE_DIR / "output" / "stp_B001K85BN2_B001K85BNC_B008KEJ1LM"
CHOICE_DIR = BASE_DIR / "output" / "choice_logit_beard_mustache"
OUT_FILE = BASE_DIR / "output" / "Urbaner_Amazon_Analysis_Report.docx"
CATALOG_PATH = BASE_DIR / "attribute_catalog.csv"


SELECTED_ATTRIBUTES = [
    "rechargeable_design",
    "blade_sharpness",
    "blade_hair_pulling",
    "motor_power_rpm",
    "waterproof_rating_ipx8",
    "adjustable_comb_settings",
    "price_value_ratio",
    "ease_of_use_overall",
]

FONT_NAME = "標楷體"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_hdr = OxmlElement("w:tblHeader")
    tbl_hdr.set(qn("w:val"), "true")
    tr_pr.append(tbl_hdr)


def set_doc_fonts(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def add_para(document: Document, text: str, bold: bool = False, italic: bool = False, align=None) -> None:
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.bold = bold
    run.italic = italic


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_number(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_table(document: Document, df: pd.DataFrame, title: str) -> None:
    document.add_paragraph(title).runs[0].bold = True
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        set_cell_shading(hdr[i], "D9EAF7")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
            for run in cells[i].paragraphs[0].runs:
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def build_attribute_note_table() -> pd.DataFrame:
    catalog = pd.read_csv(CATALOG_PATH).fillna("")
    subset = catalog[catalog["attribute_key"].isin(SELECTED_ATTRIBUTES)].copy()
    subset["sort_order"] = subset["attribute_key"].apply(lambda x: SELECTED_ATTRIBUTES.index(x))
    subset = subset.sort_values("sort_order")

    note = subset[["attribute_key", "label", "label_zh", "definition"]].copy()
    note.columns = ["屬性代碼", "英文屬性", "中文備註", "定義說明"]
    return note


def enrich_coefficient_table(coeff: pd.DataFrame) -> pd.DataFrame:
    catalog = pd.read_csv(CATALOG_PATH).fillna("")
    key_map = catalog.set_index("attribute_key")[["label", "label_zh"]].to_dict("index")

    coeff = coeff.copy()
    coeff["tier"] = coeff["term"].str.extract(r"_tier_(mid|high)", expand=False)
    coeff["attribute_key"] = coeff["term"].str.replace(r"_tier_(mid|high)$", "", regex=True)
    coeff["屬性代碼"] = coeff["attribute_key"]
    coeff["英文屬性"] = coeff["attribute_key"].map(lambda x: key_map.get(x, {}).get("label", x))
    coeff["中文備註"] = coeff["attribute_key"].map(lambda x: key_map.get(x, {}).get("label_zh", ""))
    coeff["水準"] = coeff["tier"].map({"mid": "中階(mid)", "high": "高階(high)"}).fillna("-")
    coeff["方向"] = coeff["direction"].map({"positive": "正向", "negative": "負向"}).fillna(coeff["direction"])
    coeff["係數"] = coeff["coefficient"].round(4)
    coeff["標準誤"] = coeff["std_err"].round(4)
    coeff["z 值"] = coeff["z_value"].round(3)
    coeff["p 值"] = coeff["p_value"].map(lambda x: f"{x:.4g}")
    coeff["Odds Ratio"] = coeff["odds_ratio"].round(3)

    show_cols = [
        "屬性代碼",
        "英文屬性",
        "中文備註",
        "水準",
        "係數",
        "標準誤",
        "z 值",
        "p 值",
        "Odds Ratio",
        "方向",
    ]
    return coeff[show_cols]


def add_picture(document: Document, image_path: Path, caption: str, width_inches: float = 6.3) -> None:
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(width_inches))
        caption_p = document.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.runs[0].italic = True
    else:
        add_para(document, f"[缺少圖檔] {image_path.name}")


def orient_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def main() -> None:
    doc = Document()
    set_doc_fonts(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Urbaner Amazon 評論探勘與聯合分析正式報告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Beard / Mustache Trimmers 類別｜STP、Choice Logit、產品定位與策略建議")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = FONT_NAME
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("產出日期：2026-04-22")
    date_run.font.name = FONT_NAME
    date_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    doc.add_page_break()

    doc.add_heading("1. 研究目的與分析範圍", level=1)
    add_para(doc, "本報告整合 Amazon 顧客評論探勘、STP 分析、聯合分析（Choice / Logit）與產品策略建議，目的是協助 Urbaner 在 Beard / Mustache Trimmers 類別中建立更清晰的產品定位與定價方向。")
    add_para(doc, "分析主軸包含三層：第一層是由評論探勘萃取產品屬性；第二層是透過 STP 找出市場分群、定位與差異化重點；第三層是把屬性轉成 choice data，估計效用、購買機率與產品組合偏好。")

    doc.add_heading("2. 資料來源與分析流程", level=1)
    add_bullet(doc, "評論資料：Amazon 顧客評論，Beard / Mustache Trimmers 類別，共 40 個產品、7,118 筆評論。")
    add_bullet(doc, "STP 來源：review_scoring_table、product_quality_scorecard、positioning_scorecard、segmentation_variables、targeting_anova 等輸出。")
    add_bullet(doc, "Choice / Logit 來源：將每筆評論視為一次購買決策，建立 40 × 7,118 的 long format choice data。")
    add_bullet(doc, "模型方法：以 Conditional Logit 估計屬性 dummy 的部分效用值，並以 logistic 函數轉成購買機率。")
    add_bullet(doc, "策略目標：找出最有價值的屬性、最具吸引力的產品組合，以及可行的產品線與定價策略。")

    doc.add_heading("3. STP 分析摘要", level=1)
    stp_summary = pd.DataFrame(
        [
            ["Themes discovered", 20],
            ["Attributes extracted", 114],
            ["Reviews analysed", 1058],
            ["Segment count", 2],
            ["Silhouette score", 0.475],
            ["Most discriminating feature", "shower_use_safe"],
        ],
        columns=["項目", "結果"],
    )
    add_table(doc, stp_summary, "表 1. STP 分析摘要")
    add_para(doc, "STP 的核心訊號很一致：防水、易用性、鋒利度、價值感與禮品情境，都是市場可被感知且能拉開差異的屬性。尤其 shower_use_safe 的跨產品變異最大，代表它是最強的定位差異點。")
    add_para(doc, "分群結果顯示，主要受眾可概括為兩大群：一群偏重修鬍功能、價格與基礎便利性；另一群則更重視價值與耐用，且在美國市場中禮品購買情境相對突出。")

    doc.add_heading("4. 關鍵屬性中英對照與中文備註", level=1)
    add_para(doc, "為了讓後續策略決策與跨部門溝通更一致，以下整理本次 Choice / Logit 使用之 8 個核心屬性，提供英文屬性名稱、中文備註與定義。")
    add_table(doc, build_attribute_note_table(), "表 2. 聯合分析核心屬性（含中文備註）")

    doc.add_heading("5. 聯合分析（Choice / Logit）完整流程", level=1)
    add_para(doc, "本研究以評論近似真實選擇情境，透過條件羅吉特模型估計 part-worth utility（部分效用）。流程分為八個步驟，以確保資料轉換、估計與解讀具有可追蹤性。")
    add_number(doc, "步驟 1｜評論清理與欄位標準化：整合標題、內文、星等、國家與語言欄位，並排除文字不足的評論。")
    add_number(doc, "步驟 2｜評論屬性評分：依 attribute_catalog 關鍵詞規則，將每篇評論轉為 8 個屬性分數。")
    add_number(doc, "步驟 3｜產品層級彙整：把評論層級分數聚合為產品平均分數，形成產品特徵向量。")
    add_number(doc, "步驟 4｜屬性分層（low / mid / high）：採分位或穩健切分方式，將連續分數轉為離散層級。")
    add_number(doc, "步驟 5｜虛擬變數編碼：以 low 為參照組，建立 mid/high dummy；係數可直接解讀相對於 low 的效用增減。")
    add_number(doc, "步驟 6｜建構 long-format choice data：每筆評論視為一個 choice set，40 個產品同場競爭，實際評論對應產品記為 choice=1。")
    add_number(doc, "步驟 7｜Conditional Logit 估計：以 BFGS 收斂求解，取得屬性係數、標準誤、z 值與 p 值。")
    add_number(doc, "步驟 8｜機率與策略轉換：將效用映射為選擇機率，並連結產品組合、價格與定位建議。")

    choice_metrics = pd.DataFrame(
        [
            ["Choice set size（每次比較產品數）", 40],
            ["Long-format rows（長表資料列數）", 284720],
            ["Estimated reviews（估計樣本評論數）", 2500],
            ["Log-likelihood", -8589.749],
            ["Null log-likelihood", -26257.444],
            ["McFadden pseudo R²", 0.6729],
            ["Top-1 hit rate", 0.0701],
        ],
        columns=["指標", "數值"],
    )
    add_table(doc, choice_metrics, "表 3. Choice / Logit 模型摘要")
    add_para(doc, "McFadden pseudo R² = 0.6729，代表模型相對於隨機基準有顯著解釋力；Top-1 hit rate 受 choice set 規模影響較大，應與 pseudo R² 併讀，不宜單獨判斷模型品質。")

    doc.add_heading("6. 聯合分析係數結果（含中文註解）", level=1)
    add_para(doc, "係數表已補上屬性中英文與中文備註，便於技術團隊、行銷團隊與管理層對齊。方向為正向代表相對於 low 層級可提升選擇機率；方向為負向則代表降低吸引力。")

    coeff = pd.read_csv(CHOICE_DIR / "logit_coefficients.csv")
    coeff_table = enrich_coefficient_table(coeff)
    add_table(doc, coeff_table, "表 4. Conditional Logit 係數表（中英屬性註解）")

    add_para(doc, "重點解讀 1：整體易用性（Ease of Use）高階係數最高，表示產品若能有效降低學習成本，對轉換率的提升最明顯。")
    add_para(doc, "重點解讀 2：性價比、刀片鋒利度與防水能力皆有正向效用，顯示『可感知的實用價值』比單純規格堆疊更重要。")
    add_para(doc, "重點解讀 3：可調式導梳在 mid/high 均為負向，代表該屬性在目前市場語境中存在『功能複雜化』風險，需以更直覺的溝通或結構設計化解。")

    doc.add_heading("7. 產品組合排序與購買機率", level=1)
    prob = pd.read_csv(STP_DIR / "L18_probability_table.csv")
    top5 = prob.sort_values("rank").head(5).copy()
    top5["產品組合"] = top5[["A", "B", "C", "D", "E", "F", "G", "H"]].agg("-".join, axis=1)
    top5 = top5[["rank", "Run", "產品組合", "total_utility", "purchase_prob_raw", "purchase_prob_norm", "est_price_usd"]]
    top5.columns = ["排名", "Run", "產品組合", "總效用", "logistic 機率", "標準化機率", "估計價格(USD)"]
    top5 = top5.round({"總效用": 4, "logistic 機率": 6, "標準化機率": 6, "估計價格(USD)": 1})
    add_table(doc, top5, "表 5. 前 5 名產品組合")
    add_para(doc, "最佳組合是 Run 9：A1-B3-C3-D1-E3-F2-G1-H2。它不是所有屬性都極端拉滿，而是把高權重屬性集中在 USB-C 充電、IPX8 防水、高質感機身與中度功能配置上，因此在效用與市場吸引力上都最均衡。")

    doc.add_heading("8. 產品定位與策略解讀", level=1)
    add_bullet(doc, "Value：低價入門款，適合價格敏感或首次購買者，但不應成為主力形象款。")
    add_bullet(doc, "Mainstream：中價位、均衡型產品，重點是穩定、易用、耐用，適合擴大量體。")
    add_bullet(doc, "Premium：高價高效用產品，是目前最有利的主力定位，適合承接品牌溢價。")
    add_bullet(doc, "Niche：送禮、特殊功能或高控制感需求的小眾產品，可作為補充線。")
    add_para(doc, "從模型結果來看，市場不是單純追求功能數量，而是追求『好用、值得、夠鋒利、夠方便』。因此，最值得投資的是易用性、價值感、刀片鋒利度、防水與 USB-C 充電；最應避免的是過度複雜但不一定有感的調整設計。")
    add_para(doc, "價格策略上，存在 premium pricing 空間，但前提是產品真的具備能被感知的高效用屬性。換言之，價格溢價不是靠故事，而是靠使用體驗與功能差異。")

    doc.add_heading("9. 策略建議", level=1)
    add_number(doc, "建議推出一個 Hero SKU，採 Premium 日常便利型定位：USB-C、IPX8、防拉毛、好清潔、好上手。")
    add_number(doc, "建議推出一個 Performance SKU，強調精準修剪、馬達性能與更高控制感，對重度使用者與高涉入客群有效。")
    add_number(doc, "建議推出一個 Gift SKU，強調禮品包裝、完整配備與節慶送禮情境。")
    add_number(doc, "Entry SKU 可保留，但應扮演流量入口，不宜承擔品牌主形象。")
    add_para(doc, "產品線策略建議採三層：入門引流、中階擴量、高階溢價。主戰場放在中高階的便利型與高效能型，因為目前市場的核心痛點與價值感都集中在這一帶。")

    doc.add_heading("10. 圖表與視覺化（逐圖解讀）", level=1)
    add_para(doc, "以下每張圖均附上圖像閱讀方式、觀察重點與決策含意，方便直接放入簡報報告口條。")

    charts = [
        (STP_DIR / "quality_heatmap.png", "圖 1. 產品屬性品質熱圖"),
        (STP_DIR / "perceptual_map.png", "圖 2. 產品感知定位圖"),
        (STP_DIR / "segmentation_map.png", "圖 3. 顧客分群圖"),
        (STP_DIR / "OD_main_effects.png", "圖 4. 正交設計主效果圖"),
        (STP_DIR / "OD_utility_price_space.png", "圖 5. 效用 × 價格散佈圖"),
        (CHOICE_DIR / "top5_choice_probabilities.png", "圖 6. 前 5 名產品組合購買機率"),
        (CHOICE_DIR / "logit_coefficients_bar.png", "圖 7. Conditional Logit 係數條圖"),
    ]
    chart_notes = {
        "圖 1. 產品屬性品質熱圖": "說明：橫軸為高差異屬性，縱軸為主要競品；顏色越偏綠代表該屬性品質感知越高。管理意涵：可快速辨識 Urbaner 應優先補強的屬性缺口與可主打的相對優勢。",
        "圖 2. 產品感知定位圖": "說明：以主成分將多屬性壓縮為二維定位空間，點位距離越近代表消費者感知越相似。管理意涵：可判斷是否陷入同質競爭，並找出可切入的差異化座標。",
        "圖 3. 顧客分群圖": "說明：以評論屬性顯著度進行分群，顏色區塊代表偏好結構不同的客群。管理意涵：有助於區分主力受眾與次要受眾，避免單一訴求覆蓋全部客群。",
        "圖 4. 正交設計主效果圖": "說明：顯示各因子不同水準對總效用的平均影響幅度。管理意涵：可決定優先投資的規格因子，例如防水與充電型態通常具有較大拉升空間。",
        "圖 5. 效用 × 價格散佈圖": "說明：橫軸為估計價格，縱軸為總效用；右上角代表高價高效用組合。管理意涵：用於篩選高價值 SKU，並判斷 premium 定價是否具備效用支撐。",
        "圖 6. 前 5 名產品組合購買機率": "說明：比較前五名組合的 raw / norm 機率表現。管理意涵：協助先做小規模打樣或 A/B 上架順序，降低新品組合決策風險。",
        "圖 7. Conditional Logit 係數條圖": "說明：條長顯示效用強弱，正值代表提升被選擇機率，負值代表抑制選擇。管理意涵：可直接對應產品文案與主賣點排序，先打高正向係數屬性。",
    }
    for image_path, caption in charts:
        add_picture(doc, image_path, caption)
        add_para(doc, chart_notes.get(caption, ""))
        doc.add_paragraph("")

    doc.add_heading("11. 結論", level=1)
    add_para(doc, "綜合 STP、聯合分析與 choice model 的結果，Urbaner 在 Beard / Mustache Trimmers 類別的最佳路徑不是單純低價競爭，而是以『高使用感的中高階定位』切入。")
    add_para(doc, "真正能創造價值的不是功能數量本身，而是使用者感知到的易用性、價值感、防水便利性、鋒利度與整體體驗。這些元素構成了可持續的產品溢價基礎，也最適合作為品牌主訊息。")

    doc.add_heading("12. 附註", level=1)
    add_para(doc, "本文件由既有分析輸出自動彙整而成，包含 STP、效用分析、購買機率與產品策略建議。若未來要補入真實價格係數，可再把 WTP 換算成絕對金額版的願付價格。")

    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    main()