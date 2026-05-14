"""
產生「儀表板數字統計解讀」Word 文件
按儀表板 7 個頁面順序，對每個統計數字說明：
  1. 這個數字怎麼算出來的
  2. 統計方法叫什麼
  3. 怎麼正確解讀
輸出：output_conjoint/URBANER_儀表板統計解讀.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Microsoft YaHei"
C_BLUE   = RGBColor(0x1F, 0x3A, 0x5F)
C_MID    = RGBColor(0x2E, 0x5C, 0x8A)
C_LIGHT  = RGBColor(0x4A, 0x7A, 0xB0)
C_MUTED  = RGBColor(0x66, 0x66, 0x66)
C_WARN   = RGBColor(0xCC, 0x44, 0x00)
C_GREEN  = RGBColor(0x1A, 0x6B, 0x2A)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_US     = RGBColor(0x2E, 0x5C, 0x8A)
C_JP     = RGBColor(0x9A, 0x31, 0x2F)
C_GOLD   = RGBColor(0x8B, 0x6F, 0x1A)


def _rpr(run):
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    return run._element.rPr


def sr(run, size=10, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    _rpr(run).rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shd(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_color)
    tc.append(s)


def borders(cell):
    tc = cell._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for e in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{e}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "AAAAAA")
        tb.append(b)
    tc.append(tb)


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single")
    bt.set(qn("w:sz"), "6")
    bt.set(qn("w:color"), "CCCCCC")
    pBdr.append(bt)
    pPr.append(pBdr)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    # background shading via XML
    pPr = p._p.get_or_add_pPr()
    shd_el = OxmlElement("w:shd")
    shd_el.set(qn("w:val"), "clear")
    shd_el.set(qn("w:color"), "auto")
    shd_el.set(qn("w:fill"), "1F3A5F")
    pPr.append(shd_el)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"  {text}")
    sr(run, size=14, bold=True, color=C_WHITE)
    return p


def h2(doc, text, color=C_MID):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    sr(run, size=12, bold=True, color=color)
    return p


def h3(doc, text, color=C_LIGHT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    sr(run, size=10, bold=True, color=color)
    return p


def body(doc, text, indent=0, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    sr(run, size=10, italic=italic, color=color or C_MUTED)
    return p


def bullet(doc, text, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"• {text}")
    sr(run, size=10, color=C_MUTED)


def stat_box(doc, label, value, note="", color=C_BLUE):
    """single KPI row: label | value (note)"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Cm(5), Cm(3.5), Cm(8)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
    c0, c1, c2 = table.rows[0].cells
    shd(c0, "EEF3FA"); borders(c0)
    shd(c1, "FFFFFF"); borders(c1)
    shd(c2, "FFFFFF"); borders(c2)
    r0 = c0.paragraphs[0].add_run(label)
    sr(r0, size=10, bold=True, color=color)
    r1 = c1.paragraphs[0].add_run(value)
    sr(r1, size=11, bold=True, color=color)
    r2 = c2.paragraphs[0].add_run(note)
    sr(r2, size=9, italic=True, color=C_MUTED)
    doc.add_paragraph()


def table(doc, headers, rows, hdr_hex="1F3A5F"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        shd(c, hdr_hex); borders(c)
        run = c.paragraphs[0].add_run(h)
        sr(run, size=9, bold=True, color=C_WHITE)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            shd(c, bg); borders(c)
            run = c.paragraphs[0].add_run(str(val))
            sr(run, size=9)
    doc.add_paragraph()
    return t


def method_block(doc, method_text, meaning_text):
    """淺藍底方法說明框"""
    p_m = doc.add_paragraph()
    p_m.paragraph_format.left_indent = Cm(0.4)
    p_m.paragraph_format.space_before = Pt(4)
    p_m.paragraph_format.space_after = Pt(2)
    r1 = p_m.add_run("統計方法：")
    sr(r1, size=9, bold=True, color=C_MID)
    r2 = p_m.add_run(method_text)
    sr(r2, size=9, color=C_MUTED)

    p_e = doc.add_paragraph()
    p_e.paragraph_format.left_indent = Cm(0.4)
    p_e.paragraph_format.space_before = Pt(0)
    p_e.paragraph_format.space_after = Pt(6)
    r3 = p_e.add_run("數字意義：")
    sr(r3, size=9, bold=True, color=C_MID)
    r4 = p_e.add_run(meaning_text)
    sr(r4, size=9, color=C_MUTED)


def interp_block(doc, text):
    """橘底解讀框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("解讀：")
    sr(r1, size=9, bold=True, color=C_WARN)
    r2 = p.add_run(text)
    sr(r2, size=9, italic=True, color=RGBColor(0x44, 0x33, 0x11))


def warn_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(f"⚠ {text}")
    sr(run, size=9, italic=True, color=C_WARN)


# ═══════════════════════════════════════════════════════════
def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ─ 封面 ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("URBANER 儀表板：統計數字完整解讀")
    sr(r, size=20, bold=True, color=C_BLUE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("按儀表板 7 個頁面順序 ｜ 每個數字附方法說明 + 正確解讀方式 ｜ 2026-05-14")
    sr(r2, size=10, italic=True, color=C_MUTED)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("跨境電商競賽答辯版 — 評審追問對照使用")
    sr(r3, size=10, color=C_WARN)

    divider(doc)

    # ─ 閱讀說明 ───────────────────────────────────────────
    h2(doc, "如何使用本文件", color=C_BLUE)
    body(doc,
        "本文件依照儀表板左側導覽列的 7 個頁面順序，逐一解釋每個統計數字與圖表的：\n"
        "① 統計方法（怎麼算的）　② 數字定義（代表什麼）　③ 正確解讀（不能誤解的地方）",
        color=C_MUTED
    )
    table(doc,
        ["頁面", "核心統計方法", "主要數字"],
        [
            ["執行儀表板（總覽）", "MNL + Conjoint + K-Means 彙整", "URBANER 市佔 5.21% / 11.26%、購買率 83.3% / 89.1%"],
            ["雙市場對比", "ANOVA + K-Means + RMS 距離", "F 值、Silhouette、★ 差距"],
            ["STP 策略分析", "K-Means + PCA + ANOVA", "分群佔比、Silhouette、市場定位座標"],
            ["Conjoint 偏好", "Split-model Logistic Conjoint + WTP + MNL", "屬性重要性 % / WTP 金額 / 升級後份額"],
            ["最優產品組合", "Conjoint 設計卡片 + 購買概率 Sigmoid", "83.3% / 89.1% 購買率"],
            ["社群洞察", "評論文本 NLP + 描述統計", "共識條數、品類分布"],
            ["行銷劇本", "洞察摘要（非統計推論）", "文案建議"],
        ]
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 1：執行儀表板（總覽）
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 1｜執行儀表板（總覽）")

    h2(doc, "1-1　總評論數 11,523 則 ｜ 追蹤 SKU 數 84 支")
    method_block(doc,
        "描述統計（Descriptive Statistics）。直接計數資料集中所有評論筆數與 SKU 數量。",
        "反映本分析的資料規模。US 4,360 則 / JP 7,163 則；7 個確認 URBANER ASIN + 77 個競品 SKU。"
    )
    interp_block(doc,
        "這兩個數字是「樣本母體的大小」。評論數越多，後續統計（星等均值、屬性品質分）越穩定。"
        "注意：SKU 數（84）遠小於評論數，代表平均每個 SKU 約有 137 則評論，密度尚可。"
    )

    h2(doc, "1-2　US 最佳組合預估購買率 83.3%　｜　JP 最佳組合預估購買率 89.1%")
    method_block(doc,
        "Logistic Regression 的 Sigmoid 輸出值。"
        "把最優設計卡片（Card 22）的各屬性 part-worth 加總得到效用 U，"
        "再經 Z-score 正規化後代入 Sigmoid：P = 1 / (1 + exp(−Z))。",
        "P > 0.5 代表該產品規格組合「優於市場均值」的相對偏好強度。"
        "這裡的 83.3% / 89.1% 是相對偏好強度，不是「每 100 位消費者有 83 位會買」的市場份額。"
    )
    interp_block(doc,
        "正確解讀：「在所有受評的 SKU 中，這個規格組合的偏好強度排在市場前段班（高於平均）」。"
        "不能說「有 83% 的消費者會選這個產品」——那需要真實的消費者實驗或問卷驗證。"
    )
    warn_block(doc, "購買率是「相對偏好概率」而非「市場轉換率」，答辯時注意區分。")

    h2(doc, "1-3　URBANER 市佔圓餅圖：US 5.21%　｜　JP 11.26%")
    method_block(doc,
        "Multinomial Logit（MNL）Share-of-Preference 模型。"
        "公式：P(i) = exp(U_i) / Σ exp(U_j)，j 為全選擇集所有 SKU。"
        "U_i 由 Conjoint Logit 估出的屬性 part-worth × 各 SKU 的品質分數加總得出。",
        "份額加總 = 100%。US 選擇集 76 個選項（52 SKU + 24 設計卡片），JP 60 個。"
        "URBANER 5.21% 代表：在 76 個選項的競爭中，URBANER SKU 被模型預測合計選到的機率。"
    )
    interp_block(doc,
        "US 5.21% 的直接含義：在與 52 個市場 SKU（含競品）的直接競爭下，"
        "URBANER 的 4 個美國 SKU 合計只能爭到約 1/20 的偏好份額，顯示競爭力仍弱。"
        "JP 11.26% 較好，且有單一 SKU（B0BL2YWH3N）排入全場前 6。"
    )

    h2(doc, "1-4　屬性重要性圖（US 功能合一數 51.5%、JP 前 6 項均在 10% 以上）")
    method_block(doc,
        "Range-based Importance（範圍法重要性）。"
        "公式：importance% = (max(β) − min(β)) / Σ(max(β_k) − min(β_k)) × 100。"
        "β 係數來自 Split-model Binary Logistic Regression（逐屬性分別估計 part-worth）。",
        "重要性 % 越大，該屬性對消費者「買 vs 不買」決策的拉扯力越大。"
    )
    interp_block(doc,
        "US 功能合一數 51.5%：需附加說明——此屬性出現「完全分離（Perfect Separation）」問題，"
        "即偏好=1 的 SKU 與偏好=0 的 SKU 在此屬性上完全可分，導致 β 係數異常大（>60）。"
        "51.5% 方向正確（功能合一確實是美國重要屬性），但幅度高估，答辯時要誠實說明。\n"
        "JP 前 6 屬性均勻（10–17%）：代表日本消費者是「全方位審查」型，沒有壓倒性的單一決策因子。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 2：雙市場對比
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 2｜雙市場對比")

    h2(doc, "2-1　ANOVA F 值圖（哪些屬性最能拉開不同顧客群？）")
    method_block(doc,
        "One-way ANOVA（單因子變異數分析）。"
        "因子（X）：K-Means 分出的顧客群（cluster label）。"
        "依變數（Y）：每個屬性的 Axis B 品質分（0–10）。"
        "對 114 個屬性逐一執行，篩選 p < 0.001 的高度顯著屬性。",
        "F 值 = 組間變異（群與群的均值差距）/ 組內變異（群內個體的離散度）。"
        "F 越大 + p 越小 → 該屬性在不同顧客群之間差異越大 → 「區隔力」越強。"
    )
    interp_block(doc,
        "F 值的直觀意義：若附件件數的 F = 2630（JP Top 1），代表「附件件數分數」在三個顧客群之間的差距，"
        "是同一群內個體差異的 2630 倍——群與群之間分得非常清楚。"
        "US Top 1 是「送禮場景（gift_suitability_men, F=1299.5）」：美國顧客依「要不要送禮」分成截然不同的兩群。"
    )
    warn_block(doc,
        "ANOVA 的 F 值是「區隔力」，不是「重要性」。F 值大代表這個屬性能拉開不同顧客，"
        "不代表消費者整體最在乎這個屬性（那是 Conjoint 重要性的問題）。兩者問不同問題。"
    )

    h2(doc, "2-2　顧客群規模圖（柱 = 人數佔比，金點 = Avg★）")
    method_block(doc,
        "K-Means 分群結果的描述統計。柱高 = 該 cluster 評論則數 ÷ 市場總評論；"
        "金點 = 該 cluster 內所有評論的星等算術平均值。",
        "兩個維度同時看：柱越高 = 族群規模越大；★ 越高 = 整體滿意度越好。"
    )
    interp_block(doc,
        "「大眾痛點群」= 柱很高但 ★ 偏低（US S2 日常自用大眾 76.5% / ★3.18、JP S1 CP 值優先大眾 91.6% / ★3.46）。\n"
        "「高滿意小族群」= 柱矮但 ★ 高（US S3 USB-C 高端鐵粉 4% / ★4.48）。\n"
        "策略意涵：大眾族群不是擴張對象，應先修補痛點；高端小族群值得集中行銷資源，因為他們滿意且有溢價意願。"
    )

    h2(doc, "2-3　K-Means 分群可靠度：US Silhouette = 0.358　｜　JP Silhouette = 0.570")
    method_block(doc,
        "Silhouette Score（輪廓係數）。"
        "計算每個資料點「與自己 cluster 的相似度」vs「與最近的其他 cluster 的相似度」，"
        "取值範圍 −1 到 +1。",
        "Silhouette > 0.5 視為分群品質「佳」（邊界清晰）；0.25–0.5 為「中等」；< 0.25 為「差」。"
    )
    interp_block(doc,
        "US 0.358（中等）：三個美國顧客群有一定重疊，顧客分群邊界不是非常清楚，"
        "代表美國顧客的屬性偏好向量較接近（差異較不極端）。\n"
        "JP 0.570（佳）：日本三個顧客群邊界清晰，分群結果可靠，"
        "代表日本顧客的偏好差異確實明顯，後續 Targeting 可以放心依群設計。"
    )

    h2(doc, "2-4　URBANER vs 市場頂尖的 RMS 距離差（美國差 0.89 ｜ 日本差 1.46）")
    method_block(doc,
        "RMS（Root Mean Square）距離。"
        "把每支 SKU 的 114 個屬性品質分數視為向量，"
        "計算與「全屬性 = 10」理想向量的均方根距離：d = √(mean((quality_i − 10)²))。",
        "距離越小 = 越接近顧客理想規格。用來比較 URBANER 自家旗艦 vs 市場頂尖之間的規格差距。"
    )
    interp_block(doc,
        "美國：URBANER 旗艦距理想 2.49，市場頂尖（Ufree B0FL267TCG）距理想 1.60，差距 0.89。\n"
        "日本：URBANER 旗艦距理想 3.43，市場頂尖距理想 1.97，差距 1.46。\n"
        "日本落後更大（1.46 > 0.89），且日本頂尖是 Nose/Ear 品類，而 URBANER 目前最強在 Beard——代表 JP Nose/Ear 是一個還沒站穩的空白戰場。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 3：STP 策略分析
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 3｜STP 策略分析")

    h2(doc, "3-1　顧客分群分析（K-Means，K=3）")
    method_block(doc,
        "非監督式分群流程：\n"
        "① 114 維屬性顯著度向量（每則評論）→ StandardScaler 標準化\n"
        "② PCA 降維至保留 85% 變異（降低噪音、加快運算）\n"
        "③ K-Means 分群，K=3（以 Silhouette + 最小群體 ≥ 5% 護欄選定）\n"
        "④ 用 ANOVA 找各 cluster 最具區隔力的屬性 → 人工命名 persona",
        "人數 n = 該 cluster 包含的評論則數；占比 = n / 市場總評論；Avg★ = cluster 內星等均值。"
    )
    table(doc,
        ["市場", "群", "名稱", "佔比", "Avg★", "核心特徵"],
        [
            ["US", "S1", "送禮場景主導", "19.5%", "4.05", "適合送禮給男性，節慶購買"],
            ["US", "S2", "日常自用大眾", "76.5%", "3.18", "鬍鬚 + 耳鼻多用，自用為主"],
            ["US", "S3", "USB-C 高端鐵粉", "4.0%", "4.48", "充電精品，願意付溢價"],
            ["JP", "S1", "CP 值優先大眾", "91.6%", "3.46", "乾電池款，耐用易用"],
            ["JP", "S2", "鬍鬚講究客", "7.8%", "4.01", "高精度、多附件"],
            ["JP", "S3", "靜音敏感族", "0.6%", "3.20", "靜音、輕量"],
        ]
    )
    interp_block(doc,
        "Targeting 建議：\n"
        "US → 優先 S3（高端鐵粉，★4.48，願意付溢價）；S2 大眾做痛點修補（Listing 清晰、附件齊全）。\n"
        "JP → S1 是 91.6% 的大眾核心，但 Avg★ 僅 3.46，需先解決痛點；"
        "S2 鬍鬚講究客（7.8%）是高 ARPU 潛力群，值得個別行銷。"
    )

    h2(doc, "3-2　市場定位地圖（PCA 二維投影）")
    method_block(doc,
        "PCA（Principal Component Analysis）二維投影。"
        "把每支 SKU 的 114 維品質向量降到 2 維（PC1, PC2）。"
        "PC1 通常代表「整體品質強度」，PC2 代表「品質維度間的偏向（規格側重方向）」。",
        "點的位置代表 SKU 在整體品質空間中的定位。"
        "越靠近右上角 = 越接近「全屬性都 10 分」的理想品質點。"
    )
    interp_block(doc,
        "URBANER 自家旗艦落點離右上角有一段距離，對應到 ★3.5–3.8 的滿意度水準。\n"
        "市場頂尖 SKU 集中在高密度區（右上方）。\n"
        "策略用法：盯靠近右上角的競品，拆解其規格 / 文案 / 視覺，找出可複製的做法。"
    )

    h2(doc, "3-3　屬性 × SKU 口碑熱圖")
    method_block(doc,
        "描述統計：Axis B 品質分（0–10）= 該 SKU 在該屬性的口碑評分。"
        "計算方式：星等基線 + 屬性負向關鍵字扣分 + 整體情緒微調。"
        "顯示 Top 20 區隔力屬性 × Top 12 SKU 的矩陣。",
        "顏色綠 = 品質分高（接近 10）；顏色紅 = 品質分低（接近 0）。"
        "橫向：同屬性哪個 SKU 最強。縱向：同 SKU 哪些屬性最弱。"
    )
    interp_block(doc,
        "橫向找標竿：找某屬性列中最綠的 SKU，那就是該屬性的學習對象。\n"
        "縱向找補強點：URBANER 某 SKU 在哪一行是紅色，那就是最需要改善的規格弱點。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 4：Conjoint 偏好
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 4｜Conjoint 偏好（最核心的統計頁面）")

    h2(doc, "4-1　屬性重要性圖（Revealed-Preference Logistic Conjoint）")
    method_block(doc,
        "5 步驟流程：\n"
        "① 資料：Amazon 評論 → Axis B 品質分（0–10）\n"
        "② 屬性工程：品質分離散化為 Low(0–4) / Mid(4–7) / High(7–10)，做 Dummy Encoding\n"
        "③ Response 變數：avg★ ≥ 市場平均 → preferred=1，否則 0\n"
        "④ Split-model Binary Logistic Regression（逐屬性，各自估計 β_mid, β_high）\n"
        "⑤ 重要性% = (max(β) − min(β)) / Σ ranges × 100",
        "屬性重要性 % 越大，該屬性對消費者「是否偏好此產品」的拉扯力越大。"
    )
    table(doc,
        ["市場", "排名", "屬性", "重要性 %", "說明"],
        [
            ["US", "1", "功能合一數", "51.5%*", "* 完全分離問題，幅度高估"],
            ["US", "2", "充電方式 USB-C", "10.5%", ""],
            ["US", "3", "電源類型", "10.4%", ""],
            ["JP", "1", "價格帶", "16.9%", ""],
            ["JP", "2", "附件件數", "16.0%", ""],
            ["JP", "3", "長度調整段數", "15.5%", ""],
        ]
    )
    interp_block(doc,
        "最重要的解讀差異：\n"
        "US → 功能合一（「幾合一」套組）是壓倒性決策因子，Listing 第一張主圖必須講清楚幾合一。\n"
        "JP → 前 6 個屬性都在 10% 以上，日本消費者同時審查附件數、段數、精度、續航、電源類型——"
        "JP Listing 的第一條 bullet 要把「X 個附件 / X 段 / X mm / X 分鐘」全部列出。"
    )

    h2(doc, "4-2　WTP 願付溢價圖（消費者為品質升級願意多付多少錢？）")
    method_block(doc,
        "WTP = β_attribute / |β_price|\n"
        "β_attribute：屬性 part-worth（來自 Split-model Logit，見 4-1）\n"
        "β_price：把真實 Amazon 售價（USD/JPY）加入模型重跑一條 Logit 後估出的價格係數\n"
        "WTP 代表：消費者願意為該屬性品質提升 1 個品質分（0–10）多付的金額。",
        "深色柱 = p < 0.05（統計顯著，高信心解讀）；淺色柱 = 方向性參考（p ≥ 0.05）。"
    )
    table(doc,
        ["市場", "屬性", "WTP", "p 值", "顯著？"],
        [
            ["US", "機身尺寸", "$89.81 /品質分", "0.021", "✓ 顯著"],
            ["US", "附件件數", "$79.86 /品質分", "0.017", "✓ 顯著"],
            ["US", "長度調整段數", "$78.18 /品質分", "0.011", "✓ 顯著"],
            ["US", "充電方式 USB-C", "$32.39 /品質分", "0.035", "✓ 顯著"],
            ["JP", "長度調整段數", "¥7,961 /品質分", "0.042", "✓ 顯著"],
            ["JP", "調整精度", "¥6,774 /品質分", "0.024", "✓ 顯著"],
            ["JP", "附件件數", "¥5,207 /品質分", "0.022", "✓ 顯著"],
        ]
    )
    warn_block(doc,
        "日本 β_price = +0.0004（正值，p=0.096）。正值代表高售價 SKU 評分反而較高（品質訊號效果），"
        "違反 WTP 公式的前提假設（β_price 應為負）。"
        "日本 WTP 數字只能用於「哪個屬性相對更值錢」的方向性排序，不能直接用於定價決策。"
    )
    interp_block(doc,
        "怎麼用 WTP 數字（以美國為例）：\n"
        "附件件數 WTP = $79.86/品質分。若 URBANER 的附件件數品質分目前是 4（Mid 下緣），"
        "競品頂尖是 8（High 中段），差距 4 分，則理論可溢價空間 = $79.86 × 4 = $319。\n"
        "但因 β_price 不顯著（p=0.39），此計算為方向性參考，不代表真的可漲 $319。"
    )

    h2(doc, "4-3　MNL 升級模擬圖（屬性升級後市佔能跳多少？）")
    method_block(doc,
        "MNL（Multinomial Logit）Share-of-Preference 升級模擬。\n"
        "步驟：① 維持其他屬性品質分不變　② 把某一屬性品質分設為 10（理想值）\n"
        "③ 重算 P(URBANER) = exp(U_URBANER) / Σ exp(U_所有 SKU)　④ 比較現況 vs 升級後份額",
        "Softmax 確保所有 SKU 份額加總 = 1。增量 = 升級後份額 − 現況份額。"
        "增量越大代表該屬性的邊際投入 ROI 越高。"
    )
    table(doc,
        ["市場", "升級屬性", "現況份額", "升級後份額", "增量", "意義"],
        [
            ["JP", "長度調整段數 → High", "1.6%", "94.8%", "+93.2%", "最高 ROI 的單一投資"],
            ["JP", "調整精度 → High", "1.6%", "86.2%", "+84.6%", "與段數互補"],
            ["JP", "附件件數 → High", "1.6%", "52.0%", "+50.4%", "第三優先"],
            ["US", "機身尺寸 → High", "5.2%", "5.9%", "+0.7%", "單屬性效果有限"],
            ["US", "附件件數 → High", "5.2%", "5.4%", "+0.2%", "需組合策略"],
        ]
    )
    warn_block(doc,
        "JP 升級後 94.8% 是理論極端值——單一屬性設為 10 在 softmax 中產生壓倒性優勢。"
        "實際意義是「長度段數是 JP 最有槓桿的改善點」，而非字面上的 94.8% 市佔。"
    )
    interp_block(doc,
        "US 單屬性升級增量都 < 1%，代表美國的問題不在某個單一規格不夠好，"
        "而是「整體套組形態不足」——需要同步升級附件件數 + 功能合一數 + USB-C 充電的組合，"
        "而非拚命強化單一屬性。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 5：最優產品組合
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 5｜最優產品組合")

    h2(doc, "5-1　Card 22：兩市場最佳虛擬設計方案")
    method_block(doc,
        "Orthogonal Design（正交設計）產生 24 張虛擬產品卡片，"
        "每張卡片是各屬性水準的特定組合。\n"
        "Card 22 是 24 張中 MNL 份額最高的那一張（分別在 US 和 JP 各算一次）。\n"
        "購買概率 = Z-score 正規化後的 Sigmoid 輸出，同 Page 1 的 83.3% / 89.1%。",
        "Card 22 不是真實存在的 SKU，是「如果推出這樣規格的新品，競爭力如何」的模擬。"
    )
    table(doc,
        ["市場", "Card 22 核心規格", "購買概率", "全場 MNL 排名"],
        [
            ["美國", "USB-C × 38 段 × ≥10 件套組 × 5合1+", "83.3%", "第 33 名 / 76"],
            ["日本", "IPX7 × 90 分高續航 × ≥10 件全配款 × 1mm 精度", "89.1%", "第 9 名 / 60"],
        ]
    )
    interp_block(doc,
        "兩市場骨幹相同（≥10 件套組 + 38 段），但差異點很明顯：\n"
        "US → USB-C 快充 + 5合1 多功能是差異化重點（送禮場景驅動）。\n"
        "JP → IPX7 防水 + 90 分高續航 + 1mm 精度是差異化重點（規格講究驅動）。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 6：社群洞察
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 6｜社群洞察")

    h2(doc, "6-1　跨類別共識 22 條（US 11 + JP 11）")
    method_block(doc,
        "評論文本 NLP 描述統計。對各品類（001–009）的評論文本分別做屬性提取，"
        "找出在兩市場中出現頻率都高的屬性主題，人工確認為「跨類別共識洞察」。",
        "共識條數越多代表兩市場顧客有越多共同在意的點，是跨市場產品策略的基礎。"
    )
    interp_block(doc,
        "22 條共識洞察是「兩市場都有強烈共鳴的訴求點」，優先用於跨市場通用版 Listing 文案。"
        "非共識的洞察（只在一個市場出現）則做市場專屬版本的差異化文案。"
    )

    h2(doc, "6-2　各品類平台明細條形圖")
    method_block(doc,
        "描述統計：對每個品類（001–009）按市場（US/JP）統計評論數與 SKU 數，"
        "以條形圖呈現品類間的資料規模差異。",
        "Bar 越長代表該品類的市場資料越豐富，後續統計的可靠性越高。"
    )
    interp_block(doc,
        "品類 001（Beard/Mustache Trimmers）資料量最多——所有聯合分析的核心結果都來自此品類。"
        "若要把結論套用到其他品類（如 006 Pet Clippers），需要重新對該品類跑相同分析。"
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # PAGE 7：行銷劇本
    # ══════════════════════════════════════════════════════
    h1(doc, "頁面 7｜行銷劇本")

    h2(doc, "7-1　文案建議的統計依據")
    method_block(doc,
        "行銷劇本頁面的文案建議不是直接的統計推論，而是基於前 6 頁分析結果的「洞察轉化」。"
        "每一條文案對應一個或多個統計依據。",
        "文案不需要引用 p 值或 F 值，但答辯時需要能說出每條文案的來源是哪個分析。"
    )
    table(doc,
        ["文案重點", "統計依據", "對應頁面"],
        [
            ["強調「幾合一」多功能（US）", "Conjoint 重要性 51.5%（功能合一數）", "Conjoint 偏好 4-1"],
            ["USB-C 充電符號（US）", "WTP $32.39/分、Conjoint 重要性 10.5%", "Conjoint 偏好 4-1 / 4-2"],
            ["附件件數標示（兩市場）", "US WTP $79.86 / JP WTP ¥5,207；兩市場 ANOVA Top 3", "雙市場 2-1 / Conjoint 4-2"],
            ["IPX7 badge（JP）", "JP 長度段數升級模擬 +93%（段數為 Top 1 ROI）", "Conjoint 4-3"],
            ["父親節/聖誕節送禮（US）", "US S1 送禮族群 19.5%、ANOVA Top 1（F=1299）", "雙市場 2-1 / STP 3-1"],
            ["0.5mm 刻度精度（JP）", "JP Conjoint 重要性 14.9%、WTP ¥6,774", "Conjoint 4-1 / 4-2"],
        ]
    )

    divider(doc)

    # ══════════════════════════════════════════════════════
    # 統計方法對照速查表
    # ══════════════════════════════════════════════════════
    h1(doc, "附錄｜統計方法總對照表")

    table(doc,
        ["方法名稱", "在儀表板的位置", "一句話說明", "輸出的數字"],
        [
            ["Axis B 品質分", "全頁共用", "NLP 從評論文本萃取各屬性的口碑品質分數（0–10）", "屬性品質分矩陣"],
            ["K-Means 分群", "雙市場 / STP", "非監督式聚類，把相似偏好的評論歸為同一群", "cluster 標籤、佔比、Avg★"],
            ["PCA 降維", "STP 定位地圖", "把 114 維品質向量壓縮到 2 維便於視覺化", "二維座標（PC1, PC2）"],
            ["Silhouette Score", "雙市場", "評估 K-Means 分群品質（0–1，越高越好）", "US 0.358 / JP 0.570"],
            ["One-way ANOVA", "雙市場 / STP", "檢定某屬性是否在不同顧客群間有顯著差異（F 值）", "F 值、p 值"],
            ["RMS 距離", "雙市場", "SKU 與「全屬性=10」理想點的均方根距離", "US 旗艦 2.49 / JP 旗艦 3.43"],
            ["Split-model Logit", "Conjoint", "逐屬性估計 part-worth 偏好權重（β 係數）", "β_mid, β_high"],
            ["Range-based 重要性", "Conjoint", "屬性 part-worth 全範圍÷總範圍，換算為 %", "US 51.5% / JP 16.9%…"],
            ["WTP", "Conjoint", "β_attribute / |β_price|，偏好換算為金額", "美 $89.81 / 日 ¥7,961…"],
            ["Sigmoid 選擇概率", "Conjoint / 最優產品", "Z-score 正規化效用代入 Sigmoid，得相對偏好強度", "83.3% / 89.1%"],
            ["MNL 市場份額", "總覽 / Conjoint", "Softmax 競爭選擇模型，全選擇集份額加總=1", "US 5.21% / JP 11.26%"],
            ["Hedonic Pricing OLS", "（延伸分析）", "ln(售價) = β × 屬性分 + ε，從定價行為反推屬性隱含價格", "β 係數（市場定價彈性）"],
        ]
    )

    # ─ 頁尾 ────────────────────────────────────────────
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(16)
    r_f = p_f.add_run(
        "方法參照：McFadden (1974), Green & Srinivasan (1990), MacQueen (1967) K-Means "
        "｜ FourSight Lab × URBANER 產學合作 ｜ 2026-05-14"
    )
    sr(r_f, size=9, italic=True, color=C_MUTED)

    return doc


if __name__ == "__main__":
    import os
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "URBANER_儀表板統計解讀.docx")
    doc = build()
    doc.save(out)
    print(f"Saved: {out}")
