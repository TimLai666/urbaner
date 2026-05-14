"""
產生 URBANER 聯合分析統計步驟說明文件（Word 版）
輸出：output_conjoint/URBANER_統計方法說明.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Microsoft YaHei"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_run(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    colors = {
        1: RGBColor(0x1F, 0x3A, 0x5F),
        2: RGBColor(0x2E, 0x5C, 0x8A),
        3: RGBColor(0x4A, 0x7A, 0xB0),
        4: RGBColor(0x5A, 0x5A, 0x5A),
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    return p


def add_body(doc, text, indent=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    style_run(run, size=10, italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{'  ' * level}• {text}")
    style_run(run, size=10)
    return p


def add_faq(doc, q, a):
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Cm(0.5)
    p_q.paragraph_format.space_before = Pt(4)
    p_q.paragraph_format.space_after = Pt(2)
    run_q = p_q.add_run(f"Q：{q}")
    style_run(run_q, size=10, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))

    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.8)
    p_a.paragraph_format.space_before = Pt(0)
    p_a.paragraph_format.space_after = Pt(4)
    run_a = p_a.add_run(f"A：{a}")
    style_run(run_a, size=10, italic=True, color=RGBColor(0x33, 0x33, 0x33))


def add_table(doc, headers, rows, header_color="1F3A5F", header_text_color=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        set_cell_shading(hdr_cells[i], header_color)
        set_cell_borders(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].add_run(h)
        tc = RGBColor(0xFF, 0xFF, 0xFF) if header_text_color is None else header_text_color
        style_run(run, size=9, bold=True, color=tc)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F5F7FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ""
            set_cell_shading(row_cells[c_idx], bg)
            set_cell_borders(row_cells[c_idx])
            run = row_cells[c_idx].paragraphs[0].add_run(str(val))
            style_run(run, size=9)

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x1A)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── 封面標題 ───────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(8)
    run_t = p_title.add_run("URBANER 聯合分析——統計步驟完整說明")
    style_run(run_t, size=20, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)
    run_s = p_sub.add_run("跨境電商競賽答辯版 | 2026-05-14")
    style_run(run_s, size=11, italic=True, color=RGBColor(0x88, 0x88, 0x88))

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(16)
    run_d = p_desc.add_run("適用場合：競賽答辯評審追問、團隊內部統計知識對齊")
    style_run(run_d, size=10, color=RGBColor(0x55, 0x55, 0x55))

    add_divider(doc)

    # ─── 分析鏈總覽 ──────────────────────────────────────────
    add_heading(doc, "一、分析鏈總覽", level=1)
    add_body(doc, "整條分析分為 5 個統計層次，由觀察資料出發，最終產出四種商業洞察：")

    add_code_block(doc,
        "Layer 1  資料取得（Amazon 評論 → 品質分 + 偏好標記）\n"
        "         ↓\n"
        "Layer 2  屬性工程（離散化 → Dummy Encoding）\n"
        "         ↓\n"
        "Layer 3  模型估計（Split-model Binary Logistic Regression）\n"
        "              ↓ 產出：Part-worth Utilities（β 係數）\n"
        "         ┌────┬────────────┬──────────┬──────────┐\n"
        "Layer 4a      4b           4c         4d\n"
        "屬性重要性  選擇概率     WTP 溢價   MNL 市場份額\n"
        "（Range）  （Sigmoid）  （β比值）  （Softmax）\n"
        "         └────┴────────────┴──────────┴──────────┘\n"
        "                    ↓\n"
        "Layer 5  Hedonic Pricing OLS（補充市場定價驗證）"
    )

    add_body(doc,
        "核心邏輯：用真實市場的 Amazon 評論觀察資料（Revealed-Preference），"
        "透過 Logistic Regression 量化「產品屬性 → 消費者偏好」的關係，"
        "再從係數衍生屬性重要性、願付溢價、市場競爭份額三種商業洞察。"
    )

    add_divider(doc)

    # ─── Layer 1 ─────────────────────────────────────────────
    add_heading(doc, "二、Layer 1：資料取得與前處理", level=1)
    add_heading(doc, "做了什麼", level=3)
    add_table(doc,
        ["項目", "說明"],
        [
            ["資料來源", "Amazon 競品 SKU 評論文本 + 平均星等（avg★）"],
            ["屬性量化", "使用 Axis B 品質分（0–10）代表每個 SKU 在各屬性上的品質水準"],
            ["Response 變數", "preferred = 1（若 avg★ ≥ 市場平均），否則 preferred = 0"],
            ["美國樣本", "52 個 SKU"],
            ["日本樣本", "36 個 SKU"],
        ]
    )

    add_heading(doc, "為什麼這樣做", level=3)
    add_body(doc, "為什麼用 Amazon 評論而不用問卷（Survey）？")
    add_body(doc,
        "這是 Revealed-Preference（顯示偏好）設計，優於問卷的 Stated-Preference（陳述偏好）：\n"
        "評論來自真實購買後的消費者，反映實際決策行為，外部效度高；"
        "問卷容易有社會期望偏誤（受訪者說他在乎品質，其實只看價格）。"
        " avg★ 已聚合數十到數百位購買者的綜合評價，作為偏好代理變數。"
    )

    add_heading(doc, "競賽 FAQ", level=3)
    add_faq(doc,
        "Axis B 品質分是什麼？怎麼得到的？",
        "從評論文本做 NLP 分析，針對每個屬性（防水、附件、充電方式⋯⋯）擷取消費者的評價情緒與提及強度，"
        "量化為 0–10 分。分數越高代表市場上對該屬性的認可度越高。"
    )
    add_faq(doc,
        "52 個 SKU 夠多嗎？",
        "樣本偏小是本研究的已知限制。所有統計結果均為方向性分析，不宜直接外推至全市場。"
        "若要提升推論力，理想樣本數 ≥ 200 個 SKU 或搭配正式消費者問卷。"
    )

    add_divider(doc)

    # ─── Layer 2 ─────────────────────────────────────────────
    add_heading(doc, "三、Layer 2：屬性工程", level=1)
    add_heading(doc, "做了什麼", level=3)

    add_body(doc, "步驟一：離散化（將 0–10 品質分切成 3 個水準）")
    add_table(doc,
        ["水準", "分數範圍", "意義"],
        [
            ["Low", "0–4", "市場上表現較差的規格"],
            ["Mid", "4–7", "市場上中等規格"],
            ["High", "7–10", "市場上高規格"],
        ]
    )

    add_body(doc, "步驟二：Dummy Encoding（虛擬變數編碼）——Low 作為參照組（Reference Level）")
    add_table(doc,
        ["水準", "mid_dummy", "high_dummy"],
        [
            ["Low（參照組）", "0", "0"],
            ["Mid", "1", "0"],
            ["High", "0", "1"],
        ]
    )

    add_heading(doc, "為什麼這樣做", level=3)
    add_body(doc,
        "為什麼要離散化？聯合分析的核心假設是：消費者對產品的評價是由各「屬性水準」的偏好加總而來。"
        "離散化讓每個水準有獨立的係數，不強迫假設「品質分每增加 1 分，效用就線性增加」。"
        "消費者對防水等級的感受並不一定是線性的（從無防水到 IPX4 是一大跨越，IPX4 到 IPX7 相對較小）。"
    )
    add_body(doc,
        "為什麼用 Dummy Encoding 而不是 Effect Coding？"
        "Dummy coding 的係數 β 直接代表「相較於 Low 水準，Mid/High 的效用增量」，解讀最直觀，"
        "適合答辯說明：β_mid = 從最差規格升到中等規格，消費者效用增加多少。"
    )

    add_divider(doc)

    # ─── Layer 3 ─────────────────────────────────────────────
    add_heading(doc, "四、Layer 3：Split-model Binary Logistic Regression（核心模型）", level=1)
    add_heading(doc, "做了什麼", level=3)
    add_body(doc, "對每個屬性分別執行一條 Binary Logistic Regression：")
    add_code_block(doc, "logit(P(preferred=1)) = β₀ + β_mid × mid_dummy + β_high × high_dummy")
    add_bullet(doc, "因變數：preferred（0 或 1）")
    add_bullet(doc, "自變數：該屬性的 mid_dummy, high_dummy")
    add_bullet(doc, "輸出：Part-worth Utilities（部分效用值）—— β_mid, β_high")

    add_heading(doc, "為什麼用 Logistic Regression 而不是 OLS？", level=3)
    add_table(doc,
        ["比較", "OLS 線性回歸", "Binary Logistic Regression"],
        [
            ["因變數要求", "連續值", "二元（0/1）"],
            ["預測結果", "可超出 [0,1]", "自動映射至 (0,1) 機率"],
            ["理論基礎", "最小平方法（OLS）", "最大概似法（MLE）"],
            ["本案適用", "❌ preferred 是 0/1", "✅"],
        ]
    )

    add_heading(doc, "為什麼是 Split-model（逐屬性）而不是全屬性一起跑？", level=3)
    add_body(doc,
        "我們的資料是「真實市場 SKU」，不是實驗設計的正交卡片，屬性之間有相關性"
        "（例如：USB-C 充電的機器通常防水也好、附件也多）。"
        "若把所有屬性丟進同一條迴歸，會產生多重共線性（Multicollinearity），VIF 飆高，"
        "各屬性的 β 係數互相干擾，難以解讀。"
        "Split-model 每次只看一個屬性，避開共線性，係數更穩定、更可解讀。"
        "代價是無法捕捉屬性間的交互效應，但對於方向性分析已足夠。"
    )

    add_heading(doc, "Part-worth Utility 是什麼意思？", level=3)
    add_bullet(doc, "β_mid > 0：Mid 水準比 Low 水準更受消費者偏好（log-odds 單位）")
    add_bullet(doc, "β_high > β_mid：High 水準偏好更強")
    add_bullet(doc, "β 的絕對大小反映：升一個水準，偏好程度增加多少")

    add_heading(doc, "本案數字舉例（美國市場，附件件數）", level=3)
    add_table(doc,
        ["水準", "β 係數", "p 值", "解讀"],
        [
            ["Mid（7件 vs ≤3件）", "10.53", "0.97", "正向，附件越多越偏好"],
            ["High（≥10件 vs ≤3件）", "13.81", "0.97", "正向，High 偏好更強"],
        ]
    )

    add_heading(doc, "競賽 FAQ", level=3)
    add_faq(doc,
        "p 值都很大（接近 1.0），這個模型有效嗎？",
        "樣本只有 19–52 個 SKU，統計檢定力（Power）不足，導致 p 值大多不顯著。"
        "但係數的方向（正負）仍具方向性參考價值——所有屬性在兩個市場的 β 方向都符合市場直覺。"
        "正式推論需要更大樣本（≥200 SKU）。"
    )
    add_faq(doc,
        "什麼是 McFadden Pseudo R²？",
        "Logit 模型沒有 OLS 的 R²，用 McFadden Pseudo R² 替代。"
        "計算方式：1 − (LL_model / LL_null)，值在 0.2–0.4 之間視為模型擬合良好。"
        "本案數值偏低，反映樣本小、噪音大。"
    )

    add_divider(doc)

    # ─── Layer 4a ────────────────────────────────────────────
    add_heading(doc, "五、Layer 4a：屬性重要性（Range-based Importance）", level=1)
    add_heading(doc, "公式", level=3)
    add_code_block(doc,
        "range_attr   = max(β_low, β_mid, β_high) − min(β_low, β_mid, β_high)\n"
        "importance%  = range_attr / Σ(range of all attributes) × 100"
    )
    add_body(doc, "（注意：β_low = 0，因為 Low 是參照組）")

    add_heading(doc, "為什麼這樣算而不直接比較 β 大小？", level=3)
    add_body(doc,
        "不同屬性的 β 係數尺度不可直接比較（功能合一數的 β 達 60–70，附件件數的 β 約 13，"
        "但不代表功能合一數就「6倍重要」）。"
        "Range-based Importance 正規化了尺度差異，讓各屬性在同一百分比尺度競爭。"
        "這是聯合分析領域的標準方法（Green & Srinivasan, 1990）。"
    )

    add_heading(doc, "本案結果", level=3)
    add_table(doc,
        ["市場", "排名 1", "排名 2", "排名 3"],
        [
            ["美國", "功能合一數（51.5%）*", "充電方式 USB-C（10.5%）", "電源類型（10.4%）"],
            ["日本", "價格帶（16.9%）", "附件件數（16.0%）", "長度調整段數（15.5%）"],
        ]
    )
    add_body(doc,
        "* 美國「功能合一數」出現完全分離（Perfect Separation）問題，係數異常大，"
        "51.5% 數字需附加說明，不宜直接引用於策略建議。",
        italic=True, color=RGBColor(0xCC, 0x44, 0x00)
    )

    add_divider(doc)

    # ─── Layer 4b ────────────────────────────────────────────
    add_heading(doc, "六、Layer 4b：選擇概率（相對偏好強度）", level=1)
    add_heading(doc, "公式", level=3)
    add_code_block(doc,
        "U_i  = Σ β_attr（各屬性 part-worth 加總）\n"
        "Z_i  = (U_i − μ_U) / σ_U      # Z-score 正規化\n"
        "P_i  = 1 / (1 + exp(−Z_i))    # Sigmoid 函數"
    )

    add_heading(doc, "為什麼要 Z-score 正規化？", level=3)
    add_body(doc,
        "若直接把 U 代入 Sigmoid：美國市場功能合一數的 β 達 60–70，加總 U 可能超過 100，"
        "Sigmoid(100) ≈ 1.0000000000（完全飽和），所有高分 SKU 的概率都無限接近 1.0，"
        "無法區分排名。Z-score 正規化後，所有 SKU 在同一相對尺度競爭，Sigmoid 的輸出才有區分度。"
    )

    add_heading(doc, "「選擇概率」的正確解讀", level=3)
    add_body(doc,
        "這裡的概率不是「100 位消費者中有幾位會購買」的市場份額，"
        "而是：這個 SKU 相對於市場平均水準，被消費者選擇的相對偏好強度。"
        "P > 0.5 代表該 SKU 優於市場平均；P 越高，競爭力越強。"
    )

    add_heading(doc, "本案結果", level=3)
    add_table(doc,
        ["市場", "Top 1 ASIN", "選擇概率"],
        [
            ["美國", "B0D2DRS9KY", "0.659"],
            ["日本", "B0742G961R", "0.708"],
        ]
    )

    add_divider(doc)

    # ─── Layer 4c ────────────────────────────────────────────
    add_heading(doc, "七、Layer 4c：WTP 願付溢價（Willingness to Pay）", level=1)
    add_heading(doc, "公式", level=3)
    add_code_block(doc, "WTP（屬性升級 vs Low）= β_attr / |β_price|")
    add_bullet(doc, "分子 β_attr：該屬性從 Low 升至 Mid 或 High 的效用增量（來自 Layer 3）")
    add_bullet(doc, "分母 |β_price|：售價每增加 $1，效用減少多少（補入真實售價重跑 Logit 估計）")

    add_heading(doc, "WTP 公式背後的邏輯推導", level=3)
    add_code_block(doc,
        "效用無差異條件：\n"
        "β_attr × Δ品質 + β_price × ΔWTP = 0\n\n"
        "移項得：\n"
        "ΔWTP = −(β_attr / β_price) = β_attr / |β_price|"
    )
    add_body(doc,
        "消費者多願意付 ΔWTP 元，剛好能抵消屬性升級帶來的效用增量——這是「願付溢價」的貨幣化定義。"
    )

    add_heading(doc, "β_price 可靠性", level=3)
    add_table(doc,
        ["市場", "β_price", "p 值", "可靠性"],
        [
            ["美國（USD）", "−0.0212", "0.3864", "✅ 符號正確，方向性可信，但不顯著"],
            ["日本（JPY）", "+0.0004", "0.0963", "⚠️ 符號為正，WTP 不具定價意義"],
        ]
    )
    add_body(doc,
        "日本 β_price 為正的解釋：在本案日本樣本中，高售價的 SKU 評分反而較高——"
        "這是「品質訊號效果」（Price-as-quality heuristic）。消費者把高價格當作品質保證，"
        "導致 Logit 估計出正的 β_price，違反 WTP 的計算前提。"
        "此為樣本小（21 個 SKU）且市場特殊性所致。",
        italic=True, color=RGBColor(0xCC, 0x44, 0x00)
    )

    add_heading(doc, "美國市場 WTP 結果（可引用）", level=3)
    add_table(doc,
        ["屬性", "WTP（$/品質分）", "p 值", "顯著性"],
        [
            ["機身尺寸", "$89.81", "0.021", "*"],
            ["附件件數", "$79.86", "0.017", "*"],
            ["長度調整段數", "$78.18", "0.011", "*"],
            ["充電方式 USB-C", "$32.39", "0.035", "*"],
            ["電源類型", "$26.27", "0.080", "†"],
            ["防水等級", "$24.39", "0.093", "†"],
        ]
    )

    add_heading(doc, "競賽 FAQ", level=3)
    add_faq(doc,
        "WTP $89.81/品質分 是什麼意思？",
        "品質分是 0–10 的 Axis B 分數。若機身尺寸的品質分從 3 分升到 8 分（+5 分），"
        "WTP 約為 $89.81 × 5 = $449——但此數字主要用於屬性排序（哪個屬性比哪個更值錢），"
        "而非精確定價，因為 β_price 不顯著（p=0.39），幅度估計仍有不確定性。"
    )
    add_faq(doc,
        "WTP 和屬性重要性的排名不一樣，哪個更可信？",
        "兩者問的問題不同。屬性重要性（Range）問：這個屬性對偏好的影響範圍有多大？"
        "WTP 問：消費者願意用多少錢換取這個屬性的升級？"
        "兩者都有參考價值，差異代表兩種不同維度的消費者決策機制。"
    )

    add_divider(doc)

    # ─── Layer 4d ────────────────────────────────────────────
    add_heading(doc, "八、Layer 4d：MNL 市場偏好份額（Share-of-Preference）", level=1)
    add_heading(doc, "公式", level=3)
    add_code_block(doc,
        "P(選擇產品 i) = exp(U_i) / Σ exp(U_j)     j ∈ 全選擇集\n\n"
        "U_i = Σ (β_mid × mid_dummy + β_high × high_dummy)"
    )

    add_heading(doc, "MNL 和 Binary Logit 的關係", level=3)
    add_table(doc,
        ["比較", "Binary Logit（Layer 3）", "MNL（Layer 4d）"],
        [
            ["選擇問題", "買 vs 不買（2 選 1）", "從多個產品中選一個（N 選 1）"],
            ["數學形式", "Sigmoid: 1/(1+e^{−U})", "Softmax: e^{U_i}/Σe^{U_j}"],
            ["結果解讀", "某產品被偏好的概率", "在市場競爭中的份額"],
            ["本質", "相同（MNL 是 Binary Logit 的推廣）", "←"],
        ]
    )

    add_heading(doc, "為什麼用 MNL 做市場份額分析？", level=3)
    add_body(doc,
        "競賽最終問的不只是「消費者喜不喜歡這個屬性」，"
        "而是「URBANER 在市場上能搶到多少份額？」。"
        "MNL 把 URBANER 和所有競品放在同一尺度上比較，"
        "Softmax 確保所有選項的份額加總 = 1.0（100%），符合市場份額的概念。"
        "同時可模擬：若 URBANER 提升某個屬性，市場份額會增加多少？"
    )

    add_heading(doc, "IIA 假設（Independence of Irrelevant Alternatives）", level=3)
    add_body(doc,
        "MNL 有一個已知限制：假設新競品進入市場，所有現有選項的份額等比例縮減，"
        "不允許特定產品組合之間有更強的替代效果。這在本案市場初步模擬階段是可接受的近似。"
    )

    add_heading(doc, "本案結果", level=3)
    add_table(doc,
        ["市場", "群組", "偏好份額"],
        [
            ["美國", "URBANER 合計", "5.2%"],
            ["美國", "競品合計", "94.7%"],
            ["日本", "URBANER 合計", "11.3%"],
            ["日本", "設計卡片合計", "3.1%"],
            ["日本", "競品合計", "85.6%"],
        ]
    )
    add_body(doc,
        "Card 22（虛擬設計方案）在日本全選擇集排名第 9，偏好份額 3.1%，優於多數競品 SKU，"
        "代表此規格組合在日本市場具有競爭力。"
    )

    add_heading(doc, "屬性升級模擬（What-if）", level=3)
    add_table(doc,
        ["市場", "升級屬性", "現況份額", "升級後份額", "增幅"],
        [
            ["日本", "長度調整段數（→High）", "1.6%", "94.8%", "+93.2%"],
            ["日本", "調整精度（→High）", "1.6%", "86.2%", "+84.6%"],
            ["日本", "附件件數（→High）", "1.6%", "52.0%", "+50.4%"],
        ]
    )
    add_body(doc,
        "注意：94.8% 是理論模型的極端值，因為單一屬性升至「完美 High」在 softmax 中會產生壓倒性優勢。"
        "實際意義在於：長度調整段數是日本市場 URBANER 最有槓桿的改善點，而非字面上能搶到 94.8%。",
        italic=True, color=RGBColor(0x55, 0x55, 0x55)
    )

    add_heading(doc, "競賽 FAQ", level=3)
    add_faq(doc,
        "MNL 和 Layer 4b 的選擇概率有什麼不同？",
        "Layer 4b 的選擇概率（Sigmoid）是「單一產品相對市場均值的偏好強度」，沒有和競品直接競爭。"
        "MNL（Layer 4d）是「在明確選擇集中，URBANER 和每一個競品直接競爭後的市場份額」，"
        "更接近真實競爭情境。"
    )
    add_faq(doc,
        "Card 22 是什麼？",
        "Card 22 是由正交設計（Orthogonal Design）生成的 24 張虛擬產品規格卡片之一，"
        "用來模擬「如果 URBANER 推出這樣規格的新品，競爭力如何？」。"
        "Card 22 在日本排名第 9，代表其屬性組合在日本市場具有高競爭力。"
    )

    add_divider(doc)

    # ─── Layer 5 ─────────────────────────────────────────────
    add_heading(doc, "九、Layer 5：Hedonic Pricing OLS（補充市場定價驗證）", level=1)
    add_heading(doc, "公式", level=3)
    add_code_block(doc,
        "ln(Price) = β₀ + β₁×屬性1品質分 + β₂×屬性2品質分 + … + ε"
    )
    add_body(doc,
        "每個 β 的解讀：市場上的競品，該屬性品質分每提升 1 分，售價平均高出 β×100%（半對數模型）。"
    )

    add_heading(doc, "為什麼取對數（半對數模型）？", level=3)
    add_bullet(doc, "售價是正值且分佈右偏，取對數後近似常態分佈，符合 OLS 的殘差假設")
    add_bullet(doc, "係數解讀從「元」變成「百分比漲幅」，更直觀")

    add_heading(doc, "和 WTP 的差異", level=3)
    add_table(doc,
        ["維度", "WTP（路線 A）", "Hedonic Pricing（路線 C）"],
        [
            ["問的問題", "消費者願意多付多少？", "市場實際如何用屬性定價？"],
            ["視角", "需求面（消費者偏好）", "供給面（市場定價行為）"],
            ["方法", "Logit 係數比（β_attr / β_price）", "OLS 半對數回歸"],
            ["用途", "功能升級的定價決策", "市場定位、競品定價分析"],
        ]
    )

    add_heading(doc, "本案限制（誠實揭露）", level=3)
    add_table(doc,
        ["市場", "R²", "Adjusted R²", "說明"],
        [
            ["美國", "0.425", "0.059", "Adj-R² 接近 0，模型解釋力弱"],
            ["日本", "0.272", "−0.120", "Adj-R² 為負，迴歸不穩定"],
        ]
    )
    add_body(doc,
        "原因：有效樣本僅 19（US）/ 21（JP）個 SKU，相對於 7 個自變數，自由度嚴重不足（n/p 比例 < 3）。"
        "所有係數 p 值均不顯著。結論：Hedonic Pricing 在本案中僅供定性參考（屬性與售價的相關方向），"
        "不宜用於精確的定價回歸。",
        italic=True, color=RGBColor(0xCC, 0x44, 0x00)
    )

    add_divider(doc)

    # ─── 統計限制總表 ────────────────────────────────────────
    add_heading(doc, "十、統計限制誠實揭露（競賽必答）", level=1)
    add_table(doc,
        ["限制", "說明", "對結果的影響"],
        [
            ["樣本偏小", "US 52 / JP 36 個 SKU，有效樣本更少", "係數 p 值多不顯著，結果為方向性"],
            ["完全分離", "美國「功能合一數」出現 Perfect Separation", "係數異常大（β>60），重要性 51.5% 需存疑"],
            ["β_price 符號錯誤（JP）", "日本售價係數為正", "日本 WTP 不可信，僅供方向參考"],
            ["Hedonic Adj-R² 為負", "樣本量相對變數數過少", "Hedonic 迴歸不穩定，定性參考"],
            ["IIA 假設", "MNL 假設競品替代效果對稱", "市場份額模擬為近似值"],
            ["無交互效應", "Split-model 無法捕捉屬性間交互", "組合規格的效用估計可能有偏誤"],
        ]
    )

    add_divider(doc)

    # ─── 術語速查 ────────────────────────────────────────────
    add_heading(doc, "十一、關鍵術語速查", level=1)
    add_table(doc,
        ["術語", "中文", "一句話說明"],
        [
            ["Revealed-Preference", "顯示偏好", "從真實行為（而非問卷）觀察消費者偏好"],
            ["Part-worth Utility", "部分效用值", "各屬性水準對總效用的貢獻量（β 係數）"],
            ["Split-model", "分割模型", "逐屬性分開跑迴歸，避免多重共線性"],
            ["Dummy Encoding", "虛擬變數編碼", "將類別水準轉換為 0/1 變數"],
            ["Range-based Importance", "範圍法重要性", "屬性能影響效用的最大幅度，換算成百分比"],
            ["WTP", "願付溢價", "消費者願意為屬性升級多付的金額"],
            ["MNL", "多項 Logit 模型", "在多個產品選項中計算各產品的選擇機率"],
            ["IIA", "無關替代方案獨立性", "MNL 的假設：加入新選項，現有份額等比例縮減"],
            ["Hedonic Pricing", "享樂定價法", "從市場售價反推各屬性的隱含價格"],
            ["Softmax", "Softmax 函數", "MNL 的數學核心，確保所有份額加總為 1"],
            ["Sigmoid", "Sigmoid 函數", "Binary Logit 的輸出函數，映射至 (0,1)"],
        ]
    )

    add_divider(doc)

    # ─── 頁尾 ────────────────────────────────────────────────
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(16)
    run_f = p_footer.add_run(
        "統計方法參照：McFadden (1974), Green & Srinivasan (1990) | "
        "FourSight Lab × URBANER 產學合作專案 | 2026-05-14"
    )
    style_run(run_f, size=9, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    return doc


if __name__ == "__main__":
    import os
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "URBANER_統計方法說明.docx")
    doc = build_doc()
    doc.save(out_path)
    print(f"Word saved: {out_path}")
