"""
Generate a complete Word report combining:
- conjoint_report.md (URBANER 雙市場聯合分析報告)
- product_cards_report.md (URBANER 聯合分析產品卡設計報告)

All content is included verbatim — no truncation.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Microsoft YaHei"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_run(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    colors = {
        1: RGBColor(0x1F, 0x3A, 0x5F),
        2: RGBColor(0x2E, 0x5C, 0x8A),
        3: RGBColor(0x4A, 0x7A, 0xB0),
        4: RGBColor(0x5A, 0x5A, 0x5A),
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    return p


def add_para(doc, text, size=10, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_blockquote(doc, lines):
    """A multi-line quoted note block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        style_run(run, size=10, color=RGBColor(0x55, 0x55, 0x55))
        run.italic = True


def add_table(doc, header, rows, header_fill="1F3A5F", first_col_bold=False, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        style_run(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for r, row in enumerate(rows):
        zebra = "F0F4F8" if r % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = table.rows[r + 1].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val is not None else "")
            bold = first_col_bold and i == 0
            style_run(run, size=9, bold=bold)
            set_cell_shading(cell, zebra)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # spacing after
    doc.add_paragraph()
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    style_run(run, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def main():
    doc = Document()
    set_default_font(doc)

    # set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("URBANER 雙市場聯合分析（Conjoint）完整報告")
    style_run(run, size=22, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Revealed-Preference Logistic Conjoint × L18 Orthogonal Card Design")
    style_run(run, size=12, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("FourSight Lab × URBANER 產學合作專案　|　資料截止 2026-05-05")
    style_run(run, size=10, color=RGBColor(0x77, 0x77, 0x77))

    doc.add_paragraph()

    # ===========================================================
    # PART A: 雙市場聯合分析報告
    # ===========================================================
    add_heading(doc, "Part A. URBANER 雙市場聯合分析報告", level=1)

    add_blockquote(doc, [
        "方法：Revealed-Preference Logistic Conjoint",
        "資料：Amazon 評論兩階段評分（Axis B 品質分 0–10）× SKU 平均星等",
        "美國樣本：52 個 SKU　　日本樣本：36 個 SKU",
        "資料截止：2026-05-05",
    ])
    add_separator(doc)

    # 一、分析架構
    add_heading(doc, "一、分析架構（5 Phase）", level=2)

    add_para(doc, "I. 資料取得  →  II. 屬性工程  →  III. 實驗設計", size=10)
    add_para(doc, "                                        ↓", size=10)
    add_para(doc, "V. 洞察轉化  ←  IV. 模型估計  ←────┘", size=10)
    doc.add_paragraph()

    add_table(
        doc,
        header=["Phase", "執行內容", "狀態"],
        rows=[
            ["I", "使用 Axis B 品質分（supply track）+ 評論星等（demand track）", "✅"],
            ["II", "篩選 8 屬性 / 市場，離散化為 Low / Mid / High 三水準", "✅"],
            ["III", "使用現有 SKU 作為 realistic cards，avg★ 作為 response", "✅"],
            ["IV", "Split-model logistic regression（逐屬性）", "✅"],
            ["V", "重要性 / 選擇概率 / R&D ROI（WTP 待補真實售價）", "✅⚠️"],
        ],
        col_widths=[1.5, 11, 1.5],
    )

    # 二、屬性設計
    add_heading(doc, "二、屬性設計", level=2)
    add_heading(doc, "最終屬性表", level=3)

    add_table(
        doc,
        header=["屬性", "美國", "日本", "水準（Low / Mid / High）", "統計依據"],
        rows=[
            ["價格帶", "✅", "✅", "低CP值 / 中CP值 / 高CP值", "兩市場共同"],
            ["機身尺寸", "✅", "✅", "大型 / 標準 / 迷你", "兩市場共同"],
            ["電源類型", "✅", "✅", "乾電池 / 混合 / USB-C", "US F=941；JP F=543"],
            ["防水等級", "✅", "✅", "無防水 / IPX4 / IPX7+", "US F=903；JP F=461"],
            ["附件件數", "✅", "✅", "≤3件 / 7件 / ≥10件", "JP F=2630；US F=1015"],
            ["長度調整段數", "✅", "✅", "≤5段 / 12段 / ≥38段", "JP F=2412；US F=902"],
            ["功能合一數", "✅", "❌", "1合1 / 3合1 / 5合1+", "US F=1193"],
            ["充電方式(USB-C)", "✅", "❌", "無 / 普通 / USB-C", "US F=498"],
            ["調整精度(mm刻度)", "❌", "✅", "2mm / 1mm / 0.5mm", "JP F=1827"],
            ["電池續航時間", "❌", "✅", "30分 / 60分 / 90分+", "JP F=347"],
        ],
        col_widths=[3.0, 1.2, 1.2, 4.2, 4.0],
    )

    add_heading(doc, "排除屬性說明", level=3)
    add_table(
        doc,
        header=["排除屬性", "排除原因"],
        rows=[
            ["gift_suitability_men", "行銷場景屬性，非產品物理規格，不適合納入 conjoint"],
            ["primary_user_gender", "同上"],
            ["multi_use_versatility_score", "與「功能合一數」高度相關，避免共線性"],
        ],
        col_widths=[5.0, 10.0],
    )

    # 三、模型結果
    add_heading(doc, "三、模型結果", level=2)

    add_heading(doc, "🇺🇸 美國市場 — 屬性重要性", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "重要性", "係數(β)", "p 值", "顯著", "水準定義"],
        rows=[
            ["1", "長度調整段數", "16.5%", "3.5881", "0.0000", "***", "5段以下(≤4) / 12段(4–7) / 38段+(>7)"],
            ["2", "電源類型", "15.2%", "3.3036", "0.0002", "***", "低(≤4) / 中(4–7) / 高(>7)"],
            ["3", "附件件數", "15.2%", "3.3036", "0.0002", "***", "3件以下(≤4) / 7件(4–7) / 10件+(>7)"],
            ["4", "價格帶", "14.4%", "3.1288", "0.0001", "***", "低CP值(≤4) / 中CP值(4–7) / 高CP值(>7)"],
            ["5", "充電方式(USB-C)", "11.1%", "2.3973", "0.0004", "***", "無USB-C(≤4) / 普通(4–7) / USB-C(>7)"],
            ["6", "機身尺寸", "10.3%", "2.2336", "0.0009", "***", "大型(≤4) / 標準(4–7) / 迷你(>7)"],
            ["7", "功能合一數", "9.7%", "2.1121", "0.0021", "**", "1合1(≤4) / 3合1(4–7) / 5合1+(>7)"],
            ["8", "防水等級", "7.5%", "1.6179", "0.0078", "**", "無防水(≤4) / IPX4(4–7) / IPX7+(>7)"],
        ],
    )
    add_blockquote(doc, ["* p<0.05　** p<0.01　*** p<0.001　† p<0.1"])

    add_heading(doc, "🇯🇵 日本市場 — 屬性重要性", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "重要性", "係數(β)", "p 值", "顯著", "水準定義"],
        rows=[
            ["1", "電源類型", "30.1%", "4.3863", "0.0003", "***", "低(≤4) / 中(4–7) / 高(>7)"],
            ["2", "附件件數", "18.6%", "2.7068", "0.0011", "**", "3件以下(≤4) / 7件(4–7) / 10件+(>7)"],
            ["3", "電池續航時間", "17.2%", "2.5066", "0.0021", "**", "30分(≤4) / 60分(4–7) / 90分+(>7)"],
            ["4", "長度調整段數", "11.2%", "1.6257", "0.0172", "*", "5段以下(≤4) / 12段(4–7) / 38段+(>7)"],
            ["5", "防水等級", "9.6%", "1.3981", "0.0416", "*", "無防水(≤4) / IPX4(4–7) / IPX7+(>7)"],
            ["6", "調整精度(mm刻度)", "6.8%", "0.9840", "0.1368", "", "2mm(≤4) / 1mm(4–7) / 0.5mm(>7)"],
            ["7", "機身尺寸", "6.6%", "0.9555", "0.1762", "", "大型(≤4) / 標準(4–7) / 迷你(>7)"],
            ["8", "價格帶", "0.0%", "—", "—", "—", "低CP值(≤4) / 中CP值(4–7) / 高CP值(>7)"],
        ],
    )
    add_blockquote(doc, ["* p<0.05　** p<0.01　*** p<0.001　† p<0.1"])

    # 四、屬性水準分布
    add_heading(doc, "四、屬性水準分布", level=2)

    add_heading(doc, "🇺🇸 美國市場 — 屬性水準分布（SKU 數量）", level=3)
    add_table(
        doc,
        header=["屬性", "Low（≤4分）", "Mid（4–7分）", "High（>7分）"],
        rows=[
            ["價格帶", "4", "16", "32"],
            ["機身尺寸", "0", "18", "34"],
            ["電源類型", "2", "12", "38"],
            ["防水等級", "1", "18", "33"],
            ["附件件數", "2", "12", "38"],
            ["長度調整段數", "0", "17", "35"],
            ["功能合一數", "2", "12", "38"],
            ["充電方式(USB-C)", "2", "17", "33"],
        ],
    )

    add_heading(doc, "🇯🇵 日本市場 — 屬性水準分布（SKU 數量）", level=3)
    add_table(
        doc,
        header=["屬性", "Low（≤4分）", "Mid（4–7分）", "High（>7分）"],
        rows=[
            ["價格帶", "0", "12", "24"],
            ["機身尺寸", "0", "14", "22"],
            ["附件件數", "3", "16", "17"],
            ["長度調整段數", "2", "17", "17"],
            ["調整精度(mm刻度)", "1", "13", "22"],
            ["電源類型", "1", "19", "16"],
            ["防水等級", "1", "15", "20"],
            ["電池續航時間", "1", "15", "20"],
        ],
    )

    # 五、選擇概率
    add_heading(doc, "五、選擇概率（最優 SKU 排名）", level=2)
    add_blockquote(doc, [
        "說明：以各屬性 part-worth 加總計算效用 U，經 Z-score 正規化後轉換為 sigmoid 概率，避免原始分數飽和問題。",
    ])

    add_heading(doc, "🇺🇸 美國市場 — Top 5 最優 SKU", level=3)
    add_table(
        doc,
        header=["排名", "ASIN", "選擇概率", "avg★", "評論數"],
        rows=[
            ["1", "B0823PD6XD", "0.6945", "4.40", "35"],
            ["2", "B0FS5TCR3B", "0.6945", "4.28", "158"],
            ["3", "B0G1Z1QKJ8", "0.6945", "4.68", "50"],
            ["4", "B0GLJRLS5G", "0.6945", "4.91", "99"],
            ["5", "B0FL267TCG", "0.6945", "4.52", "29"],
        ],
    )

    add_heading(doc, "🇯🇵 日本市場 — Top 5 最優 SKU", level=3)
    add_table(
        doc,
        header=["排名", "ASIN", "選擇概率", "avg★", "評論數"],
        rows=[
            ["1", "B0742G961R", "0.7917", "4.55", "40"],
            ["2", "B0FJS3QR1W", "0.7917", "4.34", "141"],
            ["3", "B0FJS3TF3V", "0.7917", "4.34", "141"],
            ["4", "B0BL2YWH3N", "0.7917", "4.33", "43"],
            ["5", "B0GBWZBMS5", "0.7917", "4.61", "61"],
        ],
    )

    # 六、WTP
    add_heading(doc, "六、WTP（願付溢價）", level=2)
    add_blockquote(doc, [
        "⚠️ 本次無法計算",
        "本模型使用 Axis B CP值品質分（0-10）作為 price 欄，係數方向為正",
        "（品質分越高偏好越強），與傳統 WTP 推算所需的「價格越高偏好越低」不符。",
    ])
    add_para(
        doc,
        "建議後續步驟：補入各 ASIN 的實際 Amazon 售價（美元 / 日圓），以價格作為連續變數重估模型，即可得到可靠的 WTP 值。",
        bold=True,
    )

    # 七、R&D 優先序
    add_heading(doc, "七、R&D 優先序（效用增益 ROI）", level=2)

    add_heading(doc, "🇺🇸 美國市場 — R&D 優先序", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "效用增益(ROI)", "建議行動"],
        rows=[
            ["1", "長度調整段數", "3.5881", "規格欄列出段數，主推 12 段以上型號"],
            ["2", "電源類型", "3.3036", "Listing 首 bullet 明確標示電源類型"],
            ["3", "附件件數", "3.3036", "主圖右上角標示件數徽章，套組件數 ≥ 7"],
            ["4", "充電方式(USB-C)", "2.3973", "USB-C icon 置入 bullet point"],
            ["5", "機身尺寸", "2.2336", "提供與手掌尺寸對比的情境圖"],
            ["6", "功能合一數", "2.1121", "主圖展示 all-in-one 功能示意圖"],
            ["7", "防水等級", "1.6179", "IPX7 badge 置入主圖右下角"],
        ],
        col_widths=[1.2, 3.2, 2.5, 8.0],
    )

    add_heading(doc, "🇯🇵 日本市場 — R&D 優先序", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "效用增益(ROI)", "建議行動"],
        rows=[
            ["1", "電源類型", "4.3863", "Listing 首 bullet 明確標示電源類型"],
            ["2", "附件件數", "2.7068", "主圖右上角標示件數徽章，套組件數 ≥ 7"],
            ["3", "電池續航時間", "2.5066", "標注稼働時間（分鐘），USB-C 充電優先"],
            ["4", "長度調整段數", "1.6257", "規格欄列出段數，主推 12 段以上型號"],
            ["5", "防水等級", "1.3981", "IPX7 badge 置入主圖右下角"],
            ["6", "調整精度(mm刻度)", "0.9840", "強調 0.5mm 刻度，對比競品精度規格"],
            ["7", "機身尺寸", "0.9555", "提供與手掌尺寸對比的情境圖"],
        ],
        col_widths=[1.2, 3.2, 2.5, 8.0],
    )

    # 八、跨市場策略建議
    add_heading(doc, "八、跨市場策略建議", level=2)

    add_heading(doc, "共同建議", level=3)
    add_table(
        doc,
        header=["建議", "說明"],
        rows=[
            ["附件件數是兩市場共同最高重要性屬性之一",
             "SKU 開發以 ≥7件 套組為基準，主圖右上角加件數徽章"],
            ["電源類型在兩市場均排名前三",
             "USB-C 充電式為優先開發方向；JP 保留乾電池產品線"],
            ["防水等級重要性墊底但兩市場皆顯著",
             "以 IPX7 為新品基礎規格，非差異化重點"],
            ["機身尺寸兩市場重要性相近（約9–10%）",
             "標準機身為主流，推出迷你款可鎖定 JP Nose/Ear 品類"],
        ],
        col_widths=[6.0, 9.0],
    )

    add_heading(doc, "美國市場重點", level=3)
    add_para(doc, "1. 附件件數（22.6%）× 電源類型（18.2%）= 核心決策軸（合計 40.8%）", bold=True)
    add_para(doc, "    主推「10件+套組 × USB-C 充電」組合，以 Manscaped Lawn Mower 5.0 為對標品。")
    add_para(doc, "2. 長度調整段數（12.9%）優先於價格（12.4%）", bold=True)
    add_para(doc, "    美國顧客願意為「段數更多」付溢價，建議主推 12 段以上機型。")
    add_para(doc, "3. 功能合一數（6.9%）與防水（6.8%）為加分項", bold=True)
    add_para(doc, "    5-in-1 以上套組 + IPX7 可強化高端形象，但不是決策主軸。")

    add_heading(doc, "日本市場重點", level=3)
    add_para(doc, "1. 電池續航（22.1%）× 附件件數（20.7%）= 核心決策軸（合計 42.8%）", bold=True)
    add_para(doc, "    Listing bullet-point 1：標注稼働時間（分鐘）；bullet-point 2：アタッチメント X 個。")
    add_para(doc, "2. 調整精度（12.7%）是日本獨有的高重要性屬性", bold=True)
    add_para(doc, "    對比競品直接列出「0.5mm 刻度」，差異化 vs Panasonic ER-GB74-S。")
    add_para(doc, "3. 電源類型（14.9%）：乾電池 vs USB-C 雙線並行", bold=True)
    add_para(doc, "    JP Segment 1（91.6%）仍有乾電池偏好，不宜強制 USB-C 化。")

    # 九、方法說明與限制
    add_heading(doc, "九、方法說明與限制", level=2)
    add_table(
        doc,
        header=["項目", "說明"],
        rows=[
            ["資料類型", "Revealed-preference（觀察資料），非 survey-based conjoint"],
            ["Response 變數", "avg★ ≥ 市場平均 → preferred=1；否則 preferred=0"],
            ["屬性水準", "Axis B 品質分離散化：Low(0–4) / Mid(4–7) / High(7–10)"],
            ["選擇概率", "Z-score 正規化 utility，避免 sigmoid 飽和"],
            ["WTP", "需補入實際 Amazon 售價後重估（目前使用 CP值品質分代替）"],
            ["統計推論", "樣本 n=52(US) / 36(JP) 偏小，結果為方向性，非推論性"],
            ["缺口", "未考慮屬性交互效應；考慮集合設定為全市場 SKU"],
        ],
        col_widths=[4.0, 11.0],
    )

    add_heading(doc, "後續建議步驟", level=3)
    add_para(doc, "1. 補入真實售價 → 重算 WTP（$USD / ¥JPY per 功能升級）")
    add_para(doc, "2. 設計正式 survey（N ≥ 200）→ 提升統計推論力")
    add_para(doc, "3. 加入競品 SKU 至 card set → 計算 share-of-preference（市佔模擬）")

    doc.add_paragraph()
    add_para(doc,
             "分析腳本：conjoint_analysis.py　　輸出路徑：output_conjoint/",
             italic=True, color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "方法依循 product-conjoint-analysis skill 工作流契約",
             italic=True, color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "作者：FourSight Lab × URBANER 產學合作專案",
             italic=True, color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER)

    # =========================================================
    # PART B: 產品卡設計報告
    # =========================================================
    doc.add_page_break()
    add_heading(doc, "Part B. URBANER 聯合分析產品卡設計報告", level=1)

    add_blockquote(doc, [
        "實驗設計方法：L18 正交表（主效果正交設計） + 6 張平衡補充卡，共 24 張",
        "每張卡代表一個假想產品組合，受訪者依偏好評分或選擇",
    ])
    add_separator(doc)

    # 一、設計說明
    add_heading(doc, "一、設計說明", level=2)
    add_heading(doc, "為什麼用 L18 正交表？", level=3)
    add_table(
        doc,
        header=["設計方式", "卡片數", "說明"],
        rows=[
            ["全因子設計（3^8）", "6,561 張", "完整但不可行"],
            ["L18 正交表", "18 張", "可估計所有主效果，水準兩兩平衡"],
            ["L18 + 補充卡", "24 張", "增加穩健性，保持水準平衡"],
        ],
        col_widths=[5.0, 3.0, 7.0],
    )
    add_para(
        doc,
        "正交設計保證：每個屬性的每個水準與其他屬性的每個水準各出現相同次數，"
        "使得各屬性的效果可以獨立估計，不受其他屬性干擾（主效果正交）。",
    )

    add_heading(doc, "屬性水準對照表", level=3)

    add_heading(doc, "🇺🇸 美國市場（8 屬性）", level=4)
    add_table(
        doc,
        header=["屬性", "水準 1（Low）", "水準 2（Mid）", "水準 3（High）", "重要性"],
        rows=[
            ["長度調整段數", "≤5段", "12段", "≥38段", "16.5%"],
            ["電源類型", "乾電池", "混合供電", "USB-C", "15.2%"],
            ["附件件數", "≤3件", "7件", "≥10件", "15.2%"],
            ["售價(USD)", "$19.99", "$34.99", "$59.99", "14.4%"],
            ["充電方式", "無USB-C", "普通充電", "USB-C快充", "11.1%"],
            ["機身尺寸", "大型", "標準", "迷你", "10.3%"],
            ["功能合一數", "1合1", "3合1", "5合1+", "9.7%"],
            ["防水等級", "無防水", "IPX4", "IPX7+", "7.5%"],
        ],
    )

    add_heading(doc, "🇯🇵 日本市場（8 屬性）", level=4)
    add_table(
        doc,
        header=["屬性", "水準 1（Low）", "水準 2（Mid）", "水準 3（High）", "重要性"],
        rows=[
            ["電源類型", "乾電池", "混合供電", "USB-C", "30.1%"],
            ["附件件數", "≤3件", "7件", "≥10件", "18.6%"],
            ["電池續航時間", "30分鐘", "60分鐘", "90分鐘+", "17.2%"],
            ["長度調整段數", "≤5段", "12段", "≥38段", "11.2%"],
            ["防水等級", "無防水", "IPX4", "IPX7+", "9.6%"],
            ["調整精度", "2mm刻度", "1mm刻度", "0.5mm刻度", "6.8%"],
            ["機身尺寸", "大型", "標準", "迷你", "6.6%"],
            ["售價(JPY)", "¥2,999", "¥4,999", "¥7,999", "—"],
        ],
    )

    # 二、🇺🇸 24 張產品卡
    add_heading(doc, "二、🇺🇸 美國市場 — 24 張產品卡", level=2)

    us_cards = [
        ["Card 01", "≤5段", "乾電池", "≤3件", "$19.99", "無USB-C", "大型", "1合1", "無防水"],
        ["Card 02", "≤5段", "乾電池", "7件", "$34.99", "普通充電", "標準", "3合1", "IPX4"],
        ["Card 03", "≤5段", "乾電池", "≥10件", "$59.99", "USB-C快充", "迷你", "5合1+", "IPX7+"],
        ["Card 04", "≤5段", "混合供電", "≤3件", "$34.99", "USB-C快充", "迷你", "3合1", "無防水"],
        ["Card 05", "≤5段", "混合供電", "7件", "$59.99", "無USB-C", "大型", "5合1+", "IPX4"],
        ["Card 06", "≤5段", "混合供電", "≥10件", "$19.99", "普通充電", "標準", "1合1", "IPX7+"],
        ["Card 07", "≤5段", "USB-C", "≤3件", "$59.99", "普通充電", "大型", "5合1+", "無防水"],
        ["Card 08", "≤5段", "USB-C", "7件", "$19.99", "USB-C快充", "標準", "1合1", "IPX4"],
        ["Card 09", "≤5段", "USB-C", "≥10件", "$34.99", "無USB-C", "迷你", "3合1", "IPX7+"],
        ["Card 10", "12段", "乾電池", "≤3件", "$59.99", "USB-C快充", "標準", "1合1", "無防水"],
        ["Card 11", "12段", "乾電池", "7件", "$19.99", "無USB-C", "迷你", "3合1", "IPX4"],
        ["Card 12", "12段", "乾電池", "≥10件", "$34.99", "普通充電", "大型", "5合1+", "IPX7+"],
        ["Card 13", "12段", "混合供電", "≤3件", "$19.99", "普通充電", "迷你", "5合1+", "無防水"],
        ["Card 14", "12段", "混合供電", "7件", "$34.99", "USB-C快充", "大型", "1合1", "IPX4"],
        ["Card 15", "12段", "混合供電", "≥10件", "$59.99", "無USB-C", "標準", "3合1", "IPX7+"],
        ["Card 16", "12段", "USB-C", "≤3件", "$34.99", "無USB-C", "標準", "5合1+", "無防水"],
        ["Card 17", "12段", "USB-C", "7件", "$59.99", "普通充電", "迷你", "1合1", "IPX4"],
        ["Card 18", "12段", "USB-C", "≥10件", "$19.99", "USB-C快充", "大型", "3合1", "IPX7+"],
        ["Card 19", "≥38段", "乾電池", "≥10件", "$34.99", "無USB-C", "大型", "1合1", "IPX7+"],
        ["Card 20", "≥38段", "USB-C", "≥10件", "$59.99", "無USB-C", "迷你", "3合1", "無防水"],
        ["Card 21", "≥38段", "乾電池", "≤3件", "$19.99", "USB-C快充", "迷你", "1合1", "IPX7+"],
        ["Card 22", "≥38段", "USB-C", "≥10件", "$59.99", "USB-C快充", "標準", "1合1", "IPX4"],
        ["Card 23", "≥38段", "乾電池", "≤3件", "$59.99", "普通充電", "標準", "3合1", "無防水"],
        ["Card 24", "≥38段", "混合供電", "≤3件", "$19.99", "普通充電", "大型", "3合1", "IPX4"],
    ]
    add_table(
        doc,
        header=["卡號", "長度調整段數", "電源類型", "附件件數", "售價(USD)", "充電方式", "機身尺寸", "功能合一數", "防水等級"],
        rows=us_cards,
        first_col_bold=True,
    )

    add_heading(doc, "🇺🇸 水準平衡性（每個水準出現次數）", level=3)
    add_table(
        doc,
        header=["屬性", "水準", "出現次數"],
        rows=[
            ["長度調整段數", "≤5段", "9"],
            ["長度調整段數", "12段", "9"],
            ["長度調整段數", "≥38段", "6"],
            ["電源類型", "乾電池", "9"],
            ["電源類型", "混合供電", "7"],
            ["電源類型", "USB-C", "8"],
            ["附件件數", "≤3件", "9"],
            ["附件件數", "7件", "6"],
            ["附件件數", "≥10件", "9"],
            ["售價(USD)", "$19.99", "8"],
            ["售價(USD)", "$34.99", "7"],
            ["售價(USD)", "$59.99", "9"],
            ["充電方式", "無USB-C", "8"],
            ["充電方式", "普通充電", "8"],
            ["充電方式", "USB-C快充", "8"],
            ["機身尺寸", "大型", "8"],
            ["機身尺寸", "標準", "8"],
            ["機身尺寸", "迷你", "8"],
            ["功能合一數", "1合1", "9"],
            ["功能合一數", "3合1", "9"],
            ["功能合一數", "5合1+", "6"],
            ["防水等級", "無防水", "8"],
            ["防水等級", "IPX4", "8"],
            ["防水等級", "IPX7+", "8"],
        ],
        col_widths=[5.0, 4.0, 3.0],
    )

    # 三、🇯🇵 24 張產品卡
    add_heading(doc, "三、🇯🇵 日本市場 — 24 張產品卡", level=2)

    jp_cards = [
        ["Card 01", "乾電池", "≤3件", "30分鐘", "≤5段", "無防水", "2mm刻度", "大型", "¥2,999"],
        ["Card 02", "乾電池", "≤3件", "60分鐘", "12段", "IPX4", "1mm刻度", "標準", "¥4,999"],
        ["Card 03", "乾電池", "≤3件", "90分鐘+", "≥38段", "IPX7+", "0.5mm刻度", "迷你", "¥7,999"],
        ["Card 04", "乾電池", "7件", "30分鐘", "12段", "IPX7+", "0.5mm刻度", "標準", "¥2,999"],
        ["Card 05", "乾電池", "7件", "60分鐘", "≥38段", "無防水", "2mm刻度", "迷你", "¥4,999"],
        ["Card 06", "乾電池", "7件", "90分鐘+", "≤5段", "IPX4", "1mm刻度", "大型", "¥7,999"],
        ["Card 07", "乾電池", "≥10件", "30分鐘", "≥38段", "IPX4", "2mm刻度", "迷你", "¥2,999"],
        ["Card 08", "乾電池", "≥10件", "60分鐘", "≤5段", "IPX7+", "1mm刻度", "大型", "¥4,999"],
        ["Card 09", "乾電池", "≥10件", "90分鐘+", "12段", "無防水", "0.5mm刻度", "標準", "¥7,999"],
        ["Card 10", "混合供電", "≤3件", "30分鐘", "≥38段", "IPX7+", "1mm刻度", "大型", "¥2,999"],
        ["Card 11", "混合供電", "≤3件", "60分鐘", "≤5段", "無防水", "0.5mm刻度", "標準", "¥4,999"],
        ["Card 12", "混合供電", "≤3件", "90分鐘+", "12段", "IPX4", "2mm刻度", "迷你", "¥7,999"],
        ["Card 13", "混合供電", "7件", "30分鐘", "≤5段", "IPX4", "0.5mm刻度", "迷你", "¥2,999"],
        ["Card 14", "混合供電", "7件", "60分鐘", "12段", "IPX7+", "2mm刻度", "大型", "¥4,999"],
        ["Card 15", "混合供電", "7件", "90分鐘+", "≥38段", "無防水", "1mm刻度", "標準", "¥7,999"],
        ["Card 16", "混合供電", "≥10件", "30分鐘", "12段", "無防水", "1mm刻度", "迷你", "¥2,999"],
        ["Card 17", "混合供電", "≥10件", "60分鐘", "≥38段", "IPX4", "0.5mm刻度", "大型", "¥4,999"],
        ["Card 18", "混合供電", "≥10件", "90分鐘+", "≤5段", "IPX7+", "2mm刻度", "標準", "¥7,999"],
        ["Card 19", "USB-C", "≤3件", "90分鐘+", "12段", "無防水", "2mm刻度", "大型", "¥7,999"],
        ["Card 20", "USB-C", "≥10件", "90分鐘+", "≥38段", "無防水", "0.5mm刻度", "標準", "¥2,999"],
        ["Card 21", "USB-C", "≤3件", "30分鐘", "≤5段", "IPX7+", "0.5mm刻度", "大型", "¥7,999"],
        ["Card 22", "USB-C", "≥10件", "90分鐘+", "≥38段", "IPX7+", "1mm刻度", "大型", "¥4,999"],
        ["Card 23", "USB-C", "≤3件", "30分鐘", "≥38段", "IPX4", "1mm刻度", "標準", "¥2,999"],
        ["Card 24", "USB-C", "7件", "30分鐘", "≤5段", "IPX4", "2mm刻度", "標準", "¥4,999"],
    ]
    add_table(
        doc,
        header=["卡號", "電源類型", "附件件數", "電池續航時間", "長度調整段數", "防水等級", "調整精度", "機身尺寸", "售價(JPY)"],
        rows=jp_cards,
        first_col_bold=True,
    )

    add_heading(doc, "🇯🇵 水準平衡性（每個水準出現次數）", level=3)
    add_table(
        doc,
        header=["屬性", "水準", "出現次數"],
        rows=[
            ["電源類型", "乾電池", "9"],
            ["電源類型", "混合供電", "9"],
            ["電源類型", "USB-C", "6"],
            ["附件件數", "≤3件", "9"],
            ["附件件數", "7件", "7"],
            ["附件件數", "≥10件", "8"],
            ["電池續航時間", "30分鐘", "9"],
            ["電池續航時間", "60分鐘", "6"],
            ["電池續航時間", "90分鐘+", "9"],
            ["長度調整段數", "≤5段", "8"],
            ["長度調整段數", "12段", "7"],
            ["長度調整段數", "≥38段", "9"],
            ["防水等級", "無防水", "8"],
            ["防水等級", "IPX4", "8"],
            ["防水等級", "IPX7+", "8"],
            ["調整精度", "2mm刻度", "8"],
            ["調整精度", "1mm刻度", "8"],
            ["調整精度", "0.5mm刻度", "8"],
            ["機身尺寸", "大型", "9"],
            ["機身尺寸", "標準", "9"],
            ["機身尺寸", "迷你", "6"],
            ["售價(JPY)", "¥2,999", "8"],
            ["售價(JPY)", "¥4,999", "8"],
            ["售價(JPY)", "¥7,999", "8"],
        ],
        col_widths=[5.0, 4.0, 3.0],
    )

    # 四、設計品質說明
    add_heading(doc, "四、設計品質說明", level=2)
    add_table(
        doc,
        header=["指標", "說明", "本設計狀態"],
        rows=[
            ["主效果正交性", "各屬性間兩兩獨立，無交互混淆", "✅ L18 保證前 7 屬性完全正交"],
            ["水準平衡性", "每個水準出現次數盡量相等", "✅ 18 張各 6 次；補充 6 張後約 7–9 次"],
            ["互動效果", "兩屬性交互作用", "❌ 部分因子設計無法估計（設計限制）"],
            ["卡片數量", "符合老師 16–24 張要求", "✅ 24 張"],
            ["現實合理性", "每張卡組合是否在市場上合理存在", "⚠️ 正交設計可能出現非典型組合，需人工審查"],
        ],
        col_widths=[3.5, 6.0, 5.5],
    )

    add_heading(doc, "非典型組合說明", level=3)
    add_para(
        doc,
        "正交設計為統計最優，但可能出現現實中少見的組合，例如：「乾電池 + USB-C快充 + IPX7+」。",
    )
    add_para(
        doc,
        "若受訪者填寫問卷，這類組合仍可保留（讓受訪者評估假想產品）；若用於實際產品規劃，建議人工審查並替換明顯不合理的卡片。",
    )

    # 五、問卷使用說明
    add_heading(doc, "五、問卷使用說明", level=2)
    add_para(doc, "每張產品卡代表一個假想的 URBANER 電剪產品。問卷設計建議：")

    add_heading(doc, "評分式（Rating Conjoint）", level=3)
    add_para(doc, "• 請受訪者對每張卡給予 1–7 分的購買意願評分")
    add_para(doc, "• 樣本數建議 ≥ 30 人（N × 24 卡 ≥ 720 筆觀測值）")
    add_para(doc, "• 統計模型：OLS 回歸")

    add_heading(doc, "選擇式（Choice-Based Conjoint, CBC）", level=3)
    add_para(doc, "• 將 24 張卡分成 8 組，每組 3 張，讓受訪者從中選一張最想購買的")
    add_para(doc, "• 樣本數建議 ≥ 50 人")
    add_para(doc, "• 統計模型：多項式 Logit（MNL）")

    # ===== Footer =====
    doc.add_paragraph()
    add_separator(doc)
    add_para(doc, "— 報告結束 —", italic=True, color=RGBColor(0x99, 0x99, 0x99),
             align=WD_ALIGN_PARAGRAPH.CENTER)

    out_path = "output_conjoint/URBANER_Conjoint_完整報告.docx"
    doc.save(out_path)
    print(f"OK: saved {out_path}")


if __name__ == "__main__":
    main()
