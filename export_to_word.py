"""
將 extended_analysis_report.md 轉換為 Word 文件
支援：標題階層、表格、粗體/斜體/程式碼、圖片嵌入、說明文字
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHART_DIR  = Path("output_conjoint/charts")

JOBS = [
    {
        "md":     Path("output_conjoint/extended_analysis_report.md"),
        "output": Path("output_conjoint/URBANER_延伸分析報告.docx"),
    },
    {
        "md":     Path("output_conjoint/conjoint_report.md"),
        "output": Path("output_conjoint/URBANER_聯合分析報告.docx"),
    },
    {
        "md":     Path("output_conjoint/card_choice_prob_report.md"),
        "output": Path("output_conjoint/URBANER_產品組合購買機率.docx"),
    },
    {
        "md":     Path("output_conjoint/card_mnl_market_report.md"),
        "output": Path("output_conjoint/URBANER_市場競爭力分析.docx"),
    },
]

# ── 顏色定義 ──────────────────────────────────────────────────────────────────
COLOR_BRAND   = RGBColor(0xE8, 0x47, 0x2A)   # Urbaner 紅
COLOR_HEADING = RGBColor(0x1A, 0x2E, 0x44)   # 深藍黑
COLOR_SUBHEAD = RGBColor(0x2C, 0x5F, 0x8A)   # 中藍
COLOR_TABLE_H = RGBColor(0x1A, 0x2E, 0x44)   # 表頭底色
COLOR_TABLE_R = RGBColor(0xF0, 0xF4, 0xF8)   # 交替列底色
COLOR_NOTE    = RGBColor(0x88, 0x88, 0x88)   # 說明文字灰

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def apply_inline(run, text: str):
    """解析 **bold**、`code`、*italic* 後套用格式"""
    run.text = text


def add_inline_para(para, text: str):
    """
    把含有 **bold**、`code`、*italic* 的行拆成多個 run 加入段落。
    """
    # 移除開頭的 > 引言符號（已在外層處理）
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        run = para.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BRAND
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part


def build_doc(md_text: str) -> Document:
    doc = Document()

    # ── 頁面設定 ──────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── 預設段落字型 ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── 水平線 ────────────────────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pb = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pb.append(bottom)
            pPr.append(pb)
            i += 1
            continue

        # ── 標題 ─────────────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = COLOR_HEADING
            run.font.name = "Microsoft YaHei"
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.space_before = Pt(6)
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = COLOR_SUBHEAD
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = COLOR_BRAND
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        if line.startswith("#### "):
            text = line[5:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = COLOR_HEADING
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── 圖片（Markdown image）────────────────────────────────────────────
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_match:
            img_alt  = img_match.group(1)
            img_path = CHART_DIR.parent / img_match.group(2)
            if img_path.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(img_path), width=Inches(6.0))
                # 圖說
                cap = doc.add_paragraph(img_alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = COLOR_NOTE
                cap.runs[0].italic = True
                cap.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # ── 程式碼區塊 ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.8)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 灰底
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F4F8")
            pPr.append(shd)
            continue

        # ── 引言（blockquote > ）─────────────────────────────────────────────
        if line.startswith("> "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(0.8)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            add_inline_para(para, text)
            for run in para.runs:
                run.font.color.rgb = COLOR_NOTE
                run.font.size = Pt(9.5)
                run.italic = True
            i += 1
            continue

        # ── 表格 ─────────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # 過濾分隔列
            data_rows = [r for r in table_lines
                         if not re.match(r"^\s*\|[-:\s|]+\|\s*$", r)]
            if not data_rows:
                continue

            def parse_row(row):
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                return cells

            rows = [parse_row(r) for r in data_rows]
            ncols = max(len(r) for r in rows)
            # 補齊欄數
            rows = [r + [""] * (ncols - len(r)) for r in rows]

            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row_data in enumerate(rows):
                row_obj = table.rows[ri]
                for ci, cell_text in enumerate(row_data):
                    cell = row_obj.cells[ci]
                    # 清空預設內容
                    cell.paragraphs[0].clear()
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)
                    # 移除 markdown bold
                    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                    clean = re.sub(r"`([^`]+)`", r"\1", clean)
                    add_inline_para(para, clean)

                    if ri == 0:
                        set_cell_bg(cell, COLOR_TABLE_H)
                        for run in para.runs:
                            run.font.bold  = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size  = Pt(9.5)
                    else:
                        if ri % 2 == 0:
                            set_cell_bg(cell, COLOR_TABLE_R)
                        for run in para.runs:
                            run.font.size = Pt(9.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            doc.add_paragraph()
            continue

        # ── 無序清單 ─────────────────────────────────────────────────────────
        if re.match(r"^[-*] ", line):
            text = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(0.6)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_para(para, text)
            for run in para.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── 有序清單 ─────────────────────────────────────────────────────────
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent  = Cm(0.6)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_para(para, text)
            for run in para.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── 空行 ─────────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── 一般段落 ─────────────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        add_inline_para(para, line.strip())
        for run in para.runs:
            run.font.size = Pt(10.5)
        i += 1

    return doc


def main():
    for job in JOBS:
        md_path = job["md"]
        output  = job["output"]
        if not md_path.exists():
            print(f"[SKIP] 找不到 {md_path}")
            continue
        print(f"轉換中：{md_path.name} ...")
        md_text = md_path.read_text(encoding="utf-8")
        doc = build_doc(md_text)
        doc.save(output)
        print(f"[OK] Word 文件已儲存：{output}")


if __name__ == "__main__":
    main()
