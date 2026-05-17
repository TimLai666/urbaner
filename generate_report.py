"""
Generate URBANER Amazon 雙市場分析報告 (Word .docx)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "output/URBANER_研究報告.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in borders.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color="1A2744"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
    return p

def add_body(doc, text, indent=False, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table_row(table, cells_data, header=False, bg=None):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.bold = header
            if header:
                run.font.color.rgb = RGBColor(255, 255, 255)
        if bg:
            set_cell_bg(cell, bg)
    return row

# ── main document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "微軟正黑體"
doc.styles["Normal"].font.size = Pt(11)

# ─── Cover ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("URBANER 奧本")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
title_run.font.name = "微軟正黑體"

sub1_p = doc.add_paragraph()
sub1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1_run = sub1_p.add_run("Amazon 美日雙市場電剪分析報告")
sub1_run.font.size = Pt(20)
sub1_run.font.bold = True
sub1_run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x6F)
sub1_run.font.name = "微軟正黑體"

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_run = sub2_p.add_run("研究脈絡・架構・核心結論")
sub2_run.font.size = Pt(14)
sub2_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
sub2_run.font.name = "微軟正黑體"

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run(
    "Biz Connect Taipei 2026 ・ 台北市數位產學實作計畫\n"
    "合作廠商：昆豪企業股份有限公司 ・ 品牌：URBANER\n"
    "報告日期：2026年5月17日"
)
meta_run.font.size = Pt(11)
meta_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
meta_run.font.name = "微軟正黑體"

doc.add_page_break()

# ─── CHAPTER 1：研究脈絡 ─────────────────────────────────────────────────────

add_heading(doc, "一、研究脈絡", level=1)

add_body(doc,
    "昆豪企業（統編 33778451）於 1977 年創立，長期以 OEM 代工模式供應歐、美、日品牌電動毛髮修剪器。"
    "2014 年接班後，決定轉型 OBM 自有品牌，並於同年創立「URBANER 奧本」品牌，目標直接進入美國、日本的 Amazon 平台零售市場。"
    "至 2020 年，品牌營收已正式超越代工業務，顯示轉型成功。"
)

add_body(doc,
    "然而，OBM 轉型帶來新課題：相較於 OEM 只需配合客戶規格開發產品，OBM 必須自行判斷「消費者真正在乎什麼」、"
    "「競爭對手的優勢在哪裡」、「如何精準定價與行銷」。這些問題無法靠直覺回答，需要系統性的數據分析。"
)

add_heading(doc, "1.1 為什麼做這個報告？", level=2)

add_body(doc, "本報告的核心問題是：", bold=True)
add_bullet(doc, "美國與日本的電剪消費者，分別最在乎什麼屬性？")
add_bullet(doc, "URBANER 現有產品是否對準這些需求？")
add_bullet(doc, "在眾多競品中，URBANER 的市佔率有多少？哪裡有機會？")
add_bullet(doc, "每個市場應該主推哪款產品、用什麼行銷語言？")

add_body(doc,
    "要回答這些問題，不能只看自家銷售數字。必須同時分析競品評論、市場偏好、消費者分群，"
    "並以統計方法量化「消費者願意為哪些特性多付多少錢」。這正是本研究的出發點。"
)

add_heading(doc, "1.2 分析範疇", level=2)

# 範疇 table
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, txt in enumerate(["維度", "美國市場（US）", "日本市場（JP）"]):
    hdr[i].text = txt
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(hdr[i], "1A2744")

rows_data = [
    ("平台", "Amazon.com", "Amazon.co.jp ＋ 楽天市場"),
    ("產品類別", "9 大類（鬍鬚/鼻耳/體毛/眉毛/電刮刀/寵物/指甲等）", "同左"),
    ("競品資料", "11,799 則評論 × 50 個競品 ASIN", "同範疇另行抽樣"),
    ("URBANER 訂單", "4,184 件（2024-03 起）", "16,744 件（2023-01 起）"),
    ("屬性矩陣", "88 個 SKU × 114 個產品屬性", "同左"),
    ("分析期間", "2024–2026 年資料", "2023–2026 年資料"),
]

for row_vals in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(row_vals):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

doc.add_page_break()

# ─── CHAPTER 2：研究架構 ─────────────────────────────────────────────────────

add_heading(doc, "二、研究架構（MMR 四階段方法論）", level=1)

add_body(doc,
    "本研究採用 Market Model Research（MMR）架構，將整個分析流程分為四個階段，"
    "從原始數據逐步提煉為可直接操作的行銷劇本。以下說明各階段的邏輯與方法。"
)

# Phase 1
add_heading(doc, "第一階段｜資料整合", level=2)

add_body(doc,
    "整合三大資料來源：Amazon 競品評論（11,799 則）、競品每月銷量銷售額、URBANER 自身訂單。"
    "透過結構化清洗，建立 88 SKU × 114 屬性的分析矩陣，作為後續所有分析的基礎。"
)

add_body(doc, "資料來源：", bold=True)
add_bullet(doc, "rawdata_Urbaner/amazon_reviews/ — 競品評論（9 類 × xlsx）")
add_bullet(doc, "rawdata_Urbaner/competitor_sales/ — 競品月銷量")
add_bullet(doc, "rawdata_Urbaner/amazon_sales/US_sales/ — URBANER 美國訂單")
add_bullet(doc, "rawdata_Urbaner/amazon_sales/JP_sales/ — URBANER 日本訂單（shift-jis 編碼）")

# Phase 2
add_heading(doc, "第二階段｜Review-Mining STP 市場細分", level=2)

add_body(doc,
    "以評論文本 + 產品屬性為輸入，對競品與 URBANER 共 11,523 則評論進行結構化分析。"
    "核心方法為雙軸評估：Axis A（屬性顯著度，0–7 分）× Axis B（品質感知，0–10 分），"
    "搭配 ANOVA 統計驗證，找出兩個市場中各自最影響購買決策的產品屬性排名。"
)

add_body(doc, "此階段回答：「消費者最在乎什麼？」", bold=True, italic=True)

add_bullet(doc, "美國 Top 屬性：禮品適合度（F=1,299）、USB-C 充電（F=975）、IPX7 防水（F=734）")
add_bullet(doc, "日本 Top 屬性：附件件數（F=2,630）、長度調整精度（F=1,847）、電源類型（F=1,203）")

add_body(doc,
    "K-Means + PCA 聚類分析（美日各分為 3 個消費者群體），識別出高價值目標族群，"
    "例如美國 S1（禮品購買族，佔 62.4%）與日本 S1（精準修剪需求族，佔 91.6%）。"
)

# Phase 3
add_heading(doc, "第三階段｜Logistic Conjoint 偏好量化", level=2)

add_body(doc,
    "採用 Revealed-Preference Logistic Conjoint 分析，以消費者實際購買行為（而非問卷）"
    "推導出各屬性對購買機率的影響係數（β），進一步計算：",
)

add_bullet(doc, "屬性重要度（Importance）：量化哪些特性驅動購買")
add_bullet(doc, "消費者願付溢價（WTP）：β_attribute / |β_price|，以金額表達屬性價值")
add_bullet(doc, "市佔率模擬（MNL Share-of-Preference）：URBANER vs 競品在現有定位下的理論市佔率")

add_body(doc, "此階段回答：「消費者願意為哪些特性多付多少錢？」", bold=True, italic=True)

# Phase 4
add_heading(doc, "第四階段｜行銷劇本與決策支援", level=2)

add_body(doc,
    "將前三階段的統計結論轉化為可執行的行銷決策，產出兩大工具：",
)

add_bullet(doc, "雙市場戰略儀表板（dashboard.py）：整合 STP、Conjoint、WTP、競品分析，以視覺化呈現")
add_bullet(doc, "11 條行銷劇本（Playbook）：針對美日兩市場各 4 條 + 跨市場共用 3 條，每條附有具體行銷文案")

doc.add_paragraph()

# Architecture diagram (text-based)
add_heading(doc, "2.1 研究架構總覽", level=2)

arch_table = doc.add_table(rows=1, cols=4)
arch_table.style = "Table Grid"

phase_data = [
    ("INPUT\n輸入", "評論資料\n競品銷售\n自身訂單\n114 個屬性", "2E5BFF"),
    ("PROCESS ①\nReview-Mining STP", "ANOVA 屬性排名\nK-Means 分群\n雙軸評估矩陣\n消費者輪廓", "C9A36F"),
    ("PROCESS ②\nLogistic Conjoint", "屬性重要度\nWTP 溢價計算\nMNL 市佔模擬\n最佳化組合", "C9A36F"),
    ("OUTPUT\n輸出", "行銷劇本×11\n儀表板\n產品優化建議\n定價策略", "D32F4D"),
]

for i, (title, body, color) in enumerate(phase_data):
    cell = arch_table.rows[0].cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p.add_run(title)
    title_run.font.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    p.add_run("\n")
    body_run = p.add_run(body)
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(cell, color)

doc.add_paragraph()
doc.add_page_break()

# ─── CHAPTER 3：核心結論 ─────────────────────────────────────────────────────

add_heading(doc, "三、核心結論", level=1)

add_heading(doc, "3.1 美日市場需求差異", level=2)

add_body(doc,
    "分析結果顯示，美國與日本消費者的購買驅動力存在根本差異，不能使用同一套行銷策略："
)

diff_table = doc.add_table(rows=1, cols=3)
diff_table.style = "Table Grid"

diff_hdr = diff_table.rows[0].cells
for i, txt in enumerate(["維度", "美國市場 🇺🇸", "日本市場 🇯🇵"]):
    diff_hdr[i].text = txt
    p = diff_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(diff_hdr[i], "1A2744")

diff_rows = [
    ("最強驅動屬性", "禮品適合度（F=1,299）", "附件件數（F=2,630）"),
    ("第二驅動", "USB-C 充電（F=975）", "長度調整精度（F=1,847）"),
    ("第三驅動", "IPX7 防水（F=734）", "電源類型（F=1,203）"),
    ("核心消費族群", "S1：禮品購買族（62.4%）", "S1：精準修剪需求族（91.6%）"),
    ("Hero SKU", "B0FL267TCG（鬍鬚修剪器）", "B0GBWZBMS5（鼻耳修剪器）"),
    ("最適定價", "USD 79（Conjoint 最優）", "JPY 3,980–4,980"),
    ("行銷核心訊息", "Gift He'll Actually Use", "プロが認める精密さ"),
    ("主要通路", "Amazon US + Father's Day 檔期", "Amazon JP + 楽天市場マラソン"),
]

for row_vals in diff_rows:
    row = diff_table.add_row().cells
    for i, val in enumerate(row_vals):
        row[i].text = val
        p = row[i].paragraphs[0]
        p.runs[0].font.size = Pt(10)
        if i == 1:
            set_cell_bg(row[i], "EEF2FF")
        elif i == 2:
            set_cell_bg(row[i], "FFF0F3")

doc.add_paragraph()

add_heading(doc, "3.2 消費者願付溢價（WTP）", level=2)

add_body(doc,
    "Conjoint 分析量化了消費者願意為各屬性額外支付的金額，提供定價與規格決策的客觀依據："
)

wtp_table = doc.add_table(rows=1, cols=4)
wtp_table.style = "Table Grid"
wtp_hdr = wtp_table.rows[0].cells
for i, txt in enumerate(["屬性", "美國 WTP（USD）", "日本 WTP（JPY）", "決策建議"]):
    wtp_hdr[i].text = txt
    p = wtp_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(wtp_hdr[i], "C9A36F")

wtp_rows = [
    ("機身尺寸（緊湊型）", "$89.81", "—", "US 主打小巧攜帶便利"),
    ("附件件數（7件以上）", "$79.86", "¥6,800", "兩市場均以 7-in-1 為主規格"),
    ("長度調整段數", "$78.18", "¥7,961（最高）", "JP 最高 WTP，必須強調精度"),
    ("IPX7 防水", "溢價顯著", "¥4,200", "兩市場新品最低規格門檻"),
    ("USB-C 充電", "$65+", "¥3,500", "US 第二強驅動，JP 亦顯著"),
    ("調整精度（0.5mm）", "—", "¥6,773", "JP 差異化核心規格"),
]

for row_vals in wtp_rows:
    row = wtp_table.add_row().cells
    for i, val in enumerate(row_vals):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "3.3 市佔率現況與機會", level=2)

add_body(doc,
    "MNL Share-of-Preference 模型模擬顯示，URBANER 在美國市場現有定位下，理論市佔率約為 5.21%，"
    "競品合計佔 94.67%。雖然市場仍由大品牌主導，但在禮品場景與 7-in-1 套組規格上，"
    "URBANER 具備明確的利基切入點。"
)

add_body(doc, "成長機會：", bold=True)
add_bullet(doc, "美國：集中資源在 Father's Day 與 Christmas 禮品季，以 B0FL267TCG 為主攻 SKU，PPC 預算佔比 70%")
add_bullet(doc, "日本：把握楽天お買い物マラソン，以 B0GBWZBMS5 + 楽天直営店優勢，強攻鼻耳修剪器類別")
add_bullet(doc, "跨市場：以 IPX7 為新品最低門檻，7-in-1 套組為主流規格，建立品牌技術一致性")

add_heading(doc, "3.4 產品規格優化建議", level=2)

add_body(doc, "基於 Conjoint 結果，建議 2026 年後新 SKU 開發遵循以下規格：", bold=True)

spec_table = doc.add_table(rows=1, cols=3)
spec_table.style = "Table Grid"
spec_hdr = spec_table.rows[0].cells
for i, txt in enumerate(["規格項目", "必要條件", "理由"]):
    spec_hdr[i].text = txt
    p = spec_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(spec_hdr[i], "1A2744")

spec_rows = [
    ("防水規格", "IPX7 以上（兩市場必要）", "兩市場 ANOVA 皆顯著，消費者視為基本條件"),
    ("充電方式", "USB-C（US 必要，JP 建議）", "US Conjoint 第二重要屬性"),
    ("套組件數", "≥ 7 件（兩市場）", "兩市場 Top 屬性，7 件視覺錨定超值感"),
    ("長度調整精度", "0.5mm（JP 必要）", "JP WTP 最高（¥7,961），核心差異化"),
    ("電源類型", "可充電式（JP S1 偏好）", "JP 91.6% 消費者屬 S1 群（充電偏好）"),
    ("定價帶", "US $79 / JP ¥3,980–4,980", "Conjoint 最優定價，兼顧 WTP 與競品壓力"),
]

for row_vals in spec_rows:
    row = spec_table.add_row().cells
    for i, val in enumerate(row_vals):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_page_break()

# ─── CHAPTER 4：行銷劇本 ─────────────────────────────────────────────────────

add_heading(doc, "四、11 條行銷劇本總覽", level=1)

add_body(doc,
    "基於上述分析，本研究產出 11 條具體行銷劇本（Playbook），並為每條劇本配置行銷文案、"
    "KPI 目標與行銷框架說明，可直接作為廣告素材或商品文案使用。"
)

playbook_data = [
    ("美國市場（US）", [
        ("US-01", "Father's Day Gift Campaign",
         "禮品觸發季節性購買，以社會認同強化選品信心。聚焦「The Gift He'll Actually Use」核心訊息，"
         "搭配 Amazon Sponsored Products + Brands 廣告集中投放 5–6 月父親節檔期。"),
        ("US-02", "USB-C + IPX7 技術訴求",
         "以產品定位理論建立技術壁壘，USB-C 快充 + IPX7 防水為標配。"
         "訴求「IPX7 isn't a feature anymore; it's the baseline」，建立標準制定者形象。"),
        ("US-03", "Christmas Bundle Campaign",
         "聖誕節禮品季再攻，以 7-in-1 套組的價值錨定創造超值感知，"
         "強調「Everything in one premium box」一站式解決送禮難題。"),
        ("US-04", "品牌傳承信任建立",
         "以 50 年日本刀片工藝為情感錨點，Maslow 自尊層次訴求，"
         "建立「50-year Japanese blade craft」品牌故事，差異化於中國大陸競品。"),
    ]),
    ("日本市場（JP）", [
        ("JP-01", "精密調整技術主打",
         "0.5mm 精密調整為 JP Hero 屬性（WTP 最高），以「プロが認める精密さ」訴求專業美容師認同，"
         "並強調楽天直営店正規品身份。"),
        ("JP-02", "楽天マラソン × 套組攻勢",
         "楽天お買い物マラソン期間集中投放，以積分倍數訴求搭配 7 件套組超值包裝，"
         "把握日本消費者積分活動敏感性。"),
        ("JP-03", "電源類型市場細分",
         "針對 S1 群體（充電式偏好，佔 91.6%）強化 USB-C 可充電訴求，"
         "以乾電池用戶（S2，7.8%）為次要市場，分群投放不同廣告素材。"),
        ("JP-04", "父の日プレミアムギフト",
         "移植 US Father's Day 模型至 JP 市場，結合日本送禮文化禮盒包裝，"
         "以「大切な方への最高の贈り物」強化禮品場景情感連結。"),
    ]),
    ("跨市場共用（COMMON）", [
        ("COMMON-01", "IPX7 新品最低門檻",
         "兩市場新 SKU 均以 IPX7 為必要規格，統一品牌技術基準。"
         "文案核心：「防水 isn't a luxury. It's the new standard.」"),
        ("COMMON-02", "套組 ≥7 件為主流規格",
         "兩市場 top 屬性均包含附件件數，7-in-1 為視覺錨定點。"
         "訴求：「7 Functions. One System. Everything Included.」"),
        ("COMMON-03", "Hero SKU 差異化集中投放",
         "US Hero：B0FL267TCG（鬍鬚修剪器，Conjoint 最優）；"
         "JP Hero：B0GBWZBMS5（鼻耳修剪器，楽天直営転換率高）。"
         "各市場 PPC 預算 70% 集中 Hero SKU，維持 ACOS ≤18%（US）/ ≤20%（JP）。"),
    ]),
]

for market_name, items in playbook_data:
    color_hex = "2E5BFF" if "US" in market_name else ("D32F4D" if "JP" in market_name else "1A2744")
    add_heading(doc, market_name, level=2)

    for code, title, desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        code_run = p.add_run(f"[{code}] ")
        code_run.font.bold = True
        code_run.font.size = Pt(11)
        code_run.font.color.rgb = RGBColor.from_string(color_hex)
        title_run = p.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent = Cm(1.0)
        desc_run = desc_p.add_run(desc)
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)

doc.add_paragraph()
doc.add_page_break()

# ─── CHAPTER 5：結語 ─────────────────────────────────────────────────────────

add_heading(doc, "五、結語與後續行動", level=1)

add_body(doc,
    "本研究透過 MMR 四階段方法論，系統性地將 Amazon 平台數據轉化為可執行的商業決策。"
    "核心貢獻在於：不僅告訴 URBANER「市場偏好什麼」，更量化了「消費者願意多付多少錢」，"
    "並據此產出具體的產品規格建議、定價策略與行銷文案。"
)

add_body(doc, "建議後續優先執行事項：", bold=True)

action_items = [
    ("立即（0–1個月）",
     "B0FL267TCG 美國 Father's Day PPC 預算提升至 US 總預算 70%；B0GBWZBMS5 楽天マラソン備庫。"),
    ("短期（1–3個月）",
     "以 Conjoint 最優配置（7件套組 + USB-C + IPX7）開發新 SKU；"
     "更新兩市場 A+ Content 加入 IPX7 防水模組與件數視覺化模組。"),
    ("中期（3–6個月）",
     "建立 Customer Review 自動監控系統，追蹤 Hero SKU 評分動態；"
     "在楽天市場擴大直営店優勢，衝刺 JP Hero SKU 銷售比率至 ≥40%。"),
    ("長期（6個月以上）",
     "以本研究方法論為基礎，每季執行一次市場屬性 ANOVA 更新，持續追蹤市場需求演變。"),
]

for period, action in action_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    period_run = p.add_run(f"▶ {period}：")
    period_run.font.bold = True
    period_run.font.size = Pt(11)
    period_run.font.color.rgb = RGBColor(0xC9, 0xA3, 0x6F)
    action_run = p.add_run(action)
    action_run.font.size = Pt(11)

doc.add_paragraph()

add_body(doc,
    "本報告所有結論均來自統計分析，建議在執行時配合實際市場反應持續校正。"
    "URBANER 從 OEM 代工到 OBM 品牌的轉型旅程，需要的正是這種「數據驅動、市場導向」的決策文化。",
    italic=True, color="5A5A5A"
)

# ─── Save ─────────────────────────────────────────────────────────────────────

os.makedirs("output", exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved → {OUTPUT_PATH}")
