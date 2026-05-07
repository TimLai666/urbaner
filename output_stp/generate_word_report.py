"""
Generate a complete Word report combining everything in output_stp/:
- dual_market_stp_report.md (雙市場 STP 完整報告)
- market_stp_us/market_stp_report.md (US 詳細報告 + PCA charts)
- market_stp_jp/market_stp_report.md (JP 詳細報告 + PCA charts)

All content is included verbatim — no truncation.
PNG charts (perceptual_map, segmentation_map, quality_heatmap) embedded.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Microsoft YaHei"


# ==================== style helpers ====================

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_run(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11, 5: 10}
    colors = {
        1: RGBColor(0x1F, 0x3A, 0x5F),
        2: RGBColor(0x2E, 0x5C, 0x8A),
        3: RGBColor(0x4A, 0x7A, 0xB0),
        4: RGBColor(0x5A, 0x5A, 0x5A),
        5: RGBColor(0x6E, 0x6E, 0x6E),
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    return p


def add_para(doc, text, size=10, bold=False, italic=False, color=None, align=None, indent_cm=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_blockquote(doc, lines, italic=True):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        style_run(run, size=10, italic=italic, color=RGBColor(0x55, 0x55, 0x55))


def add_bullet(doc, text, size=10, bold=False, indent_cm=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run("• " + text)
    style_run(run, size=size, bold=bold)
    return p


def add_evidence(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run("⤷ " + text)
    style_run(run, size=size, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    return p


def add_table(doc, header, rows, header_fill="1F3A5F", first_col_bold=False, col_widths=None,
              cell_size=9, hdr_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        style_run(run, size=hdr_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        zebra = "F0F4F8" if r % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = table.rows[r + 1].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val is not None else "")
            bold = first_col_bold and i == 0
            style_run(run, size=cell_size, bold=bold)
            set_cell_shading(cell, zebra)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_image(doc, path, caption=None, width_inches=6.5):
    if not os.path.exists(path):
        add_para(doc, f"[圖片缺失：{path}]", italic=True, color=RGBColor(0xCC, 0x44, 0x44))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(caption)
        style_run(crun, size=9, italic=True, color=RGBColor(0x77, 0x77, 0x77))


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    style_run(run, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


# ==================== content blocks ====================

def add_finding(doc, finding_id, fields, evidence_quotes=None):
    """fields = list of (label, value) tuples; evidence_quotes optional."""
    add_heading(doc, f"Finding {finding_id}", level=4)
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        r1 = p.add_run(f"{label}：")
        style_run(r1, size=10, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
        r2 = p.add_run(value)
        style_run(r2, size=10)
    if evidence_quotes:
        for q in evidence_quotes:
            add_evidence(doc, q)


# ==================== main ====================

def main():
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("URBANER 雙市場 STP 完整分析報告")
    style_run(run, size=22, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Review-Mining-STP × Segmentation × Targeting × Positioning")
    style_run(run, size=12, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("FourSight Lab × URBANER 產學合作專案　|　資料截止 2026-05-05")
    style_run(run, size=10, color=RGBColor(0x77, 0x77, 0x77))

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run("資料來源：URBANER 自家 Amazon 評論 11,523 則 + 競品評論 + 社群媒體洞察")
    style_run(run, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph()

    # =========================================================
    # PART A: 雙市場 STP 完整報告
    # =========================================================
    add_heading(doc, "Part A. URBANER 雙市場 STP 分析報告", level=1)
    add_blockquote(doc, [
        "市場：美國（US） vs 日本（JP）",
        "資料來源：URBANER 自家 Amazon 評論（11,523 則）+ 競品評論（research/amazon_us, research/amazon_jp）+ 社群媒體洞察（research/social_us, research/social_jp）",
        "方法：Review-Mining-STP 工作流（Axis A 顯著度 0–7、Axis B 品質 0–10）+ Segmentation/Targeting/Positioning 統計腳本",
    ])
    add_separator(doc)

    # --- 0. Executive Summary ---
    add_heading(doc, "0. 報告摘要 Executive Summary", level=2)
    add_table(
        doc,
        header=["指標", "US", "JP"],
        rows=[
            ["評論數", "4,360", "7,163"],
            ["自家商品數（n_reviews ≥ 3）", "52", "36"],
            ["主要類別", "Beard/Mustache + Nose/Ear + Body Groomers", "Nose/Ear + Beard/Mustache"],
            ["評論語言", "en 100%", "ja 97% / en 3%"],
            ["顯著區隔屬性數（ANOVA p<0.05）", "99 / 114", "58 / 114"],
            ["最具區隔力屬性", "gift_suitability_men (F=1299.469)", "total_attachments_count (F=2630.384)"],
            ["區隔數 (k)", "3", "3"],
        ],
        col_widths=[5.0, 5.5, 5.5],
    )

    add_heading(doc, "核心結論", level=3)
    add_para(doc,
             "1. US 顧客重視「送禮場景 × 多功能套組」：最具區隔力屬性是 gift_suitability_men（F=1299）與 num_grooming_functions（F=1193），反映美國電剪市場大量出現在父親節 / 聖誕節送禮情境，且消費者期待「一機多用」的套組形式。")
    add_para(doc,
             "2. JP 顧客則由「附件數 / 調節精度 / 電源類型」三軸主導：total_attachments_count (F=2630)、guide_comb_variety (F=2412)、adjustable_comb_settings (F=1827) 是區隔最強的屬性。日本市場決策深度高、講究刻度（0.5mm 單位）與長期使用的耗材完整度。")
    add_para(doc,
             "3. Waterproof（IPX8）在兩市場都是核心差異化軸：US ANOVA 排名 7（F=903），JP 排名 9（F=461），且社群洞察一致顯示「防水可全水洗」是高端標配。")
    add_para(doc,
             "4. 位置：JP 表現最佳的 SKU 與 US 完全不同：US hero SKU = B0FL267TCG（距理想 1.601，Beard / Mustache Trimmers），JP hero SKU = B0GBWZBMS5（距理想 1.97，Nose / Ear Trimmers）。兩市場應採差異化主打商品策略。")

    # --- 1. Attribute Extraction ---
    add_heading(doc, "1. 屬性抽取摘要 Attribute Extraction Summary", level=2)
    add_table(
        doc,
        header=["欄位", "數值"],
        rows=[
            ["目標最低屬性數", "30"],
            ["實際屬性數", "114"],
            ["屬性短缺原因", "none — pre-frozen catalog of 114 attributes reused across markets"],
            ["理論族缺口", "無 — 4 族皆覆蓋"],
        ],
        col_widths=[4.5, 11.0],
    )
    add_para(doc, "主題（Themes，20 個）：", bold=True)
    add_para(doc,
             "Power Source, Motor & Blade, Waterproof, Functions, Adjustability, Design, "
             "Portability, Accessories, Safety, Build Quality, Ease of Use, Gifting, "
             "Price & Value, Target User, Blade Specs, Brand & Trust, Sensory Experience, "
             "Hair Capability, Social Context, Eco & Sustainability",
             indent_cm=0.5)

    add_para(doc, "理論族與子主題覆蓋（4 族全覆蓋）：", bold=True)
    add_bullet(doc, "product_positioning：attributes, functions, benefits, usage_context_service_experience")
    add_bullet(doc, "maslow：physiological, safety, belongingness, esteem, self_actualization")
    add_bullet(doc, "purchase_motivation：functional, security, relational")
    add_bullet(doc, "wom_motivation：altruistic, social_identity, self_enhancement, emotional_expression")

    add_para(doc, "雙軸計分：", bold=True)
    add_bullet(doc, "Axis A（顯著度 0–7）：每則評論 × 屬性的關鍵字密度與內容長度加權。0 = 該屬性不在此評論中；7 = 該屬性是評論的中心議題。")
    add_bullet(doc, "Axis B（品質 0–10）：每商品 × 屬性的星級基線 + 屬性負向關鍵字扣分 + 整體情緒微調。從整體商品評論彙總而得，反映此商品在該屬性上的市場觀感。")

    # --- 2. Segmentation ---
    add_heading(doc, "2. 市場區隔 Segmentation", level=2)

    add_heading(doc, "2.1 美國市場（US）— k=3，silhouette=0.358", level=3)
    add_para(doc, "這部分在做什麼：以 4,360 則 URBANER 美國評論的 114-屬性 Axis A 顯著度向量，做 PCA + K-Means 顧客分群。")
    add_para(doc, "Axis 模式：Axis A（顯著度）。")
    add_para(doc, "統計方法：StandardScaler → PCA（保留 85% 變異）→ K-Means（k=3-6，搜尋最佳 silhouette + ≥5% 最小群體護欄）。")
    add_para(doc, "理論使用：product_positioning（functions, benefits, attributes）；maslow（physiological, safety, esteem）；purchase_motivation（functional, security）。")

    add_heading(doc, "Segment 1（n=850，19.5%，avg★=3.68）", level=4)
    add_bullet(doc, "核心屬性 Top 5：gift_suitability_men, primary_user_gender, ear_hair_trimming_function, attachment_fitment, total_attachments_count")
    add_bullet(doc, "類別分布：Beard / Mustache Trimmers 87.5%, Nose / Ear Trimmers 6.8%, Body Groomers 5.6%")
    add_bullet(doc, "代表商品：B0CGHWFF71 (11.8%), B008KEJ1LM (10.0%), B0BQ1GZY7H (9.1%)")

    add_heading(doc, "Segment 2（n=3,337，76.5%，avg★=3.18）", level=4)
    add_bullet(doc, "核心屬性 Top 5：ear_hair_trimming_function, product_longevity, beard_trimming_function, gift_suitability_men, rechargeable_design")
    add_bullet(doc, "類別分布：Beard / Mustache Trimmers 94.8%, Nose / Ear Trimmers 4.2%, Body Groomers 1.0%")
    add_bullet(doc, "代表商品：B08P4HHSZT (13.1%), B0CQKZ3N5G (12.9%), B0BQ1GZY7H (12.6%)")

    add_heading(doc, "Segment 3（n=173，4.0%，avg★=4.48）", level=4)
    add_bullet(doc, "核心屬性 Top 5：power_source_type, gift_suitability_men, rechargeable_design, beard_trimming_function, ear_hair_trimming_function")
    add_bullet(doc, "類別分布：Beard / Mustache Trimmers 75.7%, Body Groomers 17.9%, Nose / Ear Trimmers 6.4%")
    add_bullet(doc, "代表商品：B0FS5TCR3B (20.8%), B0FL267TCG (7.5%), B0F2RSTDSJ (5.8%)")

    add_finding(
        doc, "US-SEG-01",
        fields=[
            ("finding_statement", "US 最大區隔（Segment 2，76.5%，n=3,337）以 ear_hair_trimming_function 為主軸。"),
            ("business_implication", "美國顧客在意「鼻耳毛功能」「壽命」「鬍鬚修整」「送禮適用性」「充電式設計」等實用面，Listing 主圖與 A+ 內容應以這些訴求作為視覺主訴。"),
            ("axes_used", "salience"),
            ("methods_used", "StandardScaler → PCA(0.85) → KMeans(k=3-6, silhouette+share guardrail)"),
            ("theories_used", "product_positioning, maslow, purchase_motivation"),
            ("subtheories_used", "functions, benefits, physiological, functional, esteem"),
            ("statistical_results", "silhouette=0.358, k=3, n=4,360"),
            ("reproducibility", "review_scoring_table.csv → SAL_COLS → StandardScaler → PCA(0.85) → KMeans"),
        ],
        evidence_quotes=[
            'evidence_quote [US_R000001, 1★, B001K85BN2]: "Good awful. The only positive is that this trimmer can be used in the shower. I don\'t have a bushy beard and even so this trimmer resulted in a very uneven cut and I was constantly having to rinse the gears hairs out to make sure it kept working. Spend more and get a better trimm"',
            'evidence_quote [US_R000011, 5★, B008KEJ1LM]: "other shavers i bought had batteries didn\'t last so i figured buying a panasonic product was a safe choice as they are a major global battery manufacturer. Still going strong after several years. whenever i go on vacation i can leave the charger home not worrying the shaver will "',
        ],
    )

    add_heading(doc, "2.2 日本市場（JP）— k=3，silhouette=0.570", level=3)
    add_para(doc, "這部分在做什麼：以 7,163 則 URBANER 日本評論的 114-屬性 Axis A 顯著度向量做相同流程。")
    add_para(doc, "Axis 模式：Axis A（顯著度）。")
    add_para(doc, "統計方法：同 US 流程；JP silhouette 較高（0.57 vs 0.36）反映日本顧客分群更分明。")

    add_heading(doc, "Segment 1（n=6,559，91.6%，avg★=3.46）", level=4)
    add_bullet(doc, "核心屬性 Top 5：power_source_type, nose_hair_trimming_function, price_value_ratio, ease_of_use_overall, product_longevity")
    add_bullet(doc, "類別分布：Nose / Ear Trimmers 59.4%, Beard / Mustache Trimmers 40.6%")
    add_bullet(doc, "代表商品：B07XTLC91J (7.6%), B0C2Y71YXV (7.5%), B0CY3PWWTQ (6.8%)")

    add_heading(doc, "Segment 2（n=561，7.8%，avg★=4.01）", level=4)
    add_bullet(doc, "核心屬性 Top 5：adjustable_comb_settings, power_source_type, product_dimensions_size, total_attachments_count, attachment_fitment")
    add_bullet(doc, "類別分布：Beard / Mustache Trimmers 75.2%, Nose / Ear Trimmers 24.8%")
    add_bullet(doc, "代表商品：B07CYZH2XC (8.0%), B0DSB44YJZ (8.0%), B0G492LXS9 (8.0%)")

    add_heading(doc, "Segment 3（n=43,0.6%，avg★=3.05）", level=4)
    add_bullet(doc, "核心屬性 Top 5：operation_vibration_feel, low_vibration_design, motor_power_rpm, power_source_type, low_noise_operation")
    add_bullet(doc, "類別分布：Nose / Ear Trimmers 62.8%, Beard / Mustache Trimmers 37.2%")
    add_bullet(doc, "代表商品：B07FFGYWJ6 (16.3%), B016JNKVHI (7.0%), B07XLLBB5J (7.0%)")

    add_finding(
        doc, "JP-SEG-01",
        fields=[
            ("finding_statement", "JP 最大區隔（Segment 1，91.6%，n=6,559）以 power_source_type 為主軸；附帶議題包括 nose_hair、CP值、易用性、長壽性。"),
            ("business_implication", "日本顧客在 Nose/Ear 與 Beard 兩品類混合議題；電源型態（乾電池 vs 充電）是關鍵決策點。Amazon JP Listing 應在 bullet-point 1 即清楚標示電源類型與耗電/續航。"),
            ("axes_used", "salience"),
            ("methods_used", "StandardScaler → PCA(0.85) → KMeans(k=3-6)"),
            ("theories_used", "product_positioning, maslow, purchase_motivation"),
            ("subtheories_used", "attributes, functions, functional, security, physiological"),
            ("statistical_results", "silhouette=0.570, k=3, n=7,163"),
        ],
        evidence_quotes=[
            'evidence_quote [JP_R000010, 5★, B001K85BN2]: "よく剃れるし問題ありません。以前は充電式を使っていましたが、電池が劣化すると捨てるしかなくなるのでこういった製品は乾電池式一択です。"',
            'evidence_quote [JP_R000048, 4★, B001K85BN2]: "私はもみ上げが伸びてくると膨らんで嫌なんですが、これで簡単に整えれました！！そんなに使用頻度は無いので電池式で十分！！説明書に充電池は使えませんとありました。よって星マイナス1"',
        ],
    )

    add_heading(doc, "2.3 跨市場區隔比較", level=3)
    add_table(
        doc,
        header=["維度", "US 大眾區隔", "JP 大眾區隔"],
        rows=[
            ["主軸屬性", "ear_hair_trimming_function", "power_source_type"],
            ["第二軸屬性", "product_longevity", "nose_hair_trimming_function"],
            ["第三軸屬性", "beard_trimming_function", "price_value_ratio"],
            ["Avg★", "3.18", "3.46"],
            ["群體佔比", "76.5%", "91.6%"],
        ],
        col_widths=[4.0, 5.5, 5.5],
    )
    add_para(doc, "關鍵差異：", bold=True)
    add_para(doc, "US 的大眾區隔仍偏「功能展示 + 送禮場景」；JP 大眾區隔則更聚焦於「電源類型 + 鼻毛/鬍鬚實用」「CP 值」「易用性」「壽命」。前者像是「買來送的商品」，後者像是「買來自用、講究 CP 的耐久工具」。")

    # --- 3. Targeting ---
    add_heading(doc, "3. 目標市場 Targeting", level=2)
    add_para(doc, "這部分在做什麼：以 ANOVA 找出區隔間最有差異的屬性，並用「平均星等 × 群體佔比」排序優先區隔。")
    add_para(doc, "Axis 模式：Axis A（顯著度）為依變數；segment 為自變數。")
    add_para(doc, "統計方法：one-way ANOVA per attribute；priority/secondary/deprioritized 排名。")
    add_para(doc, "理論使用：purchase_motivation（functional, security, relational）；maslow（safety, esteem, self_actualization）；product_positioning（functions, benefits）。")

    add_heading(doc, "3.1 US — Top 10 最具區隔力屬性", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "F", "p"],
        rows=[
            ["1", "gift_suitability_men", "1299.469", "0"],
            ["2", "primary_user_gender", "1196.508", "0"],
            ["3", "num_grooming_functions", "1193.262", "0"],
            ["4", "total_attachments_count", "1014.905", "0"],
            ["5", "attachment_fitment", "975.252", "0"],
            ["6", "power_source_type", "941.026", "0"],
            ["7", "waterproof_rating_ipx8", "902.847", "0"],
            ["8", "guide_comb_variety", "902.265", "0"],
            ["9", "multi_use_versatility_score", "778.884", "0"],
            ["10", "quick_touch_up_design", "708.689", "0"],
        ],
        col_widths=[1.5, 7.0, 3.5, 2.5],
    )

    add_heading(doc, "3.2 JP — Top 10 最具區隔力屬性", level=3)
    add_table(
        doc,
        header=["排名", "屬性", "F", "p"],
        rows=[
            ["1", "total_attachments_count", "2630.384", "0"],
            ["2", "guide_comb_variety", "2412.195", "0"],
            ["3", "adjustable_comb_settings", "1826.965", "0"],
            ["4", "attachment_fitment", "1723.208", "0"],
            ["5", "product_dimensions_size", "1004.693", "0"],
            ["6", "power_source_type", "543.117", "0"],
            ["7", "rechargeable_design", "497.535", "0"],
            ["8", "beard_trimming_function", "494.451", "0"],
            ["9", "waterproof_rating_ipx8", "460.933", "0"],
            ["10", "battery_independence", "346.662", "0"],
        ],
        col_widths=[1.5, 7.0, 3.5, 2.5],
    )

    add_heading(doc, "3.3 跨市場差異 — 最具區隔力屬性對照", level=3)
    add_table(
        doc,
        header=["類別", "屬性"],
        rows=[
            ["共同（兩市場 Top 15 都有）",
             "attachment_fitment, ease_of_use_overall, guide_comb_variety, power_source_type, total_attachments_count, waterproof_rating_ipx8"],
            ["僅 US Top 15",
             "ear_hair_trimming_function, ease_of_cleaning, gift_suitability_men, grooming_kit_completeness, multi_use_versatility_score, num_grooming_functions, plastic_material_quality, primary_user_gender, quick_touch_up_design"],
            ["僅 JP Top 15",
             "adjustable_comb_settings, battery_independence, beard_trimming_function, compact_size, cordless_during_use, instruction_manual_quality, product_dimensions_size, rechargeable_design, usb_c_charging"],
        ],
        col_widths=[4.5, 11.5],
    )

    add_para(doc, "洞察：", bold=True)
    add_bullet(doc, "兩市場共同看重 total_attachments_count、guide_comb_variety、attachment_fitment、power_source_type、waterproof_rating_ipx8 等基礎工具屬性。")
    add_bullet(doc, "US 獨有的高區隔屬性偏向「ear_hair_trimming_function, ease_of_cleaning, gift_suitability_men」— 多功能/性別/送禮維度。")
    add_bullet(doc, "JP 獨有的高區隔屬性偏向「adjustable_comb_settings, battery_independence, beard_trimming_function」— 調整精度/電池獨立/耐久維度。")

    add_finding(
        doc, "TGT-01（US）",
        fields=[
            ("finding_statement", "US Top 區隔屬性 gift_suitability_men（F=1299）反映父親節、生日、聖誕節等送禮情境是 US 男士電剪購買的核心觸發。"),
            ("business_implication", "URBANER US 應在 Q4 (10–12 月)、父親節（6 月）強化「Gift Ready」訴求 — 包含禮盒包裝、銷售套組、Father's Day 主題拍攝。"),
            ("axes_used", "salience"),
            ("methods_used", "ANOVA"),
            ("theories_used", "purchase_motivation, maslow, wom_motivation"),
            ("subtheories_used", "relational, belongingness, esteem, social_identity"),
            ("statistical_results", "F=1299.469, p<0.0001, df=2, n=4,360"),
            ("reproducibility", "segmentation_variables.csv → groupby(segment) → scipy.stats.f_oneway"),
        ],
        evidence_quotes=[
            'evidence_quote [US_R000003, 1★, B001K85BN2]: "This was apparantly an export model. The in structions were in Japanese. This purchase was intended to replace a Panasonic ER2430. The unit that was shipped ER2403P has a different configeration for the heighth adjustment. My old unit had settings of 3, 4,5,6,, and so on. This"',
            'evidence_quote [US_R000010, 5★, B008KEJ1LM]: "Excelente tengo años usandolo y hasta hora no me a fallado lo volveria a comprar"',
        ],
    )

    add_finding(
        doc, "TGT-02（JP）",
        fields=[
            ("finding_statement", "JP Top 區隔屬性 total_attachments_count（F=2630）+ guide_comb_variety（F=2412）顯示日本顧客極度在意附件齊全度與長度檔位數。"),
            ("business_implication", "URBANER JP Listing 應在 bullet-point 1-2 直接列出「アタッチメント X 個 / 長さ調整 38 段階」等具體數字；單獨販售耗材包是有效的後續銷售切入點。"),
            ("axes_used", "salience"),
            ("methods_used", "ANOVA"),
            ("theories_used", "product_positioning, purchase_motivation"),
            ("subtheories_used", "attributes, functions, functional, security"),
            ("statistical_results", "F=2630.384, p<0.0001, df=2, n=7,163"),
        ],
        evidence_quotes=[
            'evidence_quote [JP_R000145, 1★, B001K85BN2]: "前のモデルのアタッチメントが壊れた（本体はまだ使える）ので進化を確かめるべく購入。結果はかっこよくなっただけ。性能はむしろ落ちたかも。パワーが落ちたのは他の人も触れているが単3から単4に変化。トリム機構は変わって無いから必要十分でスペックダウンしたわけでは無さそう。新しい方の3㎜は本当に3㎜でそろう。前のは幾分短い刈高。なので長く感じる。また完了までの時間が掛かる様になった。デザインがかっこよくなり機能を落として本質が見えていない。おそらく開発側は変える必要を感じていなかったんじゃないか？　ただヒゲトリマーは怪しいメーカーか中華品がはびこっていて、3㎜以"',
            'evidence_quote [JP_R000181, 4★, B001K85BN2]: "良い商品だとは思いますがちょっと髭を短く残したいみたいな使い方をしたかった自分にとってはアタッチメントを1番短くセットしても髭に届かず意味のないものになってしまいました。代わりに眉毛を整えるのに使ってます。"',
        ],
    )

    add_heading(doc, "3.4 優先順序建議", level=3)
    add_table(
        doc,
        header=["市場", "Priority Segment", "Secondary", "Deprioritized"],
        rows=[
            ["US", "Segment 3（avg★=4.48, 4.0%）", "Segment 1（avg★=3.68, 19.5%）", "Segment 2（avg★=3.18, 76.5%）"],
            ["JP", "Segment 2（avg★=4.01, 7.8%）", "Segment 1（avg★=3.46, 91.6%）", "Segment 3（avg★=3.05, 0.6%）"],
        ],
        col_widths=[1.5, 4.5, 4.5, 4.5],
    )
    add_blockquote(doc, [
        "權衡說明：避免只看群體大小。Priority segment 的策略意義是「滿意度高、可規模化的群體」，"
        "而最大群體（如 US 76.5% 或 JP 91.6%）若 avg★ 偏低，則優先處理痛點（救火）而不是擴張。",
    ])

    # --- 4. Positioning ---
    add_heading(doc, "4. 定位 Positioning", level=2)
    add_para(doc, "這部分在做什麼：用 Axis B（品質 0–10）矩陣做 PCA 二維投影，並計算各 SKU 與「全屬性 = 10」理想點的 RMS 距離。")
    add_para(doc, "Axis 模式：Axis B（品質）。")
    add_para(doc, "統計方法：標準化 → PCA → 理想點距離。")
    add_para(doc, "理論使用：product_positioning（attributes, functions, benefits）；maslow（esteem, safety, self_actualization）；purchase_motivation（functional, security）。")

    add_heading(doc, "4.1 US — Top 10 最接近理想點的 URBANER SKU", level=3)
    add_table(
        doc,
        header=["排名", "ASIN", "距離", "n_reviews", "類別"],
        rows=[
            ["1", "B0FL267TCG", "1.601", "29", "Beard / Mustache Trimmers"],
            ["2", "B0FX4NX6HM", "1.761", "8", "Body Groomers"],
            ["3", "B0CDNQ62ML", "1.766", "8", "Body Groomers"],
            ["4", "B0GLJRLS5G", "1.814", "99", "Beard / Mustache Trimmers"],
            ["5", "B0FHWWPLCP", "1.867", "8", "Body Groomers"],
            ["6", "B0G1Z1QKJ8", "1.896", "50", "Beard / Mustache Trimmers"],
            ["7", "B0FS5TCR3B", "1.912", "158", "Beard / Mustache Trimmers"],
            ["8", "B092H8XHDG", "2.026", "8", "Nose / Ear Trimmers"],
            ["9", "B0F66FZHY3", "2.029", "8", "Nose / Ear Trimmers"],
            ["10", "B0F2RSTDSJ", "2.065", "73", "Beard / Mustache Trimmers"],
        ],
    )

    add_heading(doc, "4.2 JP — Top 10 最接近理想點的 URBANER SKU", level=3)
    add_table(
        doc,
        header=["排名", "ASIN", "距離", "n_reviews", "類別"],
        rows=[
            ["1", "B0GBWZBMS5", "1.97", "61", "Nose / Ear Trimmers"],
            ["2", "B0742G961R", "2.23", "40", "Beard / Mustache Trimmers"],
            ["3", "B0F99F11BW", "2.316", "142", "Nose / Ear Trimmers"],
            ["4", "B0F5VJD5DQ", "2.537", "150", "Nose / Ear Trimmers"],
            ["5", "B0BL2YWH3N", "2.657", "43", "Beard / Mustache Trimmers"],
            ["6", "B0FJS3QR1W", "2.817", "141", "Beard / Mustache Trimmers"],
            ["7", "B0FJS3TF3V", "2.817", "141", "Beard / Mustache Trimmers"],
            ["8", "B016JNKVHI", "2.918", "215", "Beard / Mustache Trimmers"],
            ["9", "B0BY1KRJZ1", "2.953", "115", "Beard / Mustache Trimmers"],
            ["10", "B0FHKR48NR", "3.113", "16", "Nose / Ear Trimmers"],
        ],
    )

    add_heading(doc, "4.3 跨市場 Hero SKU 對照", level=3)
    add_table(
        doc,
        header=["市場", "Hero SKU", "距理想點", "評論數", "類別"],
        rows=[
            ["US", "B0FL267TCG", "1.601", "29", "Beard / Mustache Trimmers"],
            ["JP", "B0GBWZBMS5", "1.97", "61", "Nose / Ear Trimmers"],
        ],
    )

    add_finding(
        doc, "POS-01",
        fields=[
            ("finding_statement", "US 與 JP 的 Hero SKU 不同（US B0FL267TCG、JP B0GBWZBMS5），且兩市場的高分 SKU 排名也不同；應採差異化主打商品策略。"),
            ("business_implication", "URBANER 不應將 US 暢銷品直接當作 JP 主推；JP 高分 SKU 多集中於 Nose/Ear（B0GBWZBMS5、B0F99F11BW、B0F5VJD5DQ），US 高分 SKU 多集中於 Beard/Mustache 高 SKU 編號。"),
            ("axes_used", "quality"),
            ("methods_used", "PCA + ideal-point RMS distance"),
            ("theories_used", "product_positioning, maslow, purchase_motivation"),
            ("subtheories_used", "attributes, functions, benefits, esteem, security, functional"),
        ],
    )

    # --- 5. Competitor benchmarking ---
    add_heading(doc, "5. 競品標竿 Competitor Benchmarking", level=2)
    add_blockquote(doc, [
        "資料來源：research/amazon_us/ 與 research/amazon_jp/ 內各 9 個類別的 Top 競品 + 5-6 則最新真實評論",
    ])

    add_heading(doc, "5.1 US 競品評論精華（每類 Top 2）", level=3)

    us_comp = [
        ("001 — Beard / Mustache Trimmers", [
            ("B07XJZ2HVC - Braun Beard Trimmer 6-Piece Set with Gillette ProGlide", [
                "5.0 | 2025-10-23 | These trim coarse beard hairs easily. Easy to adjust for length, trims very close to skin without irritation. Holds charge for quite awhile.",
                "5.0 | 2020-08-06 | Coming from Philips Norelco 5000. Tried Panasonic but build was cheap, adjustable a struggle. Braun: solid feel, multiple attachments, adjustable, even razor included.",
            ]),
            ("B0FHP7DKK1 - Valano Beard Trimmer All-in-One Self-Sharpening Cordless Kit", [
                "5.0 | 2026-04-14 | Easy shaving, no cutting yourself with blade. To get better shave keep it plugged in - doesn't keep charging long. Little and neat.",
                "5.0 | 2026-04-20 | Was already charged out of box. Shaved neck/face/surrounding trimmed beard easily.",
            ]),
        ]),
        ("002 — Nose / Ear Trimmers", [
            ("B0F66FZHY3 - Ear and Nose Hair Trimmer 2026 USB Rechargeable IPX7 (532)", [
                "5.0 | 2026-04-02 | Does what it says. Very pleased, easier than previous trimmer. Good quality, love rechargeable.",
                "5.0 | 2026-03-29 | Work great in nose and ears.",
            ]),
            ("B0GDD6Y7G7 - NUMIFUN Nose Hair Trimmer Women 2025 Pink Painless (466)", [
                "5.0 | 2026-03-23 | Best quality trimmer I've seen. Made with good material. Sturdy, nothing cheap looking.",
                "5.0 | 2026-04-13 | Works great. Heavier and more sturdy than cheaper models.",
            ]),
        ]),
        ("003 — Body Groomers", [
            ("B0BLZ5PJZJ - INVJOY Manscape Body & Pubic Hair Trimmer w/Light (862)", [
                "5.0 | 2025-02-12 | Best manscaper I've ever bought. No issues, works great.",
                "5.0 | 2026-04-11 | Super nice trimmer for cost. Cuts great, flawless charging.",
            ]),
            ("B0F2SVX5G9 - Electric Body Hair Trimmer Ball Trimmer Waterproof (631)", [
                "5.0 | 2026-04-12 | Arrived early in great condition.",
                "5.0 | 2024-09-23 | Amazing. From first use impressed by how easy and comfortable. No cuts or irritation - gave confidence.",
            ]),
        ]),
        ("004 — Eyebrow Trimmers", [
            ("B0FH6TDQ9L - Eyebrow Trimmer 2-in-1 Rechargeable Painless (889)", [
                "5.0 | 2025-02-28 | Excellent. Trims facial hair. Use for under brow arch, chin, upper lip.",
                "5.0 | 2025-02-06 | Quick, precise, painless. Game-changer. Compact, lightweight.",
            ]),
        ]),
        ("005 — Foil Shavers", [
            ("B0BP5DSSLC - Braun Series 9 Pro 9460cc Wet/Dry Foil Shaver (907)", [
                "5.0 | 2023-06-18 | Best shave in years - maybe ever! For 45 years used \"best\" razors and blades. Switched to electric.",
                "5.0 | 2025-11-27 | Was sceptical of expensive electric razor. Cost of disposable razors and environmental impact convinced me.",
            ]),
        ]),
        ("006 — Pet Electric Clippers", [
            ("B0BXNVL9WM - Dog Hair Clippers Kit LED Display Low Noise Heavy Duty (861)", [
                "5.0 | 2025-10-27 | Have 2 because Chow Chows. Battery fantastic but 2 dogs hard to finish on one charge.",
                "5.0 | 2025-10-20 | Love cordless. Battery lasts long. Good price. Doesn't upset animals. Gets through matted hair.",
            ]),
        ]),
        ("007 — Pet Nail Clipper", [
            ("B07DGMGSYQ - Millers Forge Stainless Steel Plier Style (810)", [
                "5.0 | 2025-01-31 | Buy a set every 1-2 years. Stay sharp longer. Cuts blue tick/walker hound nails cleanly without crushing.",
                "5.0 | 2026-03-23 | German shepherd + pittie thick wide nails - finally a clipper that works.",
            ]),
        ]),
        ("008 — Dog Nail Grinder", [
            ("B0FCG1FFVL - Casfuy 6-Speed Dog Nail Grinder w/ LED Lights Quiet (750)", [
                "5.0 | 2026-04-13 | Dog finally lets me do nails! Need treat per nail, but lets me grind nails.",
                "5.0 | 2025-02-21 | Two years ago hurt dog with clippers. After that wouldn't let me near feet. This grinder works.",
            ]),
        ]),
        ("009 — Baby Hair Clippers", [
            ("B09F98J1GS - Bimirth Baby Hair Clipper with Vacuum (509)", [
                "5.0 | 2026-03-11 | 3yo struggles at barber shop, didn't cry. Hair picked up by machine, no mess.",
                "5.0 | 2025-05-24 | Silent baby clippers lifesaver. Haircuts used to be battle with noise/vibration. So quiet.",
            ]),
        ]),
    ]

    for cat, products in us_comp:
        add_heading(doc, cat, level=4)
        for prod_name, reviews in products:
            add_para(doc, prod_name, bold=True, size=10, indent_cm=0.3, italic=True,
                     color=RGBColor(0x2E, 0x5C, 0x8A))
            for r in reviews:
                add_evidence(doc, r)

    add_heading(doc, "5.2 JP 競品評論精華（每類 Top 2）", level=3)

    jp_comp = [
        ("001 — Beard / Mustache Trimmers", [
            ("B07663D5M1 - Philips MG3720/15 多功能美容套装 (1743)", [
                "5.0 | 24-09-13 | アタッチメントも複数あり使い勝手がよいです",
                "4.0 | 22-01-29 | 悪い点: 電源のスイッチが超硬い (-1点) 良い点: 良く剃れる、扱いやすい、バッテリー持ちも十分 (+5点)",
            ]),
        ]),
        ("002 — Nose / Ear Trimmers", [
            ("B08CD75GNY - 皇家(Royal) シルバーブレット 手動式 鼻毛カッター ET-3 (4216)", [
                "5.0 | 26-02-16 | 切れ味よく、電動よりコンパクトでメンテナンス容易。質感良く使いやすい。バッテリー交換や充電の煩わしさなし",
                "5.0 | 25-11-05 | 電動5種類試したが毛抜けや電池問題あった。これに辿り着いた",
            ]),
        ]),
        ("003 — Body Groomers", [
            ("B06XSD3N13 - Panasonic ER-GK60-W 美体修剪器 VIO適用 浴室OK (3613)", [
                "5.0 | 26-01-28 | ガッツリ剃れる！でも一人では膝裏は剃り残り。安定のパナソニック",
                "4.0 | 25-07-02 | 2021購入4年使用、2-3ヶ月に1回。すね・腕・VIO全身対応で快適。剃り心地刺激少なく、ヒリヒリしない",
            ]),
        ]),
        ("004 — Eyebrow Trimmers", [
            ("B079M8D8Y9 - Panasonic ES-WF41-P 脸部シェーバー 光美麗 ピンク (6959)", [
                "5.0 | 26-02-21 | 使いやすい。確実にシェービング、痛くない。音が心地良い。医療クリニック推奨",
                "5.0 | 26-01-08 | 使い心地とても良い",
            ]),
        ]),
        ("005 — Foil Shavers", [
            ("B0B2NHRTFN - Panasonic Lamdash ES-NLV68-K 5層刀片 (838)", [
                "5.0 | 26-04-03 | ブラウンからの買い替え。剃り味こっち、肌アタック度合いも、剃り残しもこっち良い。リニア剃り心地気持ちいい",
                "4.0 | 26-04-19 | キレイに剃れるが夕方ぐらいに少し生えてくる",
            ]),
        ]),
        ("006 — Pet Electric Clippers", [
            ("B01K73KIMO - Pateker ペット用バリカン 静音軽量 トリミング (7964)", [
                "5.0 | 26-04-09 | わんちゃんのトリミング先の方足のボソボソ何とかしたくて購入。切れ味バツグン使いやすい",
                "4.0 | 26-01-18 | 犬全身トリミング初心者でも使いやすい。動作音静かで音敏感な子も落ち着く。刃の切れ味良く毛詰まらずスムーズ。軽量手にフィット長時間OK",
            ]),
        ]),
        ("007 — Pet Nail Clipper", [
            ("B07RSM1NG6 - HAWATOUR 犬用爪切り Pro w/ガード (9964)", [
                "5.0 | 26-04-03 | Decent clippers",
                "5.0 | 22-02-25 | Cuts pet's nails from both sides - easier on dog's nails (CA)",
            ]),
        ]),
        ("008 — Dog Nail Grinder", [
            ("B07ZKZB1J9 - Boshel ペットネイルグラインダー 静音 (1249)", [
                "5.0 | 24-03-02 | Great buy. Showed to mini dachshund, started back feet 1 nail/day (CA)",
                "5.0 | 22-04-07 | Simple, good power (2 speeds), ergonomic, cable charge (MX)",
            ]),
        ]),
        ("009 — Baby Hair Clippers", [
            ("B0BRNJ14YJ - Rozally 理容師監修 静音水洗 子供/大人兼用 (1119)", [
                "5.0 | 26-02-12 | 充電後長時間使用可能、切れ味良、コスパ最高",
                "4.0 | 26-03-12 | 刈りやすい",
            ]),
        ]),
    ]

    for cat, products in jp_comp:
        add_heading(doc, cat, level=4)
        for prod_name, reviews in products:
            add_para(doc, prod_name, bold=True, size=10, indent_cm=0.3, italic=True,
                     color=RGBColor(0x2E, 0x5C, 0x8A))
            for r in reviews:
                add_evidence(doc, r)

    # --- 6. Social-media insights ---
    add_heading(doc, "6. 社群媒體洞察 Social-Media Insights", level=2)
    add_blockquote(doc, [
        "資料來源：research/social_us/insights.md（Reddit + 美國論壇 + 部落格）、",
        "research/social_jp/insights.md（知恵袋 + 価格.com + マイベスト + note + 楽天みんなのレビュー）",
    ])

    add_heading(doc, "6.1 US 社群三大共識", level=3)
    add_bullet(doc, "Beard / Mustache：使用者已從 Wahl 系列轉向 Andis、ANGFAN 等高速無刷馬達品牌；LCD 顯示電量 + 速度檔位是「現代」電剪的標配。Braun Series 9 Pro 為敏感肌首選。")
    add_bullet(doc, "Body Groomer：Manscaped Lawn Mower 5.0 Ultra 的「無 guard 也不傷」是市場黃金標準；IPX7 + 充電底座是高端標誌。")
    add_bullet(doc, "Pet Clipper / Grinder：Oneisall 是 <60dB 入門首選；Casfuy Dog Nail Grinder 是用戶評價最高、終身保固加分。LED 照血線是革命設計。")
    add_bullet(doc, "Foil Shaver：Panasonic Arc 5/Lambdash 是 Reddit 最多人推薦的「best bang for buck」。Braun Series 9 高端首選但清洗 cartridge 成本高。")

    add_heading(doc, "6.2 JP 社群三大共識", level=3)
    add_bullet(doc, "Beard / Mustache：Panasonic ER-GB74-S 是 kakaku.com / マイベスト 雙料定番；38 段、45° 鋭利刃、防水洗為核心訴求。「充電 15-16 小時」是常見負評。")
    add_bullet(doc, "Nose / Ear：パナソニック・日立・フィリップス 三強；「ほとんど痛くない」「丸洗いできる」「IPX7」是消費者心智上的標配。")
    add_bullet(doc, "Body Groomer / VIO：パナソニック ER-GK23-A 在「肌當たり痛くない」評價最佳。VIO 周りの安全性は購入決策最大關注點。")
    add_bullet(doc, "Pet / Baby：「専門家監修」「医療クリニック推奨」「分貝數值化」是日本消費者偏好的權威佐證形式 — 具有高 esteem + safety 雙重訴求價值。")
    add_bullet(doc, "跨類別文化共通：日本市場的「身嗜み（みだしなみ）」、「父の日ギフト」、「節約意識」是貫穿所有電剪購買決策的文化背景。")

    # --- 7. Strategy ---
    add_heading(doc, "7. 策略建議 Strategic Recommendations", level=2)

    add_heading(doc, "7.1 美國市場 — 「Gift-Ready × Multi-Function Pro」", level=3)
    add_para(doc, "1. 送禮場景驅動行銷日曆", bold=True)
    add_bullet(doc, "Q2（5–6 月）父親節主題禮盒（Father's Day）", indent_cm=1.0)
    add_bullet(doc, "Q4（10–12 月）聖誕節 + 黑五 + 禮物指南", indent_cm=1.0)
    add_bullet(doc, "Listing A+ 第一屏應出現「Perfect Gift for Him」訴求；包裝設計需可直接送禮（不需重新包裝）", indent_cm=1.0)

    add_para(doc, "2. 多功能套組為核心 SKU 形態", bold=True)
    add_bullet(doc, "US 顧客最看重 num_grooming_functions (F=1193) 與 total_attachments_count (F=1015)", indent_cm=1.0)
    add_bullet(doc, "主推 5-in-1 / 7-in-1 套組，附件數標明於主圖右上角徽章", indent_cm=1.0)

    add_para(doc, "3. 競品差異化（vs Wahl / Andis / Manscaped）", bold=True)
    add_bullet(doc, "強調 Japanese OEM 工藝（50 年代工背景）→ 對應 maslow/esteem", indent_cm=1.0)
    add_bullet(doc, "LCD 電量顯示 + USB-C 為標配（社群共識）", indent_cm=1.0)
    add_bullet(doc, "IPX7 + 充電底座以對抗 Manscaped Lawn Mower 5.0 Ultra", indent_cm=1.0)

    add_para(doc, "4. Pricing & Positioning", bold=True)
    add_bullet(doc, "從 Conair（B008KEJ1LM 等）等入門品向上區隔", indent_cm=1.0)
    add_bullet(doc, "價格帶：$60–$120，高於 Conair 但低於 Braun Series 9 Pro 旗艦（$300+）", indent_cm=1.0)

    add_heading(doc, "7.2 日本市場 — 「精度・耐久・分貝数値化」", level=3)
    add_para(doc, "1. 以「具體數字」標榜技術規格", bold=True)
    add_bullet(doc, "bullet-point 1：アタッチメント X 個 / 長さ調整 X 段階（0.5mm 単位）", indent_cm=1.0)
    add_bullet(doc, "bullet-point 2：稼働時間 XX 分 / IPX7 防水 / 騒音 XX dB", indent_cm=1.0)
    add_bullet(doc, "對應 Top ANOVA 屬性 total_attachments_count (F=2630)、guide_comb_variety (F=2412)、adjustable_comb_settings (F=1827)", indent_cm=1.0)

    add_para(doc, "2. 権威佐證與比較", bold=True)
    add_bullet(doc, "「専門家監修」「サロン推奨」「医療クリニック共同開発」等表述（社群偏好）", indent_cm=1.0)
    add_bullet(doc, "與 Panasonic ER-GB74-S、Maxell IZN-C240-K 直接比較規格表（attachments, modes, blade material）", indent_cm=1.0)

    add_para(doc, "3. 電源策略雙線並行", bold=True)
    add_bullet(doc, "JP 市場 47% 鼻毛電剪用戶仍偏好乾電池款（来源：Segment 1 Top attribute 為 power_source_type）", indent_cm=1.0)
    add_bullet(doc, "不應強制全 USB-C 化，保留乾電池產品線（如 B07XTLC91J 系列）作為穩定收入", indent_cm=1.0)

    add_para(doc, "4. 語言與文化在地化", bold=True)
    add_bullet(doc, "日文 listing 措辭應導入「身嗜み（みだしなみ）」、「自然な仕上がり」、「丸洗いできる」", indent_cm=1.0)
    add_bullet(doc, "父の日（6 月第 3 個週日）禮品檔期：禮盒包裝 + のし対応", indent_cm=1.0)
    add_bullet(doc, "楽天市場直營店優勢（台灣廠商稀缺資格）— 強化「URBANER 直営正規品 + 1 年保証」訴求", indent_cm=1.0)

    add_heading(doc, "7.3 跨市場共通行動", level=3)
    add_para(doc, "1. Waterproof（IPX7 / IPX8）為兩市場共同的差異化錨：US ANOVA 排名 7、JP ANOVA 排名 9，且兩市場社群皆視為高端必備。任何新 SKU 開發應以 IPX7+ 為起點。")
    add_para(doc, "2. 附件齊全度 = 兩市場共同價值：US Segment 1（送禮）與 JP Segment 2 都把 total_attachments_count、guide_comb_variety、attachment_fitment 列為 Top 屬性。SKU 開發應以 4-piece 起跳，理想 7-piece。")
    add_para(doc, "3. Hero SKU 差異化策略：")
    add_bullet(doc, "US Hero B0FL267TCG (Beard / Mustache Trimmers) → 美國推廣資源集中此 SKU", indent_cm=1.0)
    add_bullet(doc, "JP Hero B0GBWZBMS5 (Nose / Ear Trimmers) → 日本推廣資源集中此 SKU", indent_cm=1.0)
    add_para(doc, "4. 理論行銷對映：Maslow safety（IPX7、敏感肌、医療）+ esteem（職人工藝、Japanese OEM）+ belongingness（送禮、家族）為兩市場共同情感觸發。")

    # --- 8. Theory & Theme Coverage ---
    add_heading(doc, "8. 理論與主題覆蓋摘要 Theory and Theme Coverage Summary", level=2)
    add_table(
        doc,
        header=["理論族", "已被屬性涵蓋的子主題", "未被涵蓋"],
        rows=[
            ["product_positioning", "attributes, functions, benefits, usage_context_service_experience", "—"],
            ["maslow", "physiological, safety, belongingness, esteem, self_actualization", "—"],
            ["purchase_motivation", "functional, security, relational", "—"],
            ["wom_motivation", "altruistic, social_identity, self_enhancement, emotional_expression", "—"],
        ],
        col_widths=[4.0, 9.0, 2.0],
    )
    add_para(doc, "所有 4 族與 20 主題在 US 與 JP 兩市場皆有實證評論支撐。", bold=True)

    # --- 9. Reproducibility ---
    add_heading(doc, "9. 可重現性與產出工件 Reproducibility", level=2)
    add_table(
        doc,
        header=["工件", "路徑", "用途"],
        rows=[
            ["屬性目錄", "attribute_catalog.csv", "114 屬性 + 4 理論族 + 子主題 + 範例引言"],
            ["US 顯著度表", "output_stp/market_stp_us/review_scoring_table.csv", "4,360 reviews × 114 attributes (Axis A)"],
            ["US 品質表", "output_stp/market_stp_us/product_quality_scorecard.csv", "52 products × 114 attributes (Axis B)"],
            ["US 區隔變數", "output_stp/market_stp_us/segmentation_variables.csv", "review × salience + segment label"],
            ["US Targeting ANOVA", "output_stp/market_stp_us/targeting_anova.csv", "114 attrs × F, p, segment means"],
            ["US 定位散布圖", "output_stp/market_stp_us/perceptual_map.png", "品質 PCA 二維圖"],
            ["US 區隔散布圖", "output_stp/market_stp_us/segmentation_map.png", "顯著度 PCA 二維圖"],
            ["US 屬性熱圖", "output_stp/market_stp_us/quality_heatmap.png", "Top 20 屬性 × Top 12 SKU"],
            ["JP 顯著度表", "output_stp/market_stp_jp/review_scoring_table.csv", "7,163 reviews × 114 attributes"],
            ["JP 品質表", "output_stp/market_stp_jp/product_quality_scorecard.csv", "36 products × 114 attributes"],
            ["JP 區隔變數", "output_stp/market_stp_jp/segmentation_variables.csv", "同上"],
            ["JP Targeting ANOVA", "output_stp/market_stp_jp/targeting_anova.csv", "同上"],
            ["JP 定位散布圖", "output_stp/market_stp_jp/perceptual_map.png", "同上"],
            ["JP 區隔散布圖", "output_stp/market_stp_jp/segmentation_map.png", "同上"],
            ["JP 屬性熱圖", "output_stp/market_stp_jp/quality_heatmap.png", "同上"],
            ["雙市場報告", "output_stp/market_stp_dual/dual_market_stp_report.md", "本檔來源"],
        ],
        col_widths=[3.0, 6.5, 5.5],
    )

    add_para(doc, "重新執行：", bold=True)
    add_para(doc, "python market_stp.py            # 重產 US + JP 兩市場所有工件", indent_cm=0.5)
    add_para(doc, "python build_dual_market_report.py  # 重產雙市場報告", indent_cm=0.5)

    # =========================================================
    # PART B: US 詳細報告
    # =========================================================
    doc.add_page_break()
    add_heading(doc, "Part B. US 市場詳細 STP 報告（English Source）", level=1)

    add_bullet(doc, "Reviews analysed: 4,360 URBANER own-product reviews")
    add_bullet(doc, "Products in quality scorecard: 52")
    add_bullet(doc, "Attribute catalog: 114 attributes, 20 themes, 4 theory families")
    add_bullet(doc, "Languages: en=4360")
    add_bullet(doc, "Categories covered: 001_Beard_Mustache_Trimmers, 002_Nose_Ear Hair_Trimmers, 003_Body_Groomers")

    add_heading(doc, "Attribute Extraction Summary", level=2)
    add_table(
        doc,
        header=["Field", "Value"],
        rows=[
            ["Target minimum", "30"],
            ["Actual count", "114"],
            ["Shortfall reason", "none — pre-frozen catalog of 114 attributes reused across markets"],
            ["Theory gap", "none"],
        ],
        col_widths=[4.5, 11.0],
    )
    add_para(doc, "Themes discovered (20)：", bold=True)
    add_para(doc, "Power Source, Motor & Blade, Waterproof, Functions, Adjustability, Design, Portability, Accessories, Safety, Build Quality, Ease of Use, Gifting, Price & Value, Target User, Blade Specs, Brand & Trust, Sensory Experience, Hair Capability, Social Context, Eco & Sustainability", indent_cm=0.5)

    add_para(doc, "Theory family coverage：", bold=True)
    add_bullet(doc, "product_positioning: attributes, functions, benefits, usage_context_service_experience")
    add_bullet(doc, "maslow: physiological, safety, belongingness, esteem, self_actualization")
    add_bullet(doc, "purchase_motivation: functional, security, relational")
    add_bullet(doc, "wom_motivation: altruistic, social_identity, self_enhancement, emotional_expression")

    add_para(doc, "Scoring axes：", bold=True)
    add_bullet(doc, "Axis A (Salience 0–7): per review × attribute; bilingual keyword presence + length-density weighting.")
    add_bullet(doc, "Axis B (Quality 0–10): per product × attribute; star-rating baseline adjusted by attribute-specific negative signals + sentiment delta.")

    add_heading(doc, "Segmentation", level=2)
    add_para(doc, "What this section does: Clusters 4,360 reviewers in US by their salience profile across 114 attributes.")
    add_para(doc, "Axis modelling: Axis A (salience) only.")
    add_para(doc, "Statistical methods: StandardScaler → PCA (85% variance retained) → K-Means (k=3–6, silhouette + ≥5% guardrail).")
    add_para(doc, "Theories used: product_positioning (functions, benefits, attributes); maslow (physiological, safety, esteem); purchase_motivation (functional, security).")
    add_para(doc, "Result: k = 3, silhouette = 0.358. See segmentation_map.png.", bold=True)

    # Insert US segmentation map
    add_image(doc, "output_stp/market_stp_us/segmentation_map.png",
              caption="圖 B-1：US 市場區隔散布圖（Axis A 顯著度 PCA 二維投影）")

    add_heading(doc, "Segment 1 — Gift Suitability Men Focus", level=3)
    add_bullet(doc, "n = 850 reviews (19.5%)")
    add_bullet(doc, "Avg rating = 3.68★")
    add_bullet(doc, "Dominant product = B0CGHWFF71")
    add_bullet(doc, "Category mix = Beard / Mustache Trimmers 88%, Nose / Ear Trimmers 7%, Body Groomers 6%")
    add_bullet(doc, "Language mix = en 100%")
    add_bullet(doc, "Top 5 attributes = gift_suitability_men, primary_user_gender, ear_hair_trimming_function, attachment_fitment, total_attachments_count")

    add_heading(doc, "Segment 2 — Ear Hair Trimming Function Focus", level=3)
    add_bullet(doc, "n = 3,337 reviews (76.5%)")
    add_bullet(doc, "Avg rating = 3.18★")
    add_bullet(doc, "Dominant product = B08P4HHSZT")
    add_bullet(doc, "Category mix = Beard / Mustache Trimmers 95%, Nose / Ear Trimmers 4%, Body Groomers 1%")
    add_bullet(doc, "Language mix = en 100%")
    add_bullet(doc, "Top 5 attributes = ear_hair_trimming_function, product_longevity, beard_trimming_function, gift_suitability_men, rechargeable_design")

    add_heading(doc, "Segment 3 — Power Source Type Focus", level=3)
    add_bullet(doc, "n = 173 reviews (4.0%)")
    add_bullet(doc, "Avg rating = 4.48★")
    add_bullet(doc, "Dominant product = B0FS5TCR3B")
    add_bullet(doc, "Category mix = Beard / Mustache Trimmers 76%, Body Groomers 18%, Nose / Ear Trimmers 6%")
    add_bullet(doc, "Language mix = en 100%")
    add_bullet(doc, "Top 5 attributes = power_source_type, gift_suitability_men, rechargeable_design, beard_trimming_function, ear_hair_trimming_function")

    add_finding(
        doc, "SEG-01 (US)",
        fields=[
            ("finding_statement", "The largest segment in US (segment 2, 76.5% share, n=3,337) is dominated by ear_hair_trimming_function concerns."),
            ("business_implication", "URBANER US listings should lead with messaging that addresses ear_hair_trimming_function to capture the modal customer."),
            ("axes_used", "salience"),
            ("methods_used", "StandardScaler → PCA → K-Means"),
            ("theories_used", "product_positioning, maslow"),
            ("subtheories_used", "functions, benefits, physiological, safety"),
            ("statistical_results", "silhouette=0.358, k=3, n=4360"),
            ("reproducibility", "review_scoring_table.csv → SAL_COLS → StandardScaler → PCA(0.85) → KMeans(k=3..6, silhouette+share guardrail)"),
        ],
        evidence_quotes=[
            'evidence_quote [US_R000001, 1★, B001K85BN2]: "Good awful. The only positive is that this trimmer can be used in the shower. I don\'t have a bushy beard and even so this trimmer resulted in a very uneven cut and I was constantly having to rinse the gears hairs out to make sure it kept working. Spend more and get a better trimm"',
            'evidence_quote [US_R000011, 5★, B008KEJ1LM]: "other shavers i bought had batteries didn\'t last so i figured buying a panasonic product was a safe choice as they are a major global battery manufacturer. Still going strong after several years. whenever i go on vacation i can leave the charger home not worrying the shaver will "',
        ],
    )

    add_heading(doc, "Targeting", level=2)
    add_para(doc, "What this section does: Identifies which attributes most discriminate segments via one-way ANOVA, and ranks segments by avg rating and share for prioritization.")
    add_para(doc, "Axis modelling: Axis A (salience) as the dependent variable across segment groups.")
    add_para(doc, "Statistical methods: One-way ANOVA per attribute; pairwise priority ranking by avg_rating × share.")
    add_para(doc, "Theories used: purchase_motivation (functional, security, relational); maslow (safety, esteem, self_actualization); product_positioning (functions, benefits).")
    add_para(doc, "Significant attributes (p<0.05): 99 / 114. See targeting_anova.csv.", bold=True)

    add_para(doc, "Top 10 most discriminating attributes:", bold=True)
    add_table(
        doc,
        header=["Attribute", "F", "p"],
        rows=[
            ["gift_suitability_men", "1299.469", "0"],
            ["primary_user_gender", "1196.508", "0"],
            ["num_grooming_functions", "1193.262", "0"],
            ["total_attachments_count", "1014.905", "0"],
            ["attachment_fitment", "975.252", "0"],
            ["power_source_type", "941.026", "0"],
            ["waterproof_rating_ipx8", "902.847", "0"],
            ["guide_comb_variety", "902.265", "0"],
            ["multi_use_versatility_score", "778.884", "0"],
            ["quick_touch_up_design", "708.689", "0"],
        ],
    )

    add_bullet(doc, "Priority segments: Segment 3", bold=True)
    add_bullet(doc, "Secondary segments: Segment 1", bold=True)
    add_bullet(doc, "Deprioritized: Segment 2", bold=True)

    add_finding(
        doc, "TGT-01 (US)",
        fields=[
            ("finding_statement", "gift_suitability_men is the strongest segment-discriminating attribute in US (F=1299.469, p=0.0)."),
            ("business_implication", "URBANER US ad creative and Amazon listing copy should foreground gift_suitability_men because it is the axis on which this market's segments diverge most clearly."),
            ("axes_used", "salience"),
            ("methods_used", "ANOVA"),
            ("theories_used", "product_positioning, purchase_motivation"),
            ("subtheories_used", "functions, benefits, functional"),
            ("reproducibility", "segmentation_variables.csv → groupby(segment) → scipy.stats.f_oneway per salience column → sort by F."),
        ],
        evidence_quotes=[
            'evidence_quote [US_R000003, 1★, B001K85BN2]: "This was apparantly an export model. The in structions were in Japanese. This purchase was intended to replace a Panasonic ER2430. The unit that was shipped ER2403P has a different configeration for the heighth adjustment. My old unit had settings of 3, 4,5,6,, and so on. This"',
            'evidence_quote [US_R000010, 5★, B008KEJ1LM]: "Excelente tengo años usandolo y hasta hora no me a fallado lo volveria a comprar"',
        ],
    )

    add_heading(doc, "Positioning", level=2)
    add_para(doc, "What this section does: Maps URBANER products on a quality perceptual space and computes ideal-point distances.")
    add_para(doc, "Axis modelling: Axis B (quality 0–10) per product × attribute.")
    add_para(doc, "Statistical methods: Standardised quality matrix → PCA (2 components); ideal-point distance (RMS distance from 10.0 vector).")
    add_para(doc, "Theories used: product_positioning (attributes, functions, benefits); maslow (esteem, safety, self_actualization); purchase_motivation (functional, security).")
    add_para(doc, "PCA variance explained: PC1=51.1%, PC2=4.5%. See perceptual_map.png and quality_heatmap.png.", bold=True)

    add_image(doc, "output_stp/market_stp_us/perceptual_map.png",
              caption="圖 B-2：US 品質感知圖（Axis B 品質 PCA 二維投影 + 理想點）")

    add_image(doc, "output_stp/market_stp_us/quality_heatmap.png",
              caption="圖 B-3：US 屬性品質熱圖（Top 20 屬性 × Top 12 SKU）")

    add_para(doc, "Top 10 URBANER products closest to ideal (lowest RMS distance from quality=10)：", bold=True)
    add_table(
        doc,
        header=["Rank", "Product", "Distance", "n_reviews", "Category"],
        rows=[
            ["1", "B0FL267TCG", "1.601", "29", "Beard / Mustache Trimmers"],
            ["2", "B0FX4NX6HM", "1.761", "8", "Body Groomers"],
            ["3", "B0CDNQ62ML", "1.766", "8", "Body Groomers"],
            ["4", "B0GLJRLS5G", "1.814", "99", "Beard / Mustache Trimmers"],
            ["5", "B0FHWWPLCP", "1.867", "8", "Body Groomers"],
            ["6", "B0G1Z1QKJ8", "1.896", "50", "Beard / Mustache Trimmers"],
            ["7", "B0FS5TCR3B", "1.912", "158", "Beard / Mustache Trimmers"],
            ["8", "B092H8XHDG", "2.026", "8", "Nose / Ear Trimmers"],
            ["9", "B0F66FZHY3", "2.029", "8", "Nose / Ear Trimmers"],
            ["10", "B0F2RSTDSJ", "2.065", "73", "Beard / Mustache Trimmers"],
        ],
    )

    add_para(doc, "Bottom 5 furthest from ideal：", bold=True)
    add_table(
        doc,
        header=["Product", "Distance", "n_reviews", "Category"],
        rows=[
            ["B07FFGYWJ6", "6.187", "6", "Nose / Ear Trimmers"],
            ["B001K85BN2", "6.106", "4", "Beard / Mustache Trimmers"],
            ["B09XVP6YPM", "4.845", "341", "Beard / Mustache Trimmers"],
            ["B0F5VJD5DQ", "4.802", "77", "Nose / Ear Trimmers"],
            ["B08P4HHSZT", "4.783", "500", "Beard / Mustache Trimmers"],
        ],
    )

    add_finding(
        doc, "POS-01 (US)",
        fields=[
            ("finding_statement", "Product B0FL267TCG is the US portfolio leader on the quality perceptual map (RMS distance to ideal = 1.601)."),
            ("business_implication", "Treat B0FL267TCG as the US hero SKU — its attribute strengths should anchor brand-level messaging and any cross-sell bundles."),
            ("axes_used", "quality"),
            ("methods_used", "PCA + ideal-point RMS distance"),
            ("theories_used", "product_positioning, maslow, purchase_motivation"),
            ("subtheories_used", "attributes, functions, benefits, esteem, security, functional"),
        ],
    )

    add_heading(doc, "Theory and Theme Coverage Summary (US)", level=2)
    add_table(
        doc,
        header=["Theory family", "Subtheories evidenced", "Subtheories not evidenced"],
        rows=[
            ["product_positioning", "attributes, functions, benefits, usage_context_service_experience", "—"],
            ["maslow", "physiological, safety, belongingness, esteem, self_actualization", "—"],
            ["purchase_motivation", "functional, security, relational", "—"],
            ["wom_motivation", "altruistic, social_identity, self_enhancement, emotional_expression", "—"],
        ],
        col_widths=[4.0, 9.0, 2.5],
    )

    # =========================================================
    # PART C: JP 詳細報告
    # =========================================================
    doc.add_page_break()
    add_heading(doc, "Part C. JP 市場詳細 STP 報告（English Source）", level=1)

    add_bullet(doc, "Reviews analysed: 7,163 URBANER own-product reviews")
    add_bullet(doc, "Products in quality scorecard: 36")
    add_bullet(doc, "Attribute catalog: 114 attributes, 20 themes, 4 theory families")
    add_bullet(doc, "Languages: ja=6924, en=239")
    add_bullet(doc, "Categories covered: 002_Nose_Ear Hair_Trimmers, 001_Beard_Mustache_Trimmers")

    add_heading(doc, "Attribute Extraction Summary", level=2)
    add_table(
        doc,
        header=["Field", "Value"],
        rows=[
            ["Target minimum", "30"],
            ["Actual count", "114"],
            ["Shortfall reason", "none — pre-frozen catalog of 114 attributes reused across markets"],
            ["Theory gap", "none"],
        ],
        col_widths=[4.5, 11.0],
    )
    add_para(doc, "Themes discovered (20)：", bold=True)
    add_para(doc, "Power Source, Motor & Blade, Waterproof, Functions, Adjustability, Design, Portability, Accessories, Safety, Build Quality, Ease of Use, Gifting, Price & Value, Target User, Blade Specs, Brand & Trust, Sensory Experience, Hair Capability, Social Context, Eco & Sustainability", indent_cm=0.5)

    add_para(doc, "Theory family coverage：", bold=True)
    add_bullet(doc, "product_positioning: attributes, functions, benefits, usage_context_service_experience")
    add_bullet(doc, "maslow: physiological, safety, belongingness, esteem, self_actualization")
    add_bullet(doc, "purchase_motivation: functional, security, relational")
    add_bullet(doc, "wom_motivation: altruistic, social_identity, self_enhancement, emotional_expression")

    add_para(doc, "Scoring axes：", bold=True)
    add_bullet(doc, "Axis A (Salience 0–7): per review × attribute; bilingual keyword presence + length-density weighting.")
    add_bullet(doc, "Axis B (Quality 0–10): per product × attribute; star-rating baseline adjusted by attribute-specific negative signals + sentiment delta.")

    add_heading(doc, "Segmentation", level=2)
    add_para(doc, "What this section does: Clusters 7,163 reviewers in JP by their salience profile across 114 attributes.")
    add_para(doc, "Axis modelling: Axis A (salience) only.")
    add_para(doc, "Statistical methods: StandardScaler → PCA (85% variance retained) → K-Means (k=3–6, silhouette + ≥5% guardrail).")
    add_para(doc, "Theories used: product_positioning (functions, benefits, attributes); maslow (physiological, safety, esteem); purchase_motivation (functional, security).")
    add_para(doc, "Result: k = 3, silhouette = 0.57. See segmentation_map.png.", bold=True)

    add_image(doc, "output_stp/market_stp_jp/segmentation_map.png",
              caption="圖 C-1：JP 市場區隔散布圖（Axis A 顯著度 PCA 二維投影）")

    add_heading(doc, "Segment 1 — Power Source Type Focus", level=3)
    add_bullet(doc, "n = 6,559 reviews (91.6%)")
    add_bullet(doc, "Avg rating = 3.46★")
    add_bullet(doc, "Dominant product = B07XTLC91J")
    add_bullet(doc, "Category mix = Nose / Ear Trimmers 59%, Beard / Mustache Trimmers 41%")
    add_bullet(doc, "Language mix = ja 97%, en 3%")
    add_bullet(doc, "Top 5 attributes = power_source_type, nose_hair_trimming_function, price_value_ratio, ease_of_use_overall, product_longevity")

    add_heading(doc, "Segment 2 — Adjustable Comb Settings Focus", level=3)
    add_bullet(doc, "n = 561 reviews (7.8%)")
    add_bullet(doc, "Avg rating = 4.01★")
    add_bullet(doc, "Dominant product = B07CYZH2XC")
    add_bullet(doc, "Category mix = Beard / Mustache Trimmers 75%, Nose / Ear Trimmers 25%")
    add_bullet(doc, "Language mix = ja 97%, en 3%")
    add_bullet(doc, "Top 5 attributes = adjustable_comb_settings, power_source_type, product_dimensions_size, total_attachments_count, attachment_fitment")

    add_heading(doc, "Segment 3 — Operation Vibration Feel Focus", level=3)
    add_bullet(doc, "n = 43 reviews (0.6%)")
    add_bullet(doc, "Avg rating = 3.05★")
    add_bullet(doc, "Dominant product = B07FFGYWJ6")
    add_bullet(doc, "Category mix = Nose / Ear Trimmers 63%, Beard / Mustache Trimmers 37%")
    add_bullet(doc, "Language mix = ja 98%, en 2%")
    add_bullet(doc, "Top 5 attributes = operation_vibration_feel, low_vibration_design, motor_power_rpm, power_source_type, low_noise_operation")

    add_finding(
        doc, "SEG-01 (JP)",
        fields=[
            ("finding_statement", "The largest segment in JP (segment 1, 91.6% share, n=6,559) is dominated by power_source_type concerns."),
            ("business_implication", "URBANER JP listings should lead with messaging that addresses power_source_type to capture the modal customer."),
            ("axes_used", "salience"),
            ("methods_used", "StandardScaler → PCA → K-Means"),
            ("theories_used", "product_positioning, maslow"),
            ("subtheories_used", "functions, benefits, physiological, safety"),
            ("statistical_results", "silhouette=0.57, k=3, n=7163"),
            ("reproducibility", "review_scoring_table.csv → SAL_COLS → StandardScaler → PCA(0.85) → KMeans(k=3..6, silhouette+share guardrail)"),
        ],
        evidence_quotes=[
            'evidence_quote [JP_R000010, 5★, B001K85BN2]: "よく剃れるし問題ありません。以前は充電式を使っていましたが、電池が劣化すると捨てるしかなくなるのでこういった製品は乾電池式一択です。"',
            'evidence_quote [JP_R000035, 2★, B001K85BN2]: "Not satisfied about the battery life. When the battery drains the speed gets low gradually. We have to use new batteries always for better functionality. So I stopped using this."',
        ],
    )

    add_heading(doc, "Targeting", level=2)
    add_para(doc, "What this section does: Identifies which attributes most discriminate segments via one-way ANOVA, and ranks segments by avg rating and share for prioritization.")
    add_para(doc, "Axis modelling: Axis A (salience) as the dependent variable across segment groups.")
    add_para(doc, "Statistical methods: One-way ANOVA per attribute; pairwise priority ranking by avg_rating × share.")
    add_para(doc, "Theories used: purchase_motivation (functional, security, relational); maslow (safety, esteem, self_actualization); product_positioning (functions, benefits).")
    add_para(doc, "Significant attributes (p<0.05): 58 / 114. See targeting_anova.csv.", bold=True)

    add_para(doc, "Top 10 most discriminating attributes:", bold=True)
    add_table(
        doc,
        header=["Attribute", "F", "p"],
        rows=[
            ["total_attachments_count", "2630.384", "0"],
            ["guide_comb_variety", "2412.195", "0"],
            ["adjustable_comb_settings", "1826.965", "0"],
            ["attachment_fitment", "1723.208", "0"],
            ["product_dimensions_size", "1004.693", "0"],
            ["power_source_type", "543.117", "0"],
            ["rechargeable_design", "497.535", "0"],
            ["beard_trimming_function", "494.451", "0"],
            ["waterproof_rating_ipx8", "460.933", "0"],
            ["battery_independence", "346.662", "0"],
        ],
    )

    add_bullet(doc, "Priority segments: Segment 2", bold=True)
    add_bullet(doc, "Secondary segments: Segment 1", bold=True)
    add_bullet(doc, "Deprioritized: Segment 3", bold=True)

    add_finding(
        doc, "TGT-01 (JP)",
        fields=[
            ("finding_statement", "total_attachments_count is the strongest segment-discriminating attribute in JP (F=2630.384, p=0.0)."),
            ("business_implication", "URBANER JP ad creative and Amazon listing copy should foreground total_attachments_count because it is the axis on which this market's segments diverge most clearly."),
            ("axes_used", "salience"),
            ("methods_used", "ANOVA"),
            ("theories_used", "product_positioning, purchase_motivation"),
            ("subtheories_used", "functions, benefits, functional"),
            ("reproducibility", "segmentation_variables.csv → groupby(segment) → scipy.stats.f_oneway per salience column → sort by F."),
        ],
        evidence_quotes=[
            'evidence_quote [JP_R000145, 1★, B001K85BN2]: "前のモデルのアタッチメントが壊れた（本体はまだ使える）ので進化を確かめるべく購入。結果はかっこよくなっただけ。性能はむしろ落ちたかも。パワーが落ちたのは他の人も触れているが単3から単4に変化。トリム機構は変わって無いから必要十分でスペックダウンしたわけでは無さそう。新しい方の3㎜は本当に3㎜でそろう。前のは幾分短い刈高。なので長く感じる。また完了までの時間が掛かる様になった。デザインがかっこよくなり機能を落として本質が見えていない。おそらく開発側は変える必要を感じていなかったんじゃないか？　ただヒゲトリマーは怪しいメーカーか中華品がはびこっていて、3㎜以"',
            'evidence_quote [JP_R000181, 4★, B001K85BN2]: "良い商品だとは思いますがちょっと髭を短く残したいみたいな使い方をしたかった自分にとってはアタッチメントを1番短くセットしても髭に届かず意味のないものになってしまいました。代わりに眉毛を整えるのに使ってます。"',
        ],
    )

    add_heading(doc, "Positioning", level=2)
    add_para(doc, "What this section does: Maps URBANER products on a quality perceptual space and computes ideal-point distances.")
    add_para(doc, "Axis modelling: Axis B (quality 0–10) per product × attribute.")
    add_para(doc, "Statistical methods: Standardised quality matrix → PCA (2 components); ideal-point distance (RMS distance from 10.0 vector).")
    add_para(doc, "Theories used: product_positioning (attributes, functions, benefits); maslow (esteem, safety, self_actualization); purchase_motivation (functional, security).")
    add_para(doc, "PCA variance explained: PC1=29.5%, PC2=7.5%. See perceptual_map.png and quality_heatmap.png.", bold=True)

    add_image(doc, "output_stp/market_stp_jp/perceptual_map.png",
              caption="圖 C-2：JP 品質感知圖（Axis B 品質 PCA 二維投影 + 理想點）")

    add_image(doc, "output_stp/market_stp_jp/quality_heatmap.png",
              caption="圖 C-3：JP 屬性品質熱圖（Top 20 屬性 × Top 12 SKU）")

    add_para(doc, "Top 10 URBANER products closest to ideal (lowest RMS distance from quality=10)：", bold=True)
    add_table(
        doc,
        header=["Rank", "Product", "Distance", "n_reviews", "Category"],
        rows=[
            ["1", "B0GBWZBMS5", "1.97", "61", "Nose / Ear Trimmers"],
            ["2", "B0742G961R", "2.23", "40", "Beard / Mustache Trimmers"],
            ["3", "B0F99F11BW", "2.316", "142", "Nose / Ear Trimmers"],
            ["4", "B0F5VJD5DQ", "2.537", "150", "Nose / Ear Trimmers"],
            ["5", "B0BL2YWH3N", "2.657", "43", "Beard / Mustache Trimmers"],
            ["6", "B0FJS3QR1W", "2.817", "141", "Beard / Mustache Trimmers"],
            ["7", "B0FJS3TF3V", "2.817", "141", "Beard / Mustache Trimmers"],
            ["8", "B016JNKVHI", "2.918", "215", "Beard / Mustache Trimmers"],
            ["9", "B0BY1KRJZ1", "2.953", "115", "Beard / Mustache Trimmers"],
            ["10", "B0FHKR48NR", "3.113", "16", "Nose / Ear Trimmers"],
        ],
    )

    add_para(doc, "Bottom 5 furthest from ideal：", bold=True)
    add_table(
        doc,
        header=["Product", "Distance", "n_reviews", "Category"],
        rows=[
            ["B07CYQVK16", "4.966", "16", "Beard / Mustache Trimmers"],
            ["B07XTLC91J", "4.739", "500", "Nose / Ear Trimmers"],
            ["B0DKJVC5B2", "4.699", "393", "Nose / Ear Trimmers"],
            ["B08LGGMP5W", "4.658", "300", "Beard / Mustache Trimmers"],
            ["B001K85BN2", "4.587", "394", "Beard / Mustache Trimmers"],
        ],
    )

    add_finding(
        doc, "POS-01 (JP)",
        fields=[
            ("finding_statement", "Product B0GBWZBMS5 is the JP portfolio leader on the quality perceptual map (RMS distance to ideal = 1.97)."),
            ("business_implication", "Treat B0GBWZBMS5 as the JP hero SKU — its attribute strengths should anchor brand-level messaging and any cross-sell bundles."),
            ("axes_used", "quality"),
            ("methods_used", "PCA + ideal-point RMS distance"),
            ("theories_used", "product_positioning, maslow, purchase_motivation"),
            ("subtheories_used", "attributes, functions, benefits, esteem, security, functional"),
        ],
    )

    add_heading(doc, "Theory and Theme Coverage Summary (JP)", level=2)
    add_table(
        doc,
        header=["Theory family", "Subtheories evidenced", "Subtheories not evidenced"],
        rows=[
            ["product_positioning", "attributes, functions, benefits, usage_context_service_experience", "—"],
            ["maslow", "physiological, safety, belongingness, esteem, self_actualization", "—"],
            ["purchase_motivation", "functional, security, relational", "—"],
            ["wom_motivation", "altruistic, social_identity, self_enhancement, emotional_expression", "—"],
        ],
        col_widths=[4.0, 9.0, 2.5],
    )

    # ===== Footer =====
    doc.add_paragraph()
    add_separator(doc)
    add_para(doc, "資料截止：2026-05-05", italic=True, color=RGBColor(0x77, 0x77, 0x77),
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "作者：FourSight Lab × URBANER 產學合作專案", italic=True,
             color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "方法依循 review-mining-stp 0.1 工作流契約", italic=True,
             color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "— 報告結束 —", italic=True, color=RGBColor(0x99, 0x99, 0x99),
             align=WD_ALIGN_PARAGRAPH.CENTER)

    out_path = "output_stp/URBANER_STP_完整報告.docx"
    doc.save(out_path)
    print(f"OK: saved {out_path}")


if __name__ == "__main__":
    main()
